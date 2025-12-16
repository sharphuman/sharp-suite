"""Sharp JD - Professional Job Description Writer with Cross-App Auth"""
import streamlit as st
import os
import requests
from datetime import datetime, timedelta
import secrets
import json
import re
import io
import subprocess
import tempfile

SUPABASE_URL = "https://qkjtprqgblnfftrotyks.supabase.co"
SUPABASE_ANON_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InFranRwcnFnYmxuZmZ0cm90eWtzIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NjUzNTgzNDAsImV4cCI6MjA4MDkzNDM0MH0.pVzSq4M5i58zBGl7OPDhNL9qYBcg-bz8MVrBI5MQSkw"
SUPABASE_SERVICE_KEY = os.environ.get("SUPABASE_SERVICE_KEY", SUPABASE_ANON_KEY)
GOD_PASSWORD = "G0DHum@n101!!!"
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")

APP_URLS = {
    "portal": "https://demo.sharphuman.com", "jd": "https://jd.sharphuman.com",
    "screen": "https://screen.sharphuman.com", "interview": "https://interview.sharphuman.com",
    "outreach": "https://outreach.sharphuman.com", "content": "https://content.sharphuman.com",
    "sales": "https://sales.sharphuman.com",
}

# ============================================
# AUTH FUNCTIONS
# ============================================

def create_session(user_id, email):
    token = secrets.token_urlsafe(32)
    try:
        requests.post(f"{SUPABASE_URL}/rest/v1/sessions",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                     "Content-Type": "application/json", "Prefer": "return=minimal"},
            json={"user_id": user_id, "token": token, "ip_address": "unknown", "device_hash": "jd",
                  "is_active": True, "expires_at": (datetime.utcnow() + timedelta(hours=24)).isoformat()}, timeout=10)
    except: pass
    return token

def supabase_sign_in(email, password):
    try:
        r = requests.post(f"{SUPABASE_URL}/auth/v1/token?grant_type=password",
            headers={"apikey": SUPABASE_ANON_KEY, "Content-Type": "application/json"},
            json={"email": email, "password": password}, timeout=10)
        data = r.json()
        if r.status_code == 200 and data.get("access_token"):
            user = data.get("user", {})
            return {"success": True, "user": user, "session_token": create_session(user.get("id"), email)}
        return {"success": False, "message": data.get("error_description") or "Invalid credentials"}
    except Exception as e:
        return {"success": False, "message": str(e)}

def supabase_sign_up(email, password):
    try:
        r = requests.post(f"{SUPABASE_URL}/auth/v1/signup",
            headers={"apikey": SUPABASE_ANON_KEY, "Content-Type": "application/json"},
            json={"email": email, "password": password}, timeout=10)
        return {"success": True, "message": "Check email!"}
    except Exception as e:
        return {"success": False, "message": str(e)}

def validate_session_token(token):
    try:
        r = requests.get(f"{SUPABASE_URL}/rest/v1/sessions?token=eq.{token}&is_active=eq.true",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}"}, timeout=10)
        if r.status_code == 200 and r.json():
            session = r.json()[0]
            expires = datetime.fromisoformat(session["expires_at"].replace("Z", "+00:00"))
            if expires.replace(tzinfo=None) > datetime.utcnow():
                ur = requests.get(f"{SUPABASE_URL}/rest/v1/user_profiles?id=eq.{session['user_id']}",
                    headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}"}, timeout=10)
                if ur.status_code == 200 and ur.json():
                    p = ur.json()[0]
                    return {"user_id": session["user_id"], "email": p.get("email"), "plan": p.get("plan", "free"), "token": token}
    except: pass
    return None

def log_usage(user_id, session_id, app, action, tokens_used=0):
    try:
        requests.post(f"{SUPABASE_URL}/rest/v1/usage_logs",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                     "Content-Type": "application/json", "Prefer": "return=minimal"},
            json={"user_id": user_id, "session_id": session_id, "app": app, "action": action, "tokens_used": tokens_used}, timeout=5)
    except: pass

def init_session():
    defaults = [
        ('authenticated', False), ('user', None), ('is_god', False), ('session_token', None),
        ('user_plan', 'free'), ('generated_jd', None), ('jd_metadata', None), ('working_on', None),
        ('requirements_text', ''), ('enhanced_jd', None), ('copy_to_create', False),
        ('jd_versions', []), ('current_version', 0), ('rewrite_prompt', ''),
    ]
    for k, v in defaults:
        if k not in st.session_state:
            st.session_state[k] = v

def check_url_auth():
    token = st.query_params.get("token")
    if token and not st.session_state.authenticated:
        user_info = validate_session_token(token)
        if user_info:
            st.session_state.authenticated = True
            st.session_state.user = {"email": user_info["email"], "id": user_info["user_id"]}
            st.session_state.session_token = token
            st.session_state.user_plan = user_info.get("plan", "free")
            st.session_state.is_god = user_info.get("plan") == "god"
            return True
    return False

def get_user_email():
    return st.session_state.user.get("email", "User") if st.session_state.user else ("GOD" if st.session_state.is_god else "User")

def build_app_url(app_name):
    base = APP_URLS.get(app_name, "")
    token = st.session_state.get("session_token", "")
    return f"{base}?token={token}" if base and token else base

# ============================================
# FILE PROCESSING
# ============================================

def extract_text_from_file(uploaded_file):
    file_type = uploaded_file.name.split('.')[-1].lower()
    content = uploaded_file.read()
    
    if file_type == 'txt':
        return content.decode('utf-8', errors='ignore')
    elif file_type == 'json':
        try:
            data = json.loads(content.decode('utf-8'))
            if isinstance(data, dict):
                for key in ['description', 'job_description', 'content', 'text']:
                    if key in data:
                        return data[key]
                return json.dumps(data, indent=2)
            return str(data)
        except:
            return content.decode('utf-8', errors='ignore')
    elif file_type == 'pdf':
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        except:
            return "[PDF extraction failed - install pdfplumber]"
    elif file_type == 'docx':
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs)
        except:
            return "[DOCX extraction failed - install python-docx]"
    return "[Unsupported file type]"

# ============================================
# AI FUNCTIONS
# ============================================

def call_claude(prompt, max_tokens=4000, action="generate"):
    if not ANTHROPIC_API_KEY:
        return "Error: ANTHROPIC_API_KEY not set", 0
    
    try:
        r = requests.post("https://api.anthropic.com/v1/messages",
            headers={"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01", "content-type": "application/json"},
            json={"model": "claude-sonnet-4-20250514", "max_tokens": max_tokens,
                  "messages": [{"role": "user", "content": prompt}]}, timeout=120)
        
        if r.status_code == 200:
            data = r.json()
            text = data.get("content", [{}])[0].get("text", "")
            tokens = data.get("usage", {}).get("output_tokens", 0)
            
            # Log usage
            user_id = st.session_state.user.get("id") if st.session_state.user else None
            if user_id:
                log_usage(user_id, st.session_state.session_token, "jd", action, tokens)
            
            return text, tokens
        return f"Error: {r.status_code} - {r.text}", 0
    except Exception as e:
        return f"Error: {str(e)}", 0

def generate_jd_prompt(job_title, company, experience_level, industry, remote_type, tone, user_type, requirements, options):
    """Generate prompt for clean, human-written JD output"""
    
    user_context = "an external recruiter at a staffing agency" if "Recruiter" in user_type else "an internal HR/TA professional"
    
    sections = []
    if options.get('about_company'): sections.append("About [Company Name] (2-3 engaging sentences)")
    if options.get('about_role'): sections.append("About The Role (brief overview)")
    sections.append("Responsibilities (5-7 key responsibilities)")
    sections.append("Requirements - Must Have (essential qualifications)")
    sections.append("Requirements - Nice to Have (preferred qualifications)")
    if options.get('salary'): sections.append("Compensation (salary range and structure)")
    if options.get('benefits'): sections.append("Benefits & Perks")
    if options.get('culture'): sections.append("Our Culture / Why Join Us")
    if options.get('diversity'): sections.append("Equal Opportunity Statement")
    sections.append("How to Apply")
    
    return f"""Write a compelling, human-sounding job description. Write naturally like an experienced HR professional would - NO markdown formatting, NO asterisks, NO hashtags, NO bullet symbols like • or -.

CRITICAL FORMATTING RULES:
1. Use plain text only - no markdown, no special characters
2. For lists, start each item on a new line with a simple dash followed by a space
3. Section headers should be the title followed by a colon on its own line
4. Write in a warm, engaging tone that attracts candidates
5. Sound human - vary sentence length, use active voice, be specific

Context: I am {user_context} writing this JD.

Job Details:
- Title: {job_title}
- Company: {company or "A growing company"}
- Level: {experience_level}
- Industry: {industry}
- Work Type: {remote_type}
- Desired Tone: {tone}

Requirements/Context Provided:
{requirements or "Standard for this role"}

Include these sections:
{chr(10).join(f"- {s}" for s in sections)}

Target length: ~{options.get('word_target', 500)} words

Write the job description now, making it sound like a real person wrote it - not AI. Be specific, engaging, and professional."""

def rewrite_jd_prompt(current_jd, user_instructions):
    """Prompt for iterative rewriting"""
    return f"""Rewrite this job description based on the user's feedback. Maintain the same clean formatting (no markdown, no asterisks, plain text only).

CURRENT JOB DESCRIPTION:
{current_jd}

USER'S REQUESTED CHANGES:
{user_instructions}

Rewrite the job description incorporating these changes. Keep the same structure unless asked to change it. Output only the revised job description, nothing else."""

# ============================================
# EXPORT FUNCTIONS
# ============================================

def generate_docx(jd_text, metadata):
    """Generate professional DOCX using docx-js via Node"""
    job_title = metadata.get('job_title', 'Job Description')
    company = metadata.get('company', '')
    
    # Parse JD into sections
    sections = parse_jd_sections(jd_text)
    
    js_code = f'''
const {{ Document, Packer, Paragraph, TextRun, AlignmentType, HeadingLevel, LevelFormat }} = require('docx');
const fs = require('fs');

const doc = new Document({{
    styles: {{
        default: {{ document: {{ run: {{ font: "Arial", size: 22 }} }} }},
        paragraphStyles: [
            {{ id: "Title", name: "Title", basedOn: "Normal",
              run: {{ size: 48, bold: true, color: "1a1a1a", font: "Arial" }},
              paragraph: {{ spacing: {{ after: 200 }}, alignment: AlignmentType.CENTER }} }},
            {{ id: "Heading1", name: "Heading 1", basedOn: "Normal", next: "Normal", quickFormat: true,
              run: {{ size: 28, bold: true, color: "2d2d2d", font: "Arial" }},
              paragraph: {{ spacing: {{ before: 300, after: 120 }} }} }},
            {{ id: "Company", name: "Company", basedOn: "Normal",
              run: {{ size: 24, color: "db2777", font: "Arial" }},
              paragraph: {{ alignment: AlignmentType.CENTER, spacing: {{ after: 400 }} }} }}
        ]
    }},
    numbering: {{
        config: [
            {{ reference: "bullet-list",
              levels: [{{ level: 0, format: LevelFormat.BULLET, text: "•", alignment: AlignmentType.LEFT,
                style: {{ paragraph: {{ indent: {{ left: 720, hanging: 360 }} }} }} }}] }}
        ]
    }},
    sections: [{{
        properties: {{
            page: {{ margin: {{ top: 1440, right: 1440, bottom: 1440, left: 1440 }} }}
        }},
        children: [
            new Paragraph({{ heading: HeadingLevel.TITLE, children: [new TextRun("{job_title.replace('"', '\\"')}")] }}),
            {"new Paragraph({ style: 'Company', children: [new TextRun('" + company.replace('"', '\\"') + "')] })," if company else ""}
'''
    
    # Add sections
    for section_title, content in sections:
        js_code += f'''
            new Paragraph({{ heading: HeadingLevel.HEADING_1, children: [new TextRun("{section_title.replace('"', '\\"')}")] }}),
'''
        # Check if content is list-like
        lines = [l.strip() for l in content.split('\n') if l.strip()]
        for line in lines:
            if line.startswith('-') or line.startswith('•'):
                clean_line = line.lstrip('-•').strip().replace('"', '\\"').replace('\n', ' ')
                js_code += f'''
            new Paragraph({{ numbering: {{ reference: "bullet-list", level: 0 }}, children: [new TextRun("{clean_line}")] }}),
'''
            else:
                clean_line = line.replace('"', '\\"').replace('\n', ' ')
                js_code += f'''
            new Paragraph({{ children: [new TextRun("{clean_line}")] }}),
'''
    
    js_code += '''
        ]
    }]
});

Packer.toBuffer(doc).then(buffer => {
    fs.writeFileSync("/tmp/jd_output.docx", buffer);
    console.log("DOCX created successfully");
}).catch(err => {
    console.error("Error:", err);
    process.exit(1);
});
'''
    
    try:
        # Write and execute JS
        with open("/tmp/generate_docx.js", "w") as f:
            f.write(js_code)
        
        result = subprocess.run(["node", "/tmp/generate_docx.js"], capture_output=True, text=True, timeout=30)
        
        if result.returncode == 0 and os.path.exists("/tmp/jd_output.docx"):
            with open("/tmp/jd_output.docx", "rb") as f:
                return f.read()
        else:
            return None
    except Exception as e:
        return None

def generate_pdf(jd_text, metadata):
    """Generate professional PDF using reportlab"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=72, bottomMargin=72, leftMargin=72, rightMargin=72)
        
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle('CustomTitle', parent=styles['Title'], fontSize=24, textColor=HexColor('#1a1a1a'), 
                                     alignment=TA_CENTER, spaceAfter=6, fontName='Helvetica-Bold')
        company_style = ParagraphStyle('Company', parent=styles['Normal'], fontSize=14, textColor=HexColor('#db2777'),
                                       alignment=TA_CENTER, spaceAfter=30)
        heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading1'], fontSize=14, textColor=HexColor('#2d2d2d'),
                                       spaceBefore=20, spaceAfter=10, fontName='Helvetica-Bold')
        body_style = ParagraphStyle('CustomBody', parent=styles['Normal'], fontSize=11, leading=16, 
                                    textColor=HexColor('#374151'), spaceAfter=8)
        bullet_style = ParagraphStyle('Bullet', parent=body_style, leftIndent=20, bulletIndent=10)
        
        story = []
        
        # Title
        job_title = metadata.get('job_title', 'Job Description')
        story.append(Paragraph(job_title, title_style))
        
        # Company
        if metadata.get('company'):
            story.append(Paragraph(metadata['company'], company_style))
        else:
            story.append(Spacer(1, 20))
        
        # Parse and add sections
        sections = parse_jd_sections(jd_text)
        
        for section_title, content in sections:
            story.append(Paragraph(section_title, heading_style))
            
            lines = [l.strip() for l in content.split('\n') if l.strip()]
            for line in lines:
                if line.startswith('-') or line.startswith('•'):
                    clean_line = line.lstrip('-•').strip()
                    story.append(Paragraph(f"• {clean_line}", bullet_style))
                else:
                    story.append(Paragraph(line, body_style))
        
        # Footer
        story.append(Spacer(1, 30))
        footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=9, textColor=HexColor('#9ca3af'), alignment=TA_CENTER)
        story.append(Paragraph(f"Generated by Sharp JD | {datetime.now().strftime('%B %d, %Y')}", footer_style))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        return None

def generate_ats_json(jd_text, metadata):
    """Generate ATS-compatible JSON"""
    sections = parse_jd_sections(jd_text)
    
    # Extract responsibilities and requirements
    responsibilities = []
    must_have = []
    nice_to_have = []
    
    for title, content in sections:
        title_lower = title.lower()
        lines = [l.lstrip('-•').strip() for l in content.split('\n') if l.strip() and (l.strip().startswith('-') or l.strip().startswith('•'))]
        
        if 'responsibilit' in title_lower:
            responsibilities = lines
        elif 'must have' in title_lower or 'required' in title_lower:
            must_have = lines
        elif 'nice to have' in title_lower or 'preferred' in title_lower:
            nice_to_have = lines
    
    return json.dumps({
        "job_posting": {
            "title": metadata.get('job_title', ''),
            "company": metadata.get('company', ''),
            "department": "",
            "location": metadata.get('remote_type', ''),
            "employment_type": "Full-time",
            "experience_level": metadata.get('experience_level', ''),
            "industry": metadata.get('industry', ''),
            "posted_date": datetime.now().strftime('%Y-%m-%d'),
        },
        "description": {
            "full_text": jd_text,
            "responsibilities": responsibilities,
            "requirements": {
                "must_have": must_have,
                "nice_to_have": nice_to_have
            }
        },
        "metadata": {
            "source": "Sharp JD",
            "generated_at": datetime.now().isoformat(),
            "version": "1.0"
        }
    }, indent=2)

def generate_linkedin_text(jd_text, metadata):
    """Generate LinkedIn-optimized job post"""
    prompt = f"""Convert this job description into a LinkedIn job post format. Make it:
1. Engaging and scroll-stopping opening line
2. Use emojis sparingly but effectively (🚀 💼 ✨ etc.)
3. Shorter paragraphs optimized for mobile reading
4. Include relevant hashtags at the end
5. Add a clear call-to-action
6. Keep under 3000 characters

Job Title: {metadata.get('job_title', '')}
Company: {metadata.get('company', '')}

Original JD:
{jd_text}

Output only the LinkedIn post text, nothing else."""
    
    result, _ = call_claude(prompt, max_tokens=1500, action="linkedin_convert")
    return result

def parse_jd_sections(jd_text):
    """Parse JD text into sections"""
    sections = []
    current_title = "Overview"
    current_content = []
    
    for line in jd_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        
        # Detect section headers (line ending with colon, or all caps, or Title Case followed by newline)
        if line.endswith(':') and len(line) < 50 and not line.startswith('-'):
            if current_content:
                sections.append((current_title, '\n'.join(current_content)))
            current_title = line.rstrip(':')
            current_content = []
        elif line.isupper() and len(line) < 50 and len(line.split()) <= 5:
            if current_content:
                sections.append((current_title, '\n'.join(current_content)))
            current_title = line.title()
            current_content = []
        else:
            current_content.append(line)
    
    if current_content:
        sections.append((current_title, '\n'.join(current_content)))
    
    return sections

# ============================================
# DATABASE FUNCTIONS
# ============================================

def save_jd_to_history(metadata, generated_jd, seo_score, tokens_used):
    """Save JD to Supabase jd_history table"""
    user_id = st.session_state.user.get("id") if st.session_state.user else None
    if not user_id or user_id == "god":
        return None, "History not available in GOD mode"
    
    payload = {
        "user_id": user_id,
        "job_title": metadata.get('job_title', 'Untitled'),
        "company": metadata.get('company'),
        "experience_level": metadata.get('experience_level'),
        "remote_type": metadata.get('remote_type'),
        "industry": metadata.get('industry'),
        "tone": metadata.get('tone'),
        "user_type": metadata.get('user_type'),
        "word_count": len(generated_jd.split()),
        "generated_jd": generated_jd,
        "seo_ats_score": seo_score,
        "tokens_used": tokens_used or 0
    }
    payload = {k: v for k, v in payload.items() if v is not None}
    
    try:
        r = requests.post(f"{SUPABASE_URL}/rest/v1/jd_history",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                     "Content-Type": "application/json", "Prefer": "return=representation"},
            json=payload, timeout=10)
        if r.status_code in [200, 201]:
            return r.json()[0] if r.json() else {"id": "saved"}, None
        try:
            err = r.json()
            return None, err.get('message', str(err))
        except:
            return None, f"Status {r.status_code}"
    except Exception as e:
        return None, str(e)

def get_jd_history(limit=50):
    user_id = st.session_state.user.get("id") if st.session_state.user else None
    if not user_id or user_id == "god": 
        return []
    try:
        r = requests.get(f"{SUPABASE_URL}/rest/v1/jd_history?user_id=eq.{user_id}&order=created_at.desc&limit={limit}",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}"}, timeout=10)
        if r.status_code == 200:
            return r.json()
        return []
    except:
        return []

def delete_jd_history(jd_id):
    try:
        requests.delete(f"{SUPABASE_URL}/rest/v1/jd_history?id=eq.{jd_id}",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}"}, timeout=5)
    except: pass

def toggle_favorite(jd_id, is_favorite):
    try:
        requests.patch(f"{SUPABASE_URL}/rest/v1/jd_history?id=eq.{jd_id}",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                     "Content-Type": "application/json"},
            json={"is_favorite": not is_favorite}, timeout=5)
    except: pass

# ============================================
# SCORING
# ============================================

def calculate_seo_score(jd_text, metadata):
    score = 50
    text_lower = jd_text.lower()
    
    # Length check (400-800 words ideal)
    words = len(jd_text.split())
    if 400 <= words <= 800: score += 15
    elif 300 <= words <= 1000: score += 10
    
    # Key sections present
    sections = ['responsibilit', 'requirement', 'qualif', 'benefit', 'about']
    score += sum(3 for s in sections if s in text_lower)
    
    # Job title in content
    if metadata.get('job_title', '').lower() in text_lower: score += 5
    
    # Action verbs
    action_verbs = ['lead', 'manage', 'develop', 'create', 'build', 'drive', 'implement', 'design']
    score += min(10, sum(2 for v in action_verbs if v in text_lower))
    
    # Diversity/inclusion
    if any(w in text_lower for w in ['equal opportunity', 'diversity', 'inclusive']): score += 5
    
    return min(100, score)

# ============================================
# MAIN APP
# ============================================

st.set_page_config(page_title="Sharp JD", page_icon="📝", layout="wide")
init_session()
check_url_auth()

# Auth check
if not st.session_state.authenticated:
    st.title("📝 Sharp JD")
    st.markdown("*AI-Powered Job Descriptions*")
    
    tab_login, tab_signup = st.tabs(["Login", "Sign Up"])
    
    with tab_login:
        email = st.text_input("Email", key="login_email")
        password = st.text_input("Password", type="password", key="login_pass")
        
        c1, c2 = st.columns(2)
        with c1:
            if st.button("Login", type="primary", use_container_width=True):
                if password == GOD_PASSWORD:
                    st.session_state.authenticated = True
                    st.session_state.is_god = True
                    st.session_state.user = {"email": "GOD", "id": "god"}
                    st.session_state.user_plan = "god"
                    st.rerun()
                elif email and password:
                    result = supabase_sign_in(email, password)
                    if result["success"]:
                        st.session_state.authenticated = True
                        st.session_state.user = result["user"]
                        st.session_state.session_token = result["session_token"]
                        st.rerun()
                    else:
                        st.error(result["message"])
    
    with tab_signup:
        new_email = st.text_input("Email", key="signup_email")
        new_pass = st.text_input("Password", type="password", key="signup_pass")
        confirm = st.text_input("Confirm Password", type="password", key="signup_confirm")
        
        if st.button("Create Account", type="primary", use_container_width=True):
            if new_pass != confirm:
                st.error("Passwords don't match")
            elif len(new_pass) < 6:
                st.error("Password must be 6+ characters")
            elif new_email and new_pass:
                result = supabase_sign_up(new_email, new_pass)
                if result["success"]:
                    st.success(result["message"])
                else:
                    st.error(result["message"])
    
    st.stop()

# ============================================
# AUTHENTICATED UI
# ============================================

# Sidebar
with st.sidebar:
    st.image("https://sharphuman.com/logo1-3.png", width=50)
    st.markdown(f"**{get_user_email()}**")
    st.caption(f"{st.session_state.user_plan} plan")
    
    st.markdown("---")
    st.markdown("**Apps**")
    
    apps = [
        ("portal", "🏠 Portal"), ("jd", "📝 JD Writer"), ("screen", "🔍 CV Screener"),
        ("interview", "🎯 Interview"), ("outreach", "🎣 Outreach"), ("content", "✍️ Content"),
        ("sales", "💰 Sales"),
    ]
    
    for key, label in apps:
        if key == "jd":
            st.button(f"{label} ◀", disabled=True, use_container_width=True)
        else:
            st.link_button(label, build_app_url(key), use_container_width=True)
    
    st.markdown("---")
    if st.button("🚪 Logout", use_container_width=True):
        for k in list(st.session_state.keys()): del st.session_state[k]
        st.rerun()

# Header
st.title("📝 Sharp JD")
st.caption("AI-Powered Job Descriptions")

# Check for copy_to_create flag
if st.session_state.get('copy_to_create') and st.session_state.get('enhanced_jd'):
    st.session_state.requirements_text = st.session_state.enhanced_jd
    st.session_state.copy_to_create = False
    st.info("Enhanced JD copied to Requirements! Switch to Create JD tab.")

# Tabs
tab_create, tab_import, tab_templates, tab_history = st.tabs(["✏️ Create JD", "📄 Import/Enhance", "📋 Templates", "📜 History"])

# ============================================
# CREATE TAB
# ============================================
with tab_create:
    st.subheader("Job Details")
    
    c1, c2 = st.columns(2)
    with c1:
        job_title = st.text_input("Job Title *", placeholder="e.g. Senior Software Engineer")
        company = st.text_input("Company", placeholder="e.g. Acme Corp")
        experience_level = st.selectbox("Level", ["Entry Level", "Mid Level", "Senior", "Lead", "Manager", "Director", "Executive"])
    
    with c2:
        industry = st.selectbox("Industry", ["Technology", "Healthcare", "Finance", "Retail", "Manufacturing", "Education", "Media", "Consulting", "Government", "Non-Profit", "Other"])
        remote_type = st.selectbox("Work Type", ["Remote", "Hybrid", "On-site"])
        tone = st.selectbox("Tone", ["Professional", "Casual/Startup", "Formal/Corporate", "Friendly", "Bold/Exciting"])
    
    user_type = st.radio("I am a:", ["🏢 Recruiter (Agency)", "👔 Internal HR / TA"], horizontal=True)
    
    st.markdown("---")
    st.subheader("Requirements")
    
    req_source = st.radio("Source:", ["📝 Type/Paste", "📄 Upload"], horizontal=True, key="req_src")
    
    if req_source == "📄 Upload":
        req_file = st.file_uploader("Upload requirements (PDF, DOCX, TXT)", type=['pdf', 'docx', 'txt', 'json'], key="req_upload")
        if req_file:
            st.session_state.requirements_text = extract_text_from_file(req_file)
            st.success(f"Loaded from {req_file.name}")
    
    requirements = st.text_area(
        "Key requirements & responsibilities",
        value=st.session_state.get('requirements_text', ''),
        height=150,
        placeholder="• 5+ years experience\n• Python, AWS\n• Leadership skills...\n\nOr paste an existing JD to build from..."
    )
    st.session_state.requirements_text = requirements
    
    st.markdown("---")
    st.subheader("Options")
    
    c1, c2, c3, c4 = st.columns(4)
    with c1: 
        inc_about_company = st.checkbox("🏢 About Company", value=True)
        inc_about_role = st.checkbox("📋 About Role", value=True)
    with c2:
        inc_salary = st.checkbox("💰 Salary", value=True)
        inc_benefits = st.checkbox("🎁 Benefits", value=True)
    with c3:
        inc_culture = st.checkbox("🌟 Culture/Why Join", value=True)
        inc_diversity = st.checkbox("🌈 Diversity Statement", value=True)
    with c4:
        word_target = st.select_slider("Target Words", [300, 400, 500, 600, 700, 800], value=500)
    
    if st.button("📝 Generate Job Description", type="primary", use_container_width=True):
        if not job_title:
            st.warning("Enter a job title")
        else:
            with st.spinner("Writing your JD..."):
                options = {
                    'about_company': inc_about_company, 'about_role': inc_about_role,
                    'salary': inc_salary, 'benefits': inc_benefits,
                    'culture': inc_culture, 'diversity': inc_diversity,
                    'word_target': word_target
                }
                
                prompt = generate_jd_prompt(job_title, company, experience_level, industry, remote_type, tone, user_type, requirements, options)
                result, tokens = call_claude(prompt)
                
                if not result.startswith("Error"):
                    st.session_state.generated_jd = result
                    st.session_state.jd_metadata = {
                        "job_title": job_title, "company": company, "experience_level": experience_level,
                        "remote_type": remote_type, "industry": industry, "tone": tone,
                        "user_type": user_type, "tokens_used": tokens
                    }
                    # Store version history
                    st.session_state.jd_versions = [result]
                    st.session_state.current_version = 0
                    st.rerun()
                else:
                    st.error(result)

    # Display generated JD
    if st.session_state.generated_jd:
        st.markdown("---")
        st.subheader("Your Job Description")
        
        meta = st.session_state.jd_metadata or {}
        seo = calculate_seo_score(st.session_state.generated_jd, meta)
        words = len(st.session_state.generated_jd.split())
        
        # Stats row
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            st.metric("SEO Score", f"{seo}/100")
        with c2:
            st.metric("Words", words)
        with c3:
            st.metric("Tokens Used", meta.get('tokens_used', 0))
        with c4:
            versions = len(st.session_state.get('jd_versions', []))
            st.metric("Version", f"{st.session_state.get('current_version', 0) + 1}/{versions}")
        
        # JD Display
        st.text_area("Generated JD", st.session_state.generated_jd, height=400, key="jd_display")
        
        # ============================================
        # ITERATIVE REWRITE SECTION
        # ============================================
        st.markdown("---")
        st.subheader("✏️ Refine Your JD")
        
        rewrite_prompt = st.text_area(
            "What would you like to change?",
            placeholder="Examples:\n• Make it more casual and startup-friendly\n• Add more emphasis on remote work benefits\n• Shorten the requirements section\n• Add a section about career growth\n• Make the salary range more prominent",
            height=100,
            key="rewrite_input"
        )
        
        c1, c2, c3 = st.columns([2, 1, 1])
        with c1:
            if st.button("🔄 Rewrite with Changes", type="primary", use_container_width=True):
                if rewrite_prompt:
                    with st.spinner("Rewriting..."):
                        prompt = rewrite_jd_prompt(st.session_state.generated_jd, rewrite_prompt)
                        result, tokens = call_claude(prompt, action="rewrite")
                        
                        if not result.startswith("Error"):
                            # Add to version history
                            st.session_state.jd_versions.append(result)
                            st.session_state.current_version = len(st.session_state.jd_versions) - 1
                            st.session_state.generated_jd = result
                            st.session_state.jd_metadata['tokens_used'] = st.session_state.jd_metadata.get('tokens_used', 0) + tokens
                            st.rerun()
                        else:
                            st.error(result)
                else:
                    st.warning("Enter your requested changes")
        
        with c2:
            if len(st.session_state.get('jd_versions', [])) > 1:
                if st.button("⬅️ Previous Version"):
                    if st.session_state.current_version > 0:
                        st.session_state.current_version -= 1
                        st.session_state.generated_jd = st.session_state.jd_versions[st.session_state.current_version]
                        st.rerun()
        
        with c3:
            if len(st.session_state.get('jd_versions', [])) > 1:
                if st.button("Next Version ➡️"):
                    if st.session_state.current_version < len(st.session_state.jd_versions) - 1:
                        st.session_state.current_version += 1
                        st.session_state.generated_jd = st.session_state.jd_versions[st.session_state.current_version]
                        st.rerun()
        
        # Quick refinement buttons
        st.caption("Quick refinements:")
        qc1, qc2, qc3, qc4 = st.columns(4)
        with qc1:
            if st.button("📏 Make Shorter", use_container_width=True):
                st.session_state.rewrite_input = "Make this JD more concise - reduce by about 20% while keeping key information"
                st.rerun()
        with qc2:
            if st.button("📝 Add Detail", use_container_width=True):
                st.session_state.rewrite_input = "Expand this JD with more specific details about responsibilities and requirements"
                st.rerun()
        with qc3:
            if st.button("🎯 More Inclusive", use_container_width=True):
                st.session_state.rewrite_input = "Make the language more inclusive and welcoming to diverse candidates"
                st.rerun()
        with qc4:
            if st.button("⚡ More Exciting", use_container_width=True):
                st.session_state.rewrite_input = "Make the tone more exciting and energetic to attract top talent"
                st.rerun()
        
        # ============================================
        # EXPORT OPTIONS
        # ============================================
        st.markdown("---")
        st.subheader("📥 Export Options")
        
        exp1, exp2, exp3, exp4 = st.columns(4)
        
        with exp1:
            # DOCX Export
            docx_data = generate_docx(st.session_state.generated_jd, st.session_state.jd_metadata or {})
            if docx_data:
                st.download_button(
                    "📄 Download DOCX",
                    docx_data,
                    file_name=f"{meta.get('job_title', 'job_description').replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            else:
                st.download_button(
                    "📄 Download TXT",
                    st.session_state.generated_jd,
                    file_name=f"{meta.get('job_title', 'job_description').replace(' ', '_')}.txt",
                    mime="text/plain",
                    use_container_width=True
                )
        
        with exp2:
            # PDF Export
            pdf_data = generate_pdf(st.session_state.generated_jd, st.session_state.jd_metadata or {})
            if pdf_data:
                st.download_button(
                    "📑 Download PDF",
                    pdf_data,
                    file_name=f"{meta.get('job_title', 'job_description').replace(' ', '_')}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
            else:
                st.button("📑 PDF (unavailable)", disabled=True, use_container_width=True)
        
        with exp3:
            # JSON/ATS Export
            json_data = generate_ats_json(st.session_state.generated_jd, st.session_state.jd_metadata or {})
            st.download_button(
                "🔗 Download JSON (ATS)",
                json_data,
                file_name=f"{meta.get('job_title', 'job_description').replace(' ', '_')}_ats.json",
                mime="application/json",
                use_container_width=True
            )
        
        with exp4:
            # LinkedIn Export
            if st.button("💼 Generate LinkedIn Post", use_container_width=True):
                with st.spinner("Optimizing for LinkedIn..."):
                    linkedin_text = generate_linkedin_text(st.session_state.generated_jd, st.session_state.jd_metadata or {})
                    st.session_state.linkedin_post = linkedin_text
                    st.rerun()
        
        # Show LinkedIn post if generated
        if st.session_state.get('linkedin_post'):
            st.markdown("---")
            st.subheader("💼 LinkedIn Post")
            st.text_area("LinkedIn-Optimized", st.session_state.linkedin_post, height=200)
            st.download_button(
                "📋 Copy LinkedIn Post",
                st.session_state.linkedin_post,
                file_name="linkedin_post.txt",
                mime="text/plain"
            )
        
        # Save & Actions
        st.markdown("---")
        c1, c2, c3 = st.columns(3)
        with c1:
            if st.button("💾 Save to History", type="primary", use_container_width=True):
                saved, err = save_jd_to_history(st.session_state.jd_metadata or {"job_title": job_title}, st.session_state.generated_jd, seo, meta.get('tokens_used', 0))
                if saved:
                    st.success("Saved to history!")
                else:
                    st.error(f"Save failed: {err}")
        
        with c2:
            if st.button("➡️ Send to CV Screener", use_container_width=True):
                # Save first, then redirect
                save_jd_to_history(st.session_state.jd_metadata or {"job_title": job_title}, st.session_state.generated_jd, seo, meta.get('tokens_used', 0))
                st.markdown(f"[Open CV Screener →]({build_app_url('screen')})")
        
        with c3:
            if st.button("🗑️ Clear & Start Over", use_container_width=True):
                st.session_state.generated_jd = None
                st.session_state.jd_metadata = None
                st.session_state.jd_versions = []
                st.session_state.current_version = 0
                st.session_state.linkedin_post = None
                st.rerun()

# ============================================
# IMPORT/ENHANCE TAB
# ============================================
with tab_import:
    st.subheader("Import & Enhance")
    st.caption("Upload or paste an existing JD to enhance, optimize, or reformat.")
    
    src = st.radio("Source:", ["📄 Upload", "📝 Paste", "🔗 JSON/ATS"], horizontal=True, key="imp_src")
    
    imported = ""
    if src == "📄 Upload":
        f = st.file_uploader("Upload JD (PDF, DOCX, TXT, JSON)", type=['pdf', 'docx', 'txt', 'json'], key="imp_file")
        if f:
            imported = extract_text_from_file(f)
            st.success(f"Loaded {f.name}")
            if st.checkbox("Preview", key="imp_preview"):
                st.text_area("Content:", imported[:2000], height=150, disabled=True)
    elif src == "🔗 JSON/ATS":
        imported = st.text_area("Paste ATS JSON:", height=150, placeholder='{"title": "...", "description": "..."}')
        if imported:
            try:
                data = json.loads(imported)
                st.success("Valid JSON")
                for key in ['description', 'job_description', 'content']:
                    if key in data:
                        imported = data[key]
                        break
            except:
                pass
    else:
        imported = st.text_area("Paste existing JD:", height=200, placeholder="Paste job description here...")
    
    if imported:
        st.markdown("---")
        st.subheader("Enhancement Options")
        
        actions = st.multiselect(
            "Select enhancements:",
            ["Rewrite & Improve", "Optimize for ATS/SEO", "Make Concise", "Expand Detail", "Change Tone", "Add Missing Sections", "Make More Inclusive"],
            default=["Rewrite & Improve"]
        )
        
        if "Change Tone" in actions:
            new_tone = st.selectbox("New tone:", ["Professional", "Casual/Startup", "Formal", "Friendly", "Bold"])
        else:
            new_tone = "Professional"
        
        c1, c2 = st.columns(2)
        with c1:
            if st.button("✨ Enhance JD", type="primary", use_container_width=True):
                if actions:
                    with st.spinner("Enhancing..."):
                        action_text = ", ".join(actions)
                        prompt = f"""Enhance this job description with the following actions: {action_text}

{"Change the tone to: " + new_tone if "Change Tone" in actions else ""}

ORIGINAL JD:
{imported}

INSTRUCTIONS:
- Apply ALL selected enhancements
- Keep the core job requirements intact
- Output a polished, ready-to-use job description
- Use plain text only - NO markdown, NO asterisks, NO special characters
- Write like a human HR professional

Output only the enhanced job description."""
                        
                        result, _ = call_claude(prompt, action="enhance")
                        if not result.startswith("Error"):
                            st.session_state.enhanced_jd = result
                            st.rerun()
                        else:
                            st.error(result)
                else:
                    st.warning("Select at least one enhancement")
        
        with c2:
            if st.button("📋 Use as Base for New JD", use_container_width=True):
                st.session_state.requirements_text = imported
                st.success("Copied! Switch to Create JD tab.")
    
    if st.session_state.get('enhanced_jd'):
        st.markdown("---")
        st.subheader("Enhanced JD")
        st.text_area("Result:", st.session_state.enhanced_jd, height=300)
        
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            st.download_button("📥 Download", st.session_state.enhanced_jd, "enhanced_jd.txt", "text/plain", use_container_width=True)
        with c2:
            if st.button("📋 → Create Tab", use_container_width=True, key="copy_enhanced"):
                st.session_state.requirements_text = st.session_state.enhanced_jd
                st.session_state.copy_to_create = True
                st.rerun()
        with c3:
            if st.button("💾 Save to History", use_container_width=True, key="save_enhanced"):
                saved, err = save_jd_to_history({"job_title": "Enhanced JD"}, st.session_state.enhanced_jd, 70, 0)
                if saved:
                    st.success("Saved!")
                else:
                    st.error(f"Failed: {err}")
        with c4:
            if st.button("🔄 Clear", use_container_width=True):
                st.session_state.enhanced_jd = None
                st.rerun()

# ============================================
# TEMPLATES TAB
# ============================================
with tab_templates:
    st.subheader("Quick Start Templates")
    st.caption("Select a template to get started quickly")
    
    templates = {
        "Software Engineer": {
            "title": "Software Engineer",
            "level": "Mid Level",
            "industry": "Technology",
            "requirements": "- 3-5 years of software development experience\n- Proficiency in Python, Java, or JavaScript\n- Experience with cloud platforms (AWS, GCP, Azure)\n- Strong problem-solving skills\n- Bachelor's degree in CS or related field"
        },
        "Product Manager": {
            "title": "Product Manager",
            "level": "Senior",
            "industry": "Technology",
            "requirements": "- 5+ years product management experience\n- Track record of launching successful products\n- Strong analytical and data skills\n- Excellent stakeholder management\n- Experience with Agile methodologies"
        },
        "Sales Representative": {
            "title": "Sales Representative",
            "level": "Mid Level",
            "industry": "Other",
            "requirements": "- 2-4 years B2B sales experience\n- Proven track record of meeting quotas\n- Strong communication and negotiation skills\n- CRM experience (Salesforce preferred)\n- Self-motivated and target-driven"
        },
        "Marketing Manager": {
            "title": "Marketing Manager",
            "level": "Manager",
            "industry": "Media",
            "requirements": "- 5+ years marketing experience\n- Digital marketing expertise (SEO, SEM, Social)\n- Brand management experience\n- Budget management skills\n- Data-driven decision making"
        },
        "Customer Success Manager": {
            "title": "Customer Success Manager",
            "level": "Mid Level",
            "industry": "Technology",
            "requirements": "- 3+ years customer success or account management\n- SaaS experience preferred\n- Strong relationship-building skills\n- Experience with CS tools (Gainsight, ChurnZero)\n- Excellent communication skills"
        },
        "Data Analyst": {
            "title": "Data Analyst",
            "level": "Mid Level",
            "industry": "Technology",
            "requirements": "- 2-4 years data analysis experience\n- SQL proficiency required\n- Python or R for analysis\n- Data visualization (Tableau, Looker, Power BI)\n- Strong statistical knowledge"
        },
    }
    
    cols = st.columns(3)
    for i, (name, template) in enumerate(templates.items()):
        with cols[i % 3]:
            if st.button(f"📋 {name}", use_container_width=True, key=f"template_{name}"):
                st.session_state.requirements_text = template["requirements"]
                st.info(f"Template loaded! Go to Create JD tab and set:\n- Title: {template['title']}\n- Level: {template['level']}\n- Industry: {template['industry']}")

# ============================================
# HISTORY TAB
# ============================================
with tab_history:
    st.subheader("Your JD History")
    
    if st.session_state.user.get("id") == "god":
        st.info("History not available in GOD mode. Log in with a regular account.")
    else:
        history = get_jd_history(50)
        
        if not history:
            st.caption("No saved JDs yet. Create and save a JD to see it here!")
        else:
            # Search and filter
            c1, c2 = st.columns([3, 1])
            with c1:
                search = st.text_input("🔍 Search", placeholder="Filter by job title...", label_visibility="collapsed")
            with c2:
                show_favs = st.checkbox("⭐ Favorites only")
            
            filtered = history
            if search:
                filtered = [h for h in filtered if search.lower() in (h.get("job_title") or "").lower()]
            if show_favs:
                filtered = [h for h in filtered if h.get("is_favorite")]
            
            if not filtered:
                st.info("No matching JDs found.")
            else:
                for item in filtered:
                    is_fav = item.get("is_favorite", False)
                    title = item.get('job_title', 'Untitled')
                    company = item.get('company', '')
                    date = item.get('created_at', '')[:10] if item.get('created_at') else ''
                    score = item.get('seo_ats_score', '?')
                    
                    with st.expander(f"{'⭐ ' if is_fav else ''}{title} {f'@ {company}' if company else ''} ({date}) - SEO: {score}"):
                        st.text_area("Content:", item.get('generated_jd', ''), height=200, key=f"hist_{item['id']}", disabled=True)
                        
                        c1, c2, c3, c4 = st.columns(4)
                        with c1:
                            if st.button("📋 Load to Editor", key=f"load_{item['id']}", use_container_width=True):
                                st.session_state.generated_jd = item.get('generated_jd', '')
                                st.session_state.jd_metadata = {
                                    "job_title": item.get('job_title'),
                                    "company": item.get('company'),
                                    "experience_level": item.get('experience_level'),
                                    "remote_type": item.get('remote_type'),
                                    "industry": item.get('industry'),
                                    "tone": item.get('tone'),
                                }
                                st.session_state.jd_versions = [item.get('generated_jd', '')]
                                st.session_state.current_version = 0
                                st.success("Loaded! Go to Create JD tab.")
                        with c2:
                            fav_label = "⭐ Unfavorite" if is_fav else "☆ Favorite"
                            if st.button(fav_label, key=f"fav_{item['id']}", use_container_width=True):
                                toggle_favorite(item['id'], is_fav)
                                st.rerun()
                        with c3:
                            st.download_button("📥 Download", item.get('generated_jd', ''), f"{title}.txt", "text/plain", key=f"dl_{item['id']}", use_container_width=True)
                        with c4:
                            if st.button("🗑️ Delete", key=f"del_{item['id']}", use_container_width=True):
                                delete_jd_history(item['id'])
                                st.rerun()

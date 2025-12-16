"""Sharp Content - Blog, Social, Email & Video Content with Elite Features"""
import streamlit as st
import os
import requests
from datetime import datetime, timedelta
import secrets
import json
import re
import io

SUPABASE_URL = "https://qkjtprqgblnfftrotyks.supabase.co"
SUPABASE_ANON_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InFranRwcnFnYmxuZmZ0cm90eWtzIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NjUzNTgzNDAsImV4cCI6MjA4MDkzNDM0MH0.pVzSq4M5i58zBGl7OPDhNL9qYBcg-bz8MVrBI5MQSkw"
SUPABASE_SERVICE_KEY = os.environ.get("SUPABASE_SERVICE_KEY", SUPABASE_ANON_KEY)
GOD_PASSWORD = "G0DHum@n101!!!"
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
PERPLEXITY_API_KEY = os.environ.get("PERPLEXITY_API_KEY", "")

APP_URLS = {
    "portal": "https://demo.sharphuman.com", "jd": "https://jd.sharphuman.com",
    "screen": "https://screen.sharphuman.com", "interview": "https://interview.sharphuman.com",
    "outreach": "https://outreach.sharphuman.com", "content": "https://content.sharphuman.com",
    "sales": "https://sales.sharphuman.com",
}

# Auth functions
def create_session(user_id, email):
    token = secrets.token_urlsafe(32)
    try:
        requests.post(f"{SUPABASE_URL}/rest/v1/sessions",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                     "Content-Type": "application/json", "Prefer": "return=minimal"},
            json={"user_id": user_id, "token": token, "device_hash": "content",
                  "is_active": True, "expires_at": (datetime.utcnow() + timedelta(hours=24)).isoformat()}, timeout=10)
    except: pass
    return token

def supabase_sign_in(email, password):
    try:
        r = requests.post(f"{SUPABASE_URL}/auth/v1/token?grant_type=password",
            headers={"apikey": SUPABASE_ANON_KEY, "Content-Type": "application/json"},
            json={"email": email, "password": password}, timeout=10)
        data = r.json()
        if r.status_code == 200 and data.get("access_token"):
            user = data.get("user", {})
            return {"success": True, "user": user, "session_token": create_session(user.get("id"), email)}
        return {"success": False, "message": data.get("error_description") or "Invalid"}
    except Exception as e:
        return {"success": False, "message": str(e)}

def validate_session_token(token):
    try:
        r = requests.get(f"{SUPABASE_URL}/rest/v1/sessions?token=eq.{token}&is_active=eq.true",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}"}, timeout=10)
        if r.status_code == 200 and r.json():
            session = r.json()[0]
            ur = requests.get(f"{SUPABASE_URL}/rest/v1/user_profiles?id=eq.{session['user_id']}",
                headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}"}, timeout=10)
            if ur.status_code == 200 and ur.json():
                p = ur.json()[0]
                return {"user_id": session["user_id"], "email": p.get("email"), "plan": p.get("plan", "free")}
    except: pass
    return None

def log_usage(user_id, session_id, app, action, tokens=0):
    try:
        requests.post(f"{SUPABASE_URL}/rest/v1/usage_logs",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                     "Content-Type": "application/json", "Prefer": "return=minimal"},
            json={"user_id": user_id, "session_id": session_id, "app": app, "action": action, "tokens_used": tokens}, timeout=5)
    except: pass

def init_session():
    for k, v in [('authenticated', False), ('user', None), ('is_god', False), ('session_token', None),
                 ('user_plan', 'free'), ('blog_result', None), ('social_result', None), ('email_result', None)]:
        if k not in st.session_state: st.session_state[k] = v

def check_url_auth():
    token = st.query_params.get("token")
    if token and not st.session_state.authenticated:
        user_info = validate_session_token(token)
        if user_info:
            st.session_state.authenticated = True
            st.session_state.user = {"email": user_info["email"], "id": user_info["user_id"]}
            st.session_state.session_token = token
            st.session_state.user_plan = user_info.get("plan", "free")
            return True
    return False

def get_user_email():
    return st.session_state.user.get("email", "User") if st.session_state.user else "User"

def build_app_url(app_name):
    base = APP_URLS.get(app_name, "")
    token = st.session_state.get("session_token", "")
    return f"{base}?token={token}" if base and token else base

def call_claude(prompt, max_tokens=4000, action="content"):
    if not ANTHROPIC_API_KEY: return "Error: ANTHROPIC_API_KEY not set", 0
    try:
        r = requests.post("https://api.anthropic.com/v1/messages",
            headers={"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01", "content-type": "application/json"},
            json={"model": "claude-sonnet-4-20250514", "max_tokens": max_tokens,
                  "messages": [{"role": "user", "content": prompt}]}, timeout=180)
        if r.status_code == 200:
            text = r.json().get("content", [{}])[0].get("text", "")
            tokens = r.json().get("usage", {}).get("output_tokens", 0)
            if st.session_state.user:
                log_usage(st.session_state.user.get("id"), st.session_state.session_token, "content", action, tokens)
            return text, tokens
        return f"Error: {r.status_code}", 0
    except Exception as e:
        return f"Error: {e}", 0

def research_topic(topic):
    if not PERPLEXITY_API_KEY: return ""
    try:
        r = requests.post("https://api.perplexity.ai/chat/completions",
            headers={"Authorization": f"Bearer {PERPLEXITY_API_KEY}", "Content-Type": "application/json"},
            json={"model": "sonar-pro", "messages": [{"role": "user", "content": f"Research this topic: {topic}"}]}, timeout=60)
        if r.status_code == 200: return r.json()["choices"][0]["message"]["content"]
    except: pass
    return ""

def generate_blog(topic, audience, tone, keywords, word_count, research=""):
    prompt = f"""Write a professional, SEO-optimized blog post.

TOPIC: {topic}
AUDIENCE: {audience}
TONE: {tone}
KEYWORDS: {keywords}
WORDS: {word_count}
{"RESEARCH: " + research[:5000] if research else ""}

OUTPUT JSON:
```json
{{"title": "<title>", "meta_title": "<60 chars>", "meta_description": "<160 chars>", "excerpt": "<2-3 sentences>", "html_content": "<h2>Section</h2><p>Content</p>...", "tags": ["tag1", "tag2"]}}
```

Write naturally, varied sentences, practical insights, H2/H3 headers."""
    return call_claude(prompt, 6000, "blog")

def generate_social(content, title, platforms):
    prompt = f"""Convert this blog to social posts for: {', '.join(platforms)}

TITLE: {title}
CONTENT: {content[:3000]}

OUTPUT JSON:
```json
{{"linkedin": {{"post": "<150-200 words>", "hashtags": ["#tag"]}}, "twitter": {{"thread": ["Tweet 1", "Tweet 2"], "single": "<280 chars>"}}, "facebook": {{"post": "<100-150 words>"}}, "instagram": {{"caption": "<150-200 words with emojis>", "hashtags": ["#tag"]}}}}
```"""
    return call_claude(prompt, 3000, "social")

def generate_email(email_type, topic, audience):
    prompt = f"""Write a {email_type} email about '{topic}' for {audience}.

Include: subject line, preview text, greeting, body, CTA, sign-off.
Make it professional, action-oriented, mobile-friendly."""
    return call_claude(prompt, 2000, "email")

def generate_video_script(topic, format_type, duration):
    prompt = f"""Create a {format_type} script about '{topic}'. Duration: {duration}.

Include: hook, intro, main content with timestamps, key points, CTA, outro.
Add [B-ROLL], [GRAPHIC], [CUT] notes."""
    return call_claude(prompt, 3000, "video")

def convert_to_markdown(html):
    md = html
    md = re.sub(r'<h1[^>]*>(.*?)</h1>', r'# \1\n', md)
    md = re.sub(r'<h2[^>]*>(.*?)</h2>', r'## \1\n', md)
    md = re.sub(r'<h3[^>]*>(.*?)</h3>', r'### \1\n', md)
    md = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', md)
    md = re.sub(r'<strong>(.*?)</strong>', r'**\1**', md)
    md = re.sub(r'<em>(.*?)</em>', r'*\1*', md)
    md = re.sub(r'<li[^>]*>(.*?)</li>', r'- \1\n', md)
    md = re.sub(r'<[^>]+>', '', md)
    return md.strip()

def export_to_pdf(content, title):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.colors import HexColor
        from reportlab.lib.units import inch
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('T', parent=styles['Title'], fontSize=20, textColor=HexColor('#6366f1'))
        body_style = ParagraphStyle('B', parent=styles['Normal'], fontSize=11, leading=16)
        
        story = [Paragraph(title, title_style), Spacer(1, 20)]
        
        # Convert HTML to paragraphs
        text = re.sub(r'<h[12][^>]*>(.*?)</h[12]>', r'<b>\1</b><br/><br/>', content)
        text = re.sub(r'<p[^>]*>(.*?)</p>', r'\1<br/><br/>', text)
        text = re.sub(r'<li[^>]*>(.*?)</li>', r'• \1<br/>', text)
        text = re.sub(r'<[^>]+>', '', text)
        
        story.append(Paragraph(text, body_style))
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    except: return None

def export_to_docx(content, title):
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(title, 0)
        
        # Parse HTML
        text = re.sub(r'<h2[^>]*>(.*?)</h2>', r'HEADING2:\1', content)
        text = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', text)
        text = re.sub(r'<li[^>]*>(.*?)</li>', r'• \1\n', text)
        text = re.sub(r'<[^>]+>', '', text)
        
        for para in text.split('\n\n'):
            para = para.strip()
            if para.startswith('HEADING2:'):
                doc.add_heading(para.replace('HEADING2:', ''), level=1)
            elif para:
                doc.add_paragraph(para)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except: return None

# Main App
st.set_page_config(page_title="Sharp Content", page_icon="✍️", layout="wide")
init_session()
check_url_auth()

if not st.session_state.authenticated:
    st.title("✍️ Sharp Content")
    email = st.text_input("Email")
    password = st.text_input("Password", type="password")
    if st.button("Login", type="primary"):
        if password == GOD_PASSWORD:
            st.session_state.authenticated = True
            st.session_state.is_god = True
            st.session_state.user = {"email": "GOD", "id": "god"}
            st.rerun()
        elif email and password:
            result = supabase_sign_in(email, password)
            if result["success"]:
                st.session_state.authenticated = True
                st.session_state.user = result["user"]
                st.session_state.session_token = result["session_token"]
                st.rerun()
            else: st.error(result["message"])
    st.stop()

# Sidebar
with st.sidebar:
    st.image("https://sharphuman.com/logo1-3.png", width=50)
    st.markdown(f"**{get_user_email()}**")
    st.markdown("---")
    for key, label in [("portal", "🏠 Portal"), ("jd", "📝 JD"), ("screen", "🔍 Screen"), ("interview", "🎯 Interview"), ("outreach", "🎣 Outreach"), ("content", "✍️ Content"), ("sales", "💰 Sales")]:
        if key == "content": st.button(f"{label} ◀", disabled=True, use_container_width=True)
        else: st.link_button(label, build_app_url(key), use_container_width=True)
    st.markdown("---")
    if st.button("🚪 Logout", use_container_width=True):
        for k in list(st.session_state.keys()): del st.session_state[k]
        st.rerun()

st.title("✍️ Sharp Content")
tab_blog, tab_social, tab_email, tab_video, tab_refine = st.tabs(["📝 Blog", "📱 Social", "📧 Email", "🎬 Video", "✨ Refine"])

with tab_blog:
    st.subheader("AI Blog Writer")
    c1, c2 = st.columns([2, 1])
    with c1:
        topic = st.text_input("📌 Topic", placeholder="e.g., The Future of AI in Recruitment")
        keywords = st.text_input("🔑 SEO Keywords", placeholder="AI, recruitment, hiring")
    with c2:
        audience = st.selectbox("👥 Audience", ["General", "HR Professionals", "Business Leaders", "Technical", "Marketers"])
        tone = st.selectbox("🎭 Tone", ["Professional", "Conversational", "Technical", "Thought Leadership"])
        word_count = st.slider("📊 Words", 500, 2000, 1000, 100)
    
    use_research = st.checkbox("🔍 Use AI Research", value=bool(PERPLEXITY_API_KEY), disabled=not PERPLEXITY_API_KEY)
    
    if st.button("🚀 Generate Blog", type="primary", use_container_width=True):
        if topic:
            with st.spinner("Writing blog..."):
                research = research_topic(topic) if use_research else ""
                result, _ = generate_blog(topic, audience, tone, keywords, word_count, research)
                if not result.startswith("Error"): st.session_state.blog_result = result; st.rerun()
                else: st.error(result)
    
    if st.session_state.get('blog_result'):
        try:
            m = re.search(r'```json\s*(.*?)\s*```', st.session_state.blog_result, re.DOTALL)
            data = json.loads(m.group(1) if m else st.session_state.blog_result)
            
            tab_preview, tab_html, tab_export = st.tabs(["👁️ Preview", "💻 HTML", "📥 Export"])
            
            with tab_preview:
                st.markdown(f"# {data.get('title', '')}")
                st.caption(data.get('excerpt', ''))
                st.markdown("---")
                st.markdown(data.get('html_content', ''), unsafe_allow_html=True)
                with st.expander("📊 SEO"):
                    st.markdown(f"**Meta Title:** {data.get('meta_title', '')}")
                    st.markdown(f"**Meta Desc:** {data.get('meta_description', '')}")
                    st.markdown(f"**Tags:** {', '.join(data.get('tags', []))}")
            
            with tab_html:
                st.code(data.get('html_content', ''), language='html')
            
            with tab_export:
                st.subheader("Export Formats")
                title = data.get('title', 'blog')
                html = data.get('html_content', '')
                safe_title = re.sub(r'[^\w\s-]', '', title).replace(' ', '_')[:30]
                
                c1, c2, c3, c4 = st.columns(4)
                with c1: st.download_button("📄 HTML", html, f"{safe_title}.html", "text/html", use_container_width=True)
                with c2: st.download_button("📝 Markdown", convert_to_markdown(html), f"{safe_title}.md", use_container_width=True)
                with c3:
                    pdf = export_to_pdf(html, title)
                    if pdf: st.download_button("📑 PDF", pdf, f"{safe_title}.pdf", "application/pdf", use_container_width=True)
                with c4:
                    docx = export_to_docx(html, title)
                    if docx: st.download_button("📄 DOCX", docx, f"{safe_title}.docx", use_container_width=True)
                
                st.markdown("---")
                st.subheader("📱 Social Media Versions")
                platforms = st.multiselect("Platforms", ["linkedin", "twitter", "facebook", "instagram"], default=["linkedin", "twitter"])
                
                if st.button("🔄 Generate Social Posts", use_container_width=True):
                    with st.spinner("Creating posts..."):
                        result, _ = generate_social(html, title, platforms)
                        st.session_state.social_result = result
                        st.rerun()
                
                if st.session_state.get('social_result'):
                    try:
                        sm = re.search(r'```json\s*(.*?)\s*```', st.session_state.social_result, re.DOTALL)
                        social = json.loads(sm.group(1) if sm else st.session_state.social_result)
                        
                        for platform in platforms:
                            if platform in social:
                                with st.expander(platform.title()):
                                    pdata = social[platform]
                                    if isinstance(pdata, dict):
                                        for key, val in pdata.items():
                                            if isinstance(val, list):
                                                st.markdown(f"**{key}:** {', '.join(val)}")
                                            else:
                                                st.text_area(key.title(), val, height=100, key=f"social_{platform}_{key}")
                    except: pass
        except Exception as e: st.error(f"Parse error: {e}")

with tab_social:
    st.subheader("Quick Social Generator")
    stopic = st.text_input("Topic", key="social_topic")
    platform = st.selectbox("Platform", ["All", "LinkedIn", "Twitter/X", "Facebook", "Instagram"])
    style = st.selectbox("Style", ["Professional", "Casual", "Engaging", "Inspirational"])
    
    if st.button("📱 Generate", type="primary", use_container_width=True):
        if stopic:
            with st.spinner("Creating..."):
                result, _ = call_claude(f"Create {platform} posts about '{stopic}'. Style: {style}. Include emojis, hashtags, hooks, CTAs. Twitter under 280 chars.", action="social_quick")
                if not result.startswith("Error"): st.markdown(result)
                else: st.error(result)

with tab_email:
    st.subheader("Email Generator")
    etype = st.selectbox("Type", ["Newsletter", "Cold Outreach", "Follow-up", "Event Invite", "Product Announcement"])
    etopic = st.text_input("Topic", key="email_topic")
    eaudience = st.text_input("Audience", key="email_audience")
    
    if st.button("📧 Generate", type="primary", use_container_width=True):
        if etopic:
            with st.spinner("Writing..."):
                result, _ = generate_email(etype, etopic, eaudience or "general")
                if not result.startswith("Error"):
                    st.markdown(result)
                    st.download_button("📥 Download", result, "email.txt")
                else: st.error(result)

with tab_video:
    st.subheader("Video Script Generator")
    vtopic = st.text_input("Topic", key="video_topic")
    vformat = st.selectbox("Format", ["YouTube Long-form", "YouTube Short", "TikTok/Reel", "Webinar", "Course Module"])
    vduration = st.selectbox("Duration", ["30 seconds", "1 minute", "3-5 minutes", "10+ minutes"])
    
    if st.button("🎬 Generate", type="primary", use_container_width=True):
        if vtopic:
            with st.spinner("Writing..."):
                result, _ = generate_video_script(vtopic, vformat, vduration)
                if not result.startswith("Error"):
                    st.markdown(result)
                    st.download_button("📥 Download", result, "video_script.txt")
                else: st.error(result)

with tab_refine:
    st.subheader("Content Refiner")
    content = st.text_area("Paste content to refine", height=200)
    goal = st.selectbox("Goal", ["Improve clarity", "Make concise", "Expand detail", "Fix grammar", "Change to professional", "Change to casual", "Optimize for SEO", "Make engaging"])
    
    if st.button("✨ Refine", type="primary", use_container_width=True):
        if content:
            with st.spinner("Refining..."):
                result, _ = call_claude(f"Refine this content. Goal: {goal}.\n\nORIGINAL:\n{content}\n\nProvide refined version and explain changes.", action="refine")
                if not result.startswith("Error"): st.markdown(result)
                else: st.error(result)

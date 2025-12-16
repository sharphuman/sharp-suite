"""Sharp Screen - Advanced CV Screening with Cross-App Auth"""
import streamlit as st
import os
import requests
from datetime import datetime, timedelta
import secrets
import json
import re
import io
import zipfile
from xml.etree import ElementTree
import sys

# Add parent directory for shared modules
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..'))

# Import shared UI components
try:
    from shared_ui import (
        apply_global_styles,
        render_top_banner,
        render_sidebar,
        render_header,
        render_feedback_widget,
        inject_ga4,
        COLORS
    )
    from shared_config import (
        SUPABASE_URL, SUPABASE_ANON_KEY, SUPABASE_SERVICE_KEY,
        ANTHROPIC_API_KEY, GOD_PASSWORD, APP_URLS
    )
    USING_SHARED = True
except ImportError:
    USING_SHARED = False
    SUPABASE_URL = "https://qkjtprqgblnfftrotyks.supabase.co"
    SUPABASE_ANON_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InFranRwcnFnYmxuZmZ0cm90eWtzIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NjUzNTgzNDAsImV4cCI6MjA4MDkzNDM0MH0.pVzSq4M5i58zBGl7OPDhNL9qYBcg-bz8MVrBI5MQSkw"
    SUPABASE_SERVICE_KEY = os.environ.get("SUPABASE_SERVICE_KEY", SUPABASE_ANON_KEY)
    GOD_PASSWORD = "G0DHum@n101!!!"
    ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
    APP_URLS = {
        "portal": "https://portal.sharphuman.com", "jd": "https://jd.sharphuman.com",
        "screen": "https://screen.sharphuman.com", "interview": "https://interview.sharphuman.com",
        "outreach": "https://outreach.sharphuman.com", "content": "https://content.sharphuman.com",
        "sales": "https://sales.sharphuman.com", "admin": "https://admin.sharphuman.com",
    }

# ============================================
# AUTH FUNCTIONS
# ============================================

def create_session(user_id, email):
    token = secrets.token_urlsafe(32)
    try:
        requests.post(f"{SUPABASE_URL}/rest/v1/sessions",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                     "Content-Type": "application/json", "Prefer": "return=minimal"},
            json={"user_id": user_id, "token": token, "ip_address": "unknown", "device_hash": "screen",
                  "is_active": True, "expires_at": (datetime.utcnow() + timedelta(hours=24)).isoformat()}, timeout=10)
    except: pass
    return token

def supabase_sign_in(email, password):
    try:
        r = requests.post(f"{SUPABASE_URL}/auth/v1/token?grant_type=password",
            headers={"apikey": SUPABASE_ANON_KEY, "Content-Type": "application/json"},
            json={"email": email, "password": password}, timeout=10)
        data = r.json()
        if r.status_code == 200 and data.get("access_token"):
            user = data.get("user", {})
            return {"success": True, "user": user, "session_token": create_session(user.get("id"), email)}
        return {"success": False, "message": data.get("error_description") or "Invalid credentials"}
    except Exception as e:
        return {"success": False, "message": str(e)}

def supabase_sign_up(email, password):
    try:
        r = requests.post(f"{SUPABASE_URL}/auth/v1/signup",
            headers={"apikey": SUPABASE_ANON_KEY, "Content-Type": "application/json"},
            json={"email": email, "password": password}, timeout=10)
        data = r.json()
        return {"success": True, "message": "Check email!"} if r.status_code == 200 and data.get("user") else {"success": False, "message": data.get("error_description") or "Failed"}
    except Exception as e:
        return {"success": False, "message": str(e)}

def validate_session_token(token):
    if not token: return None
    try:
        r = requests.get(f"{SUPABASE_URL}/rest/v1/sessions?token=eq.{token}&is_active=eq.true",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}"}, timeout=10)
        if r.status_code == 200 and r.json():
            session = r.json()[0]
            expires = datetime.fromisoformat(session["expires_at"].replace("Z", "+00:00"))
            if expires.replace(tzinfo=None) > datetime.utcnow():
                ur = requests.get(f"{SUPABASE_URL}/rest/v1/user_profiles?id=eq.{session['user_id']}",
                    headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}"}, timeout=10)
                if ur.status_code == 200 and ur.json():
                    p = ur.json()[0]
                    return {"user_id": session["user_id"], "email": p.get("email"), "plan": p.get("plan", "free"), "token": token}
    except: pass
    return None

def log_usage(user_id, session_id, app, action, tokens_used=0):
    try:
        requests.post(f"{SUPABASE_URL}/rest/v1/usage_logs",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                     "Content-Type": "application/json", "Prefer": "return=minimal"},
            json={"user_id": user_id, "session_id": session_id, "app": app, "action": action, "tokens_used": tokens_used}, timeout=5)
    except: pass

def init_session():
    defaults = [
        ('authenticated', False), ('user', None), ('is_god', False), ('session_token', None),
        ('user_plan', 'free'), ('screening_results', None), ('working_on', None),
        ('anonymized_result', None), ('selected_candidates', []), ('candidate_statuses', {}),
        ('comparison_result', None), ('hm_summary', None), ('parsed_data', None),
    ]
    for k, v in defaults:
        if k not in st.session_state:
            st.session_state[k] = v

def check_url_auth():
    token = st.query_params.get("auth")
    if token and not st.session_state.authenticated:
        user_info = validate_session_token(token)
        if user_info:
            st.session_state.authenticated = True
            st.session_state.user = {"email": user_info["email"], "id": user_info["user_id"]}
            st.session_state.session_token = token
            st.session_state.user_plan = user_info.get("plan", "free")
            st.session_state.is_god = user_info.get("plan") == "god"
            return True
    return False

def get_user_email():
    return st.session_state.user.get("email", "User") if st.session_state.user else ("GOD" if st.session_state.is_god else "User")

def build_app_url(app_name):
    base = APP_URLS.get(app_name, "")
    token = st.session_state.get("session_token", "")
    return f"{base}?token={token}" if base and token else base

# ============================================
# DATABASE FUNCTIONS
# ============================================

def get_jd_history(limit=50):
    user_id = st.session_state.user.get("id") if st.session_state.user else None
    if not user_id or user_id == "god": return []
    try:
        r = requests.get(f"{SUPABASE_URL}/rest/v1/jd_history?user_id=eq.{user_id}&order=created_at.desc&limit={limit}",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}"}, timeout=10)
        return r.json() if r.status_code == 200 else []
    except: return []

def save_screen_history(jd_id, jd_text, candidates_input, result, candidates_count, tokens_used):
    user_id = st.session_state.user.get("id") if st.session_state.user else None
    if not user_id or user_id == "god": return None
    try:
        r = requests.post(f"{SUPABASE_URL}/rest/v1/screen_history",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                     "Content-Type": "application/json", "Prefer": "return=representation"},
            json={"user_id": user_id, "jd_history_id": jd_id, "job_description": jd_text[:5000] if jd_text else None,
                  "candidates_input": candidates_input[:5000] if candidates_input else None, "screening_result": result,
                  "candidates_count": candidates_count, "tokens_used": tokens_used}, timeout=10)
        return r.json()[0] if r.status_code in [200, 201] and r.json() else None
    except: return None

def submit_feedback(app, feedback_type, message):
    try:
        requests.post(f"{SUPABASE_URL}/rest/v1/user_feedback",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                     "Content-Type": "application/json", "Prefer": "return=minimal"},
            json={"user_id": st.session_state.user.get("id") if st.session_state.user else None,
                  "app": app, "feedback_type": feedback_type, "rating": 4, "message": message, "email": get_user_email()}, timeout=5)
        return True
    except: return False

# ============================================
# FILE PROCESSING
# ============================================

def is_readable_text(text):
    """Check if extracted text is actually readable"""
    if not text or len(text.strip()) < 10:
        return False
    alpha_chars = sum(1 for c in text if c.isalpha())
    alpha_ratio = alpha_chars / len(text) if len(text) > 0 else 0
    return alpha_ratio > 0.3

def clean_extracted_text(text):
    """Clean up extracted text"""
    if not text:
        return ""
    cleaned = ''.join(c if c.isprintable() or c in '\n\r\t' else ' ' for c in text)
    cleaned = re.sub(r'[ \t]+', ' ', cleaned)
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
    return cleaned.strip()

def extract_text_from_file(uploaded_file):
    file_type = uploaded_file.name.split('.')[-1].lower()
    content = uploaded_file.read()
    uploaded_file.seek(0)
    
    if file_type == 'txt':
        return clean_extracted_text(content.decode('utf-8', errors='ignore'))
    
    elif file_type == 'pdf':
        extracted_text = ""
        
        # Try PyMuPDF with multiple methods
        try:
            import fitz
            pdf = fitz.open(stream=content, filetype="pdf")
            text_parts = []
            for page in pdf:
                page_text = page.get_text("text")
                if not page_text or not page_text.strip():
                    page_text = page.get_text("text", flags=fitz.TEXT_PRESERVE_WHITESPACE)
                if not page_text or not page_text.strip():
                    blocks = page.get_text("blocks")
                    page_text = "\n".join([b[4] for b in blocks if b[6] == 0])
                if page_text and page_text.strip():
                    text_parts.append(page_text)
            pdf.close()
            extracted_text = "\n".join(text_parts)
        except:
            pass
        
        if extracted_text and extracted_text.strip():
            cleaned = clean_extracted_text(extracted_text)
            if is_readable_text(cleaned):
                return cleaned
        
        # Try pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text_parts = [page.extract_text() for page in pdf.pages if page.extract_text()]
                extracted_text = "\n".join(text_parts)
                if extracted_text:
                    cleaned = clean_extracted_text(extracted_text)
                    if is_readable_text(cleaned):
                        return cleaned
        except:
            pass
        
        # Last resort
        try:
            import fitz
            pdf = fitz.open(stream=content, filetype="pdf")
            text_parts = [page.get_text() for page in pdf]
            pdf.close()
            result = clean_extracted_text("\n".join(text_parts))
            if len(result) > 50:
                return result
        except:
            pass
        
        return "[PDF extraction failed - please paste content directly]"
    
    elif file_type in ['docx', 'doc']:
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            if paragraphs:
                return clean_extracted_text('\n\n'.join(paragraphs))
        except:
            pass
        
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as z:
                if 'word/document.xml' in z.namelist():
                    xml_content = z.read('word/document.xml')
                    tree = ElementTree.fromstring(xml_content)
                    texts = [elem.text for elem in tree.iter() if elem.text and elem.text.strip()]
                    if texts:
                        return clean_extracted_text(' '.join(texts))
        except:
            pass
        
        return "[DOCX extraction failed - please paste content directly]"
    elif file_type == 'json':
        try:
            data = json.loads(content.decode('utf-8'))
            if isinstance(data, list):
                return '\n---\n'.join([json.dumps(item, indent=2) for item in data])
            return json.dumps(data, indent=2)
        except:
            return content.decode('utf-8', errors='ignore')
    return content.decode('utf-8', errors='ignore')

def parse_json_candidates(json_text):
    try:
        data = json.loads(json_text)
        candidates = []
        items = data if isinstance(data, list) else data.get('candidates', [data])
        for item in items:
            candidates.append({
                'name': item.get('name') or item.get('candidateName') or item.get('full_name', 'Unknown'),
                'email': item.get('email') or item.get('emailAddress', ''),
                'skills': item.get('skills') or item.get('skillSet', []),
                'summary': item.get('summary') or item.get('professionalSummary', ''),
            })
        return candidates
    except:
        return None

# ============================================
# CLAUDE API
# ============================================

def call_claude(prompt, max_tokens=4000, action="screen"):
    if not ANTHROPIC_API_KEY:
        return "Error: ANTHROPIC_API_KEY not set", 0
    try:
        r = requests.post("https://api.anthropic.com/v1/messages",
            headers={"x-api-key": ANTHROPIC_API_KEY, "Content-Type": "application/json", "anthropic-version": "2023-06-01"},
            json={"model": "claude-sonnet-4-20250514", "max_tokens": max_tokens, "messages": [{"role": "user", "content": prompt}]}, timeout=180)
        if r.status_code == 200:
            text = r.json()["content"][0]["text"]
            tokens = (len(prompt) + len(text)) // 4
            if st.session_state.user:
                log_usage(st.session_state.user.get("id"), st.session_state.get("session_token"), "screen", action, tokens)
            return text, tokens
        return f"Error: {r.status_code}", 0
    except Exception as e:
        return f"Error: {str(e)}", 0

# ============================================
# SCREENING FUNCTIONS
# ============================================

def screen_candidates(jd_text, cvs_text, options):
    prompt = f"""You are an expert technical recruiter. Analyze candidates against this JD with DETAILED categorized skill breakdowns.

## JOB DESCRIPTION:
{jd_text}

## CANDIDATES (separated by ---):
{cvs_text}

## OPTIONS:
- Bias-Free Mode: {options.get('bias_free', True)}
- AI Detection: {options.get('ai_detection', True)}
- Salary Estimate: {options.get('salary_estimate', True)}

## OUTPUT (JSON):
```json
{{
    "screening_summary": {{
        "total_candidates": <number>,
        "recommended_for_interview": <number>,
        "maybe": <number>,
        "not_recommended": <number>,
        "time_saved_minutes": <number>,
        "bias_mitigation_score": <0-100>
    }},
    "candidates": [
        {{
            "rank": <number>,
            "identifier": "<name or Candidate A/B/C>",
            "match_score": <0-100>,
            "skills_breakdown": {{
                "technical": {{"matched": <n>, "total": <n>, "details": ["skill1"]}},
                "leadership": {{"matched": <n>, "total": <n>, "details": []}},
                "soft_skills": {{"matched": <n>, "total": <n>, "details": []}},
                "industry_experience": {{"matched": <n>, "total": <n>, "details": []}}
            }},
            "experience_years": <number>,
            "salary_alignment": "<Within Range | Above | Below | Unknown>",
            "why_this_score": "<2-3 sentences>",
            "strengths": ["<str1>", "<str2>"],
            "concerns": ["<concern1>"],
            "ai_written_probability": <0-100>,
            "recommended_action": "<Phone Screen | Assessment | Reject>",
            "rejection_reason": "<reason if rejected>",
            "next_steps": "<steps>",
            "hiring_manager_summary": "<3 bullet points>"
        }}
    ],
    "batch_insights": {{
        "strongest_candidate_summary": "<summary>",
        "common_gaps": ["<gap1>"],
        "hiring_recommendation": "<recommendation>"
    }}
}}
```"""
    return call_claude(prompt, max_tokens=6000, action="screen_batch")

def compare_candidates(candidates_data, jd_summary=""):
    prompt = f"""Create a side-by-side comparison:

## CANDIDATES:
{json.dumps(candidates_data, indent=2)}

## OUTPUT (JSON):
```json
{{
    "comparison_table": {{
        "headers": ["Metric", "Candidate 1", "Candidate 2", "Candidate 3"],
        "rows": [
            ["Experience", "<val>", "<val>", "<val>"],
            ["Match Score", "<val>%", "<val>%", "<val>%"],
            ["Technical", "<X/Y>", "<X/Y>", "<X/Y>"],
            ["Leadership", "<X/Y>", "<X/Y>", "<X/Y>"],
            ["Salary", "<status>", "<status>", "<status>"],
            ["AI Score", "<val>%", "<val>%", "<val>%"]
        ]
    }},
    "winner_recommendation": "<which and why>",
    "interview_strategy": "<how to differentiate>"
}}
```"""
    return call_claude(prompt, max_tokens=2000, action="compare")

def generate_email_template(candidate, template_type, job_title=""):
    templates = {
        "phone_screen": f"Draft phone screen invite for {candidate.get('identifier')} for {job_title}. Include: personalized line about their strengths ({', '.join(candidate.get('strengths', [])[:2])}), request for 20-30 min call, [CALENDAR_LINK] placeholder.",
        "skills_test": f"Draft skills assessment email for {candidate.get('identifier')}. Areas to validate: {', '.join(candidate.get('concerns', [])[:2])}. Professional, encouraging tone.",
        "rejection": f"Draft GDPR-compliant rejection for {candidate.get('identifier')}. Reason: {candidate.get('rejection_reason', 'skills mismatch')}. Kind, professional.",
        "linkedin": f"Draft SHORT LinkedIn message (<300 chars) for {candidate.get('identifier')}. Reference: {candidate.get('strengths', ['experience'])[0]}."
    }
    return call_claude(templates.get(template_type, templates["phone_screen"]), max_tokens=800, action=f"email_{template_type}")

def generate_hiring_manager_summary(candidates_data, bias_free=True):
    prompt = f"""Generate a Hiring Manager briefing summary. Return ONLY the formatted summary, NOT JSON.

## CANDIDATES:
{json.dumps(candidates_data, indent=2)}

{"Remove all PII (use Candidate A, B, C)" if bias_free else "Include names"}

## FORMAT YOUR RESPONSE LIKE THIS:

EXECUTIVE SUMMARY
[1-2 sentence overview of the candidate pool]

CANDIDATE RECOMMENDATIONS

[For each candidate:]
CANDIDATE [A/B/C or Name]
• Fit: [What makes them a good/poor fit]
• Gaps: [Key concerns or missing skills]  
• Action: [Recommended next step]
• Interview Focus: [What to probe in interview]

OVERALL RECOMMENDATION
[Which candidate(s) to prioritize and why]

RISK ASSESSMENT
[Any hiring risks to consider]

Write in clear, professional prose. No JSON."""
    return call_claude(prompt, max_tokens=2000, action="hm_summary")

def anonymize_cv(cv_text, options):
    prompt = f"""Anonymize this CV:

{cv_text}

{"Convert dates to relative time" if options.get('time_based') else "Keep dates"}
Replace universities with tiers, companies with descriptors.

## OUTPUT (JSON):
```json
{{
    "anonymized_cv": "<text>",
    "items_removed": {{"names": [], "emails": [], "phones": [], "universities": [], "companies": []}},
    "anonymization_score": <0-100>
}}
```"""
    return call_claude(prompt, max_tokens=3000, action="anonymize")

def estimate_salary(cv_text, jd_text, location):
    prompt = f"""Estimate salary:

JOB: {jd_text[:2000]}
CV: {cv_text[:3000]}
LOCATION: {location}

## OUTPUT (JSON):
```json
{{
    "estimated_range": {{"low": <num>, "mid": <num>, "high": <num>}},
    "confidence": "<Low|Medium|High>",
    "factors_increasing": ["<factor>"],
    "factors_decreasing": ["<factor>"],
    "market_context": "<note>"
}}
```"""
    return call_claude(prompt, max_tokens=1500, action="salary")

# ============================================
# PDF GENERATION (Real PDF using reportlab)
# ============================================

def generate_pdf_bytes(data, title="CV Screening Report"):
    """Generate actual PDF file using reportlab"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        
        if isinstance(data, str):
            data = json.loads(data)
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle('CustomTitle', parent=styles['Title'], fontSize=24, textColor=colors.HexColor('#6366f1'), spaceAfter=12)
        heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading1'], fontSize=16, textColor=colors.HexColor('#1f2937'), spaceBefore=20, spaceAfter=10)
        subheading_style = ParagraphStyle('CustomSubheading', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#6366f1'), spaceBefore=15, spaceAfter=8)
        body_style = ParagraphStyle('CustomBody', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#374151'), spaceAfter=6)
        small_style = ParagraphStyle('Small', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#6b7280'))
        green_style = ParagraphStyle('Green', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#059669'))
        red_style = ParagraphStyle('Red', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#dc2626'))
        score_high = ParagraphStyle('ScoreHigh', parent=styles['Normal'], fontSize=20, textColor=colors.HexColor('#10b981'), alignment=TA_CENTER)
        score_med = ParagraphStyle('ScoreMed', parent=styles['Normal'], fontSize=20, textColor=colors.HexColor('#f59e0b'), alignment=TA_CENTER)
        score_low = ParagraphStyle('ScoreLow', parent=styles['Normal'], fontSize=20, textColor=colors.HexColor('#ef4444'), alignment=TA_CENTER)
        
        story = []
        
        # Header
        story.append(Paragraph("📊 CV Screening Report", title_style))
        story.append(Paragraph(f"{title} | Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", small_style))
        story.append(Spacer(1, 20))
        story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor('#6366f1')))
        story.append(Spacer(1, 20))
        
        summary = data.get('screening_summary', {})
        candidates = data.get('candidates', [])
        insights = data.get('batch_insights', {})
        
        # Summary table
        story.append(Paragraph("Executive Summary", heading_style))
        
        summary_data = [
            ['Total Candidates', 'Recommended', 'Maybe', 'Not Recommended', 'Time Saved'],
            [
                str(summary.get('total_candidates', '?')),
                str(summary.get('recommended_for_interview', '?')),
                str(summary.get('maybe', '?')),
                str(summary.get('not_recommended', '?')),
                f"~{summary.get('time_saved_minutes', '?')} min"
            ]
        ]
        
        summary_table = Table(summary_data, colWidths=[1.4*inch]*5)
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#6366f1')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('FONTSIZE', (0, 1), (-1, -1), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('TOPPADDING', (0, 1), (-1, -1), 15),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 15),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f3f4f6')),
            ('TEXTCOLOR', (1, 1), (1, 1), colors.HexColor('#10b981')),
            ('TEXTCOLOR', (2, 1), (2, 1), colors.HexColor('#f59e0b')),
            ('TEXTCOLOR', (3, 1), (3, 1), colors.HexColor('#ef4444')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb')),
            ('ROUNDEDCORNERS', [5, 5, 5, 5]),
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 30))
        
        # Candidates
        story.append(Paragraph("Candidate Rankings", heading_style))
        
        for c in candidates:
            score = c.get('match_score', 0)
            skills = c.get('skills_breakdown', {})
            
            # Candidate header
            story.append(Paragraph(f"#{c.get('rank', '?')} — {c.get('identifier', 'Unknown')}", subheading_style))
            
            # Score and skills table
            score_color = '#10b981' if score >= 70 else '#f59e0b' if score >= 50 else '#ef4444'
            
            skills_data = [
                ['Match Score', 'Technical', 'Leadership', 'Soft Skills', 'Experience'],
                [
                    f"{score}%",
                    f"{skills.get('technical', {}).get('matched', '?')}/{skills.get('technical', {}).get('total', '?')}",
                    f"{skills.get('leadership', {}).get('matched', '?')}/{skills.get('leadership', {}).get('total', '?')}",
                    f"{skills.get('soft_skills', {}).get('matched', '?')}/{skills.get('soft_skills', {}).get('total', '?')}",
                    f"{c.get('experience_years', '?')} yrs"
                ]
            ]
            
            skills_table = Table(skills_data, colWidths=[1.4*inch]*5)
            skills_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f9fafb')),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('FONTSIZE', (0, 1), (-1, -1), 12),
                ('TEXTCOLOR', (0, 1), (0, 1), colors.HexColor(score_color)),
                ('FONTNAME', (0, 1), (0, 1), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb')),
            ]))
            story.append(skills_table)
            story.append(Spacer(1, 10))
            
            # Assessment
            story.append(Paragraph(f"<b>Assessment:</b> {c.get('why_this_score', 'N/A')}", body_style))
            
            # Strengths
            strengths = c.get('strengths', [])
            if strengths:
                story.append(Paragraph(f"<b>✓ Strengths:</b> {', '.join(strengths)}", green_style))
            
            # Concerns
            concerns = c.get('concerns', [])
            if concerns:
                story.append(Paragraph(f"<b>⚠ Concerns:</b> {', '.join(concerns)}", red_style))
            
            # Recommendation
            story.append(Paragraph(f"<b>→ Recommendation:</b> {c.get('recommended_action', 'N/A')}", body_style))
            story.append(Paragraph(f"<b>Next Steps:</b> {c.get('next_steps', 'N/A')}", small_style))
            
            story.append(Spacer(1, 5))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e5e7eb')))
            story.append(Spacer(1, 15))
        
        # Insights
        if insights:
            story.append(Paragraph("Batch Insights", heading_style))
            story.append(Paragraph(f"<b>Top Candidate:</b> {insights.get('strongest_candidate_summary', 'N/A')}", body_style))
            story.append(Paragraph(f"<b>Common Gaps:</b> {', '.join(insights.get('common_gaps', []))}", body_style))
            story.append(Paragraph(f"<b>Hiring Recommendation:</b> {insights.get('hiring_recommendation', 'N/A')}", body_style))
        
        # Footer
        story.append(Spacer(1, 40))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e5e7eb')))
        story.append(Spacer(1, 10))
        story.append(Paragraph("Generated by Sharp Screen | Sharp Human AI Recruiting Suite", small_style))
        story.append(Paragraph("This report is confidential and intended for internal use only.", small_style))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    except ImportError:
        # Fallback: Return HTML if reportlab not installed
        return None
    except Exception as e:
        return None

def generate_csv_report(data):
    try:
        if isinstance(data, str):
            data = json.loads(data)
        candidates = data.get('candidates', [])
        lines = ["Rank,Identifier,Score,Technical,Leadership,Soft Skills,Experience,Salary,AI Score,Action"]
        for c in candidates:
            sk = c.get('skills_breakdown', {})
            lines.append(f"{c.get('rank','')},{c.get('identifier','')},{c.get('match_score','')},{sk.get('technical',{}).get('matched','?')}/{sk.get('technical',{}).get('total','?')},{sk.get('leadership',{}).get('matched','?')}/{sk.get('leadership',{}).get('total','?')},{sk.get('soft_skills',{}).get('matched','?')}/{sk.get('soft_skills',{}).get('total','?')},{c.get('experience_years','')},{c.get('salary_alignment','')},{c.get('ai_written_probability','')},{c.get('recommended_action','')}")
        return '\n'.join(lines)
    except:
        return "Error"

# ============================================
# DISPLAY HELPERS
# ============================================

def display_candidate_rich(c):
    """Display a candidate with rich formatting"""
    score = c.get('match_score', 0)
    skills = c.get('skills_breakdown', {})
    identifier = c.get('identifier', 'Unknown')
    
    # Score color
    if score >= 70:
        score_color = "#10b981"
        score_bg = "rgba(16,185,129,0.1)"
        recommendation_badge = "🟢 Recommended"
    elif score >= 50:
        score_color = "#f59e0b"
        score_bg = "rgba(245,158,11,0.1)"
        recommendation_badge = "🟡 Maybe"
    else:
        score_color = "#ef4444"
        score_bg = "rgba(239,68,68,0.1)"
        recommendation_badge = "🔴 Not Recommended"
    
    # Main card
    st.markdown(f"""
    <div style="background:#12121a;border:1px solid rgba(99,102,241,0.3);border-radius:16px;padding:24px;margin:20px 0;">
        <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:20px;">
            <div style="display:flex;align-items:center;gap:12px;">
                <span style="background:linear-gradient(135deg,#6366f1,#8b5cf6);color:white;padding:8px 16px;border-radius:20px;font-weight:bold;">#{c.get('rank', '?')}</span>
                <span style="font-size:20px;font-weight:bold;color:white;">{identifier}</span>
                <span style="background:{score_bg};color:{score_color};padding:4px 12px;border-radius:12px;font-size:12px;">{recommendation_badge}</span>
            </div>
            <div style="text-align:right;">
                <span style="font-size:36px;font-weight:bold;color:{score_color};">{score}%</span>
                <p style="color:#9ca3af;margin:0;font-size:12px;">Match Score</p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Skills breakdown in columns
    sk_cols = st.columns(5)
    skill_items = [
        ("💻 Technical", skills.get('technical', {})),
        ("👔 Leadership", skills.get('leadership', {})),
        ("🤝 Soft Skills", skills.get('soft_skills', {})),
        ("🏢 Industry", skills.get('industry_experience', {})),
        ("📅 Experience", {"value": f"{c.get('experience_years', '?')} yrs"})
    ]
    
    for col, (label, skill_data) in zip(sk_cols, skill_items):
        with col:
            if 'value' in skill_data:
                value = skill_data['value']
            else:
                matched = skill_data.get('matched', '?')
                total = skill_data.get('total', '?')
                value = f"{matched}/{total}"
            
            st.markdown(f"""
            <div style="background:rgba(99,102,241,0.1);border-radius:12px;padding:16px;text-align:center;">
                <p style="color:#9ca3af;margin:0;font-size:11px;">{label}</p>
                <p style="color:white;font-size:20px;font-weight:bold;margin:4px 0;">{value}</p>
            </div>
            """, unsafe_allow_html=True)
    
    # Assessment section
    st.markdown("#### 📋 Assessment")
    st.markdown(f"**Why this score:** {c.get('why_this_score', 'N/A')}")
    
    # Strengths and Concerns in columns
    str_col, con_col = st.columns(2)
    with str_col:
        st.markdown("**✅ Strengths**")
        for s in c.get('strengths', []):
            st.markdown(f"<span style='background:rgba(16,185,129,0.2);color:#10b981;padding:4px 10px;border-radius:8px;margin:2px;display:inline-block;font-size:13px;'>{s}</span>", unsafe_allow_html=True)
    
    with con_col:
        st.markdown("**⚠️ Concerns**")
        for s in c.get('concerns', []):
            st.markdown(f"<span style='background:rgba(239,68,68,0.2);color:#ef4444;padding:4px 10px;border-radius:8px;margin:2px;display:inline-block;font-size:13px;'>{s}</span>", unsafe_allow_html=True)
    
    # Additional info
    info_col1, info_col2 = st.columns(2)
    with info_col1:
        st.markdown(f"**💰 Salary Alignment:** {c.get('salary_alignment', 'Unknown')}")
    with info_col2:
        ai_prob = c.get('ai_written_probability', 0)
        ai_color = "#ef4444" if ai_prob > 50 else "#f59e0b" if ai_prob > 30 else "#10b981"
        st.markdown(f"**🤖 AI Detection:** <span style='color:{ai_color}'>{ai_prob}%</span>", unsafe_allow_html=True)
    
    # Recommendation box
    st.markdown(f"""
    <div style="background:linear-gradient(135deg,rgba(99,102,241,0.1),rgba(139,92,246,0.1));border-left:4px solid #6366f1;padding:16px;border-radius:0 12px 12px 0;margin:16px 0;">
        <p style="color:#a5b4fc;margin:0 0 4px 0;font-size:12px;text-transform:uppercase;">Recommended Action</p>
        <p style="color:white;font-size:16px;font-weight:bold;margin:0;">{c.get('recommended_action', 'N/A')}</p>
        <p style="color:#9ca3af;margin:8px 0 0 0;font-size:13px;">{c.get('next_steps', '')}</p>
    </div>
    """, unsafe_allow_html=True)
    
    return identifier, score

def display_hm_summary_rich(summary_text):
    """Display hiring manager summary with rich formatting"""
    if not summary_text:
        return
    
    # Clean up any JSON artifacts
    text = summary_text.strip()
    if text.startswith('```'):
        text = re.sub(r'```json?\s*', '', text)
        text = re.sub(r'```\s*$', '', text)
    
    # Try to parse as JSON first
    try:
        data = json.loads(text)
        # Format JSON data nicely
        st.markdown("#### 📊 Executive Summary")
        st.info(data.get('executive_summary', ''))
        
        st.markdown("#### 👥 Candidate Recommendations")
        for cand in data.get('candidates', []):
            with st.container():
                st.markdown(f"**{cand.get('identifier', 'Candidate')}**")
                st.markdown(f"• **Fit:** {cand.get('fit', cand.get('fit_assessment', 'N/A'))}")
                st.markdown(f"• **Gaps:** {cand.get('gaps', cand.get('key_gaps', 'N/A'))}")
                st.markdown(f"• **Action:** {cand.get('action', cand.get('recommended_action', 'N/A'))}")
                st.markdown("---")
        
        st.markdown("#### 🎯 Overall Recommendation")
        st.success(data.get('overall_recommendation', ''))
        
        if data.get('risk_assessment'):
            st.markdown("#### ⚠️ Risk Assessment")
            st.warning(data.get('risk_assessment', ''))
    except:
        # Display as formatted text
        st.markdown(f"""
        <div style="background:#12121a;border:1px solid rgba(99,102,241,0.2);border-radius:12px;padding:24px;">
            <pre style="white-space:pre-wrap;color:#e5e5e5;font-family:inherit;margin:0;">{text}</pre>
        </div>
        """, unsafe_allow_html=True)

# ============================================
# MAIN APP
# ============================================

st.set_page_config(page_title="Sharp Screen", page_icon="🔍", layout="wide")
init_session()
check_url_auth()

# Apply shared UI if available
if USING_SHARED:
    inject_ga4()
    apply_global_styles()
    render_top_banner(show_cta=True, cta_text="Book a Demo")

# Styles
st.markdown("""<style>
@import url('https://fonts.googleapis.com/css2?family=Nunito:wght@400;600;700&display=swap');

*, *::before, *::after { font-family: 'Nunito', sans-serif !important; box-sizing: border-box; }

.stApp, [data-testid="stAppViewContainer"] { background: #0a0a0f !important; }
[data-testid="stHeader"] { background: transparent !important; }

section[data-testid="stSidebar"] { background: #0d0d14 !important; border-right: 1px solid rgba(99,102,241,0.2); }
section[data-testid="stSidebar"] > div { background: #0d0d14 !important; }
section[data-testid="stSidebar"] * { color: #e5e5e5 !important; }

/* Hide orphaned text */
section[data-testid="stSidebar"] > div > div:first-child > div:first-child { display: none !important; }

h1, h2, h3, h4, h5, h6 { color: #ffffff !important; }
p, span, label, div, li { color: #e5e5e5; }

.stTextInput > div > div > input,
.stTextArea > div > div > textarea,
.stSelectbox > div > div,
.stMultiSelect > div > div,
[data-baseweb="select"] > div { background: #12121a !important; border: 1px solid rgba(99,102,241,0.3) !important; color: #ffffff !important; border-radius: 8px !important; }

[data-baseweb="tag"] { background: rgba(99,102,241,0.3) !important; color: white !important; }

.stButton > button { background: linear-gradient(135deg, #6366f1, #8b5cf6) !important; color: white !important; border: none !important; border-radius: 8px !important; font-weight: 600 !important; }
.stDownloadButton > button { background: #1a1a2e !important; border: 1px solid rgba(99,102,241,0.3) !important; color: white !important; }

.stTabs [data-baseweb="tab-list"] { background: transparent !important; gap: 8px; border-bottom: 1px solid rgba(99,102,241,0.2); }
.stTabs [data-baseweb="tab"] { background: transparent !important; color: #9ca3af !important; border: none !important; border-bottom: 2px solid transparent; padding: 12px 20px !important; }
.stTabs [aria-selected="true"] { background: transparent !important; color: #fff !important; border-bottom: 2px solid #6366f1 !important; }

.stRadio > div { flex-direction: row !important; gap: 12px; flex-wrap: wrap; }
.stRadio > div > label { background: #12121a !important; padding: 10px 16px !important; border-radius: 8px !important; border: 1px solid rgba(99,102,241,0.2) !important; }

.stCheckbox > label { color: #e5e5e5 !important; }

[data-testid="stFileUploader"] { background: #12121a !important; border: 1px dashed rgba(99,102,241,0.3) !important; border-radius: 8px !important; padding: 20px !important; }

[data-baseweb="popover"], [data-baseweb="menu"] { background: #12121a !important; }
[role="option"] { color: #e5e5e5 !important; }
[role="option"]:hover { background: rgba(99,102,241,0.2) !important; }

.stSlider > div > div { background: rgba(99,102,241,0.3) !important; }

[data-testid="stMetricValue"] { color: #fff !important; }
[data-testid="stMetricLabel"] { color: #9ca3af !important; }

.stSuccess { background: rgba(16,185,129,0.1) !important; border: 1px solid #10b981 !important; }
.stError { background: rgba(239,68,68,0.1) !important; border: 1px solid #ef4444 !important; }
.stWarning { background: rgba(245,158,11,0.1) !important; border: 1px solid #f59e0b !important; }
.stInfo { background: rgba(99,102,241,0.1) !important; border: 1px solid #6366f1 !important; }

.user-card { background: rgba(99,102,241,0.1); border-radius: 12px; padding: 16px; margin-bottom: 20px; }
.metric-card { background: linear-gradient(135deg, rgba(99,102,241,0.1), rgba(139,92,246,0.1)); border: 1px solid rgba(99,102,241,0.2); border-radius: 12px; padding: 20px; text-align: center; }
.output-box { background: #12121a; border: 1px solid rgba(99,102,241,0.2); border-radius: 12px; padding: 20px; margin: 16px 0; white-space: pre-wrap; color: #e5e5e5; }
.status-badge { position: fixed; top: 70px; right: 20px; background: linear-gradient(135deg, #6366f1, #8b5cf6); color: white; padding: 10px 20px; border-radius: 25px; font-weight: 600; z-index: 999; animation: pulse 2s infinite; }
@keyframes pulse { 0%, 100% { opacity: 1; } 50% { opacity: 0.7; } }
</style>""", unsafe_allow_html=True)

# Auth
if not st.session_state.authenticated:
    c1, c2, c3 = st.columns([1, 1.2, 1])
    with c2:
        st.markdown("""<div style="text-align:center;padding:60px 0 30px;">
            <img src="https://sharphuman.com/logo1-3.png" style="width:80px;margin-bottom:20px;">
            <h1 style="margin:0;">Sharp Screen</h1>
            <p style="color:#9ca3af;">AI CV Screening & Analysis</p>
        </div>""", unsafe_allow_html=True)
        
        t1, t2 = st.tabs(["Log In", "Sign Up"])
        with t1:
            email = st.text_input("Email", key="l_email")
            pwd = st.text_input("Password", type="password", key="l_pwd")
            if st.button("Log In", use_container_width=True):
                if pwd == GOD_PASSWORD:
                    st.session_state.authenticated = True
                    st.session_state.is_god = True
                    st.session_state.user = {"email": "GOD", "id": "god"}
                    st.session_state.session_token = secrets.token_urlsafe(32)
                    st.rerun()
                elif email and pwd:
                    r = supabase_sign_in(email, pwd)
                    if r["success"]:
                        st.session_state.authenticated = True
                        st.session_state.user = r["user"]
                        st.session_state.session_token = r.get("session_token")
                        st.rerun()
                    else:
                        st.error(r["message"])
        with t2:
            s_email = st.text_input("Email", key="s_email")
            s_pwd = st.text_input("Password", type="password", key="s_pwd")
            s_conf = st.text_input("Confirm", type="password", key="s_conf")
            if st.button("Create Account", use_container_width=True):
                if s_pwd != s_conf: st.error("Passwords don't match")
                elif len(s_pwd) < 6: st.warning("6+ characters")
                elif s_email and s_pwd:
                    r = supabase_sign_up(s_email, s_pwd)
                    st.success(r["message"]) if r["success"] else st.error(r["message"])
    st.stop()

# ============================================
# AUTHENTICATED UI
# ============================================

if st.session_state.working_on:
    st.markdown(f'<div class="status-badge">{st.session_state.working_on}</div>', unsafe_allow_html=True)

# Sidebar with ICONS
# Sidebar - Use shared UI if available
if USING_SHARED:
    render_sidebar(
        current_app="screen",
        user_email=get_user_email(),
        user_plan=st.session_state.get('user_plan', 'free'),
        session_token=st.session_state.get('session_token', '')
    )
else:
    with st.sidebar:
        st.markdown(f"""<div class="user-card">
            <p style="color:#9ca3af;margin:0;font-size:12px;">Logged in as</p>
            <p style="color:#fff;margin:4px 0;font-weight:600;">{get_user_email()}</p>
            <p style="color:#6366f1;margin:0;font-size:12px;text-transform:uppercase;">{st.session_state.get('user_plan','free')} plan</p>
        </div>""", unsafe_allow_html=True)
        
        st.markdown("**Apps**")
        apps = [
            ("portal", "🏠 Portal"), ("jd", "📝 JD Writer"), ("screen", "🔍 CV Screener"),
            ("interview", "🎯 Interview"), ("outreach", "🚀 Outreach"), ("content", "✍️ Content"),
            ("sales", "💰 Sales"),
        ]
        for key, label in apps:
            if key == "screen":
                st.markdown(f"<div style='background:linear-gradient(135deg,#6366f1,#8b5cf6);padding:10px 16px;border-radius:8px;text-align:center;margin:4px 0;color:white;font-weight:600;'>{label} ◀</div>", unsafe_allow_html=True)
            else:
                st.link_button(label, build_app_url(key), use_container_width=True)
        if st.session_state.get("is_god"):
            st.markdown("---")
            st.link_button("⚙️ Admin", build_app_url("admin"), use_container_width=True)
        st.markdown("---")
        if st.button("🚪 Logout", use_container_width=True):
            for k in list(st.session_state.keys()): del st.session_state[k]
            st.rerun()

# Header
st.markdown("""<div style="display:flex;align-items:center;gap:16px;padding:20px 0;border-bottom:1px solid rgba(99,102,241,0.2);margin-bottom:30px;">
    <img src="https://sharphuman.com/logo1-3.png" style="width:50px;">
    <div><h1 style="margin:0;font-size:28px;">Sharp Screen</h1><p style="color:#9ca3af;margin:0;">AI-Powered CV Screening</p></div>
</div>""", unsafe_allow_html=True)

# Tabs
tab_screen, tab_blind, tab_salary = st.tabs(["🔍 Screen & Rank", "👤 Blind Resume", "💰 Salary Estimator"])

with tab_screen:
    # Input section
    st.markdown("### 📋 Job Description")
    jd_src = st.radio("Source:", ["📝 Paste", "📜 History", "📄 Upload"], horizontal=True, key="jd_src")
    
    jd_text = ""
    jd_id = None
    
    if jd_src == "📜 History":
        history = get_jd_history(20)
        if history:
            opts = {f"{j['job_title']} ({j['created_at'][:10]})": j for j in history}
            sel = st.selectbox("Select:", list(opts.keys()))
            if sel:
                jd_text = opts[sel].get('generated_jd', '')
                jd_id = opts[sel].get('id')
                st.success(f"✅ Loaded: {opts[sel].get('job_title')}")
        else:
            st.info("No saved JDs")
            jd_text = st.text_area("Paste JD:", height=120)
    elif jd_src == "📄 Upload":
        f = st.file_uploader("Upload JD", type=['pdf', 'docx', 'txt'])
        if f:
            jd_text = extract_text_from_file(f)
            st.success(f"✅ Loaded {f.name}")
    else:
        jd_text = st.text_area("Paste Job Description:", height=150, 
            placeholder="Paste the full job description here including:\n• Required qualifications & skills\n• Responsibilities\n• Experience level\n• Salary range (if available)")
    
    st.markdown("---")
    
    st.markdown("### 📄 Candidate CVs")
    cv_src = st.radio("Source:", ["📝 Paste", "📁 Upload", "🔗 JSON/ATS"], horizontal=True, key="cv_src")
    
    cvs_text = ""
    if cv_src == "📁 Upload":
        files = st.file_uploader("Upload CVs", type=['pdf', 'docx', 'txt'], accept_multiple_files=True)
        if files:
            st.success(f"📎 {len(files)} files uploaded")
            cvs_text = "\n\n---\n\n".join([f"=== {f.name} ===\n{extract_text_from_file(f)}" for f in files])
    elif cv_src == "🔗 JSON/ATS":
        json_in = st.text_area("Paste JSON:", height=120)
        if json_in:
            cands = parse_json_candidates(json_in)
            if cands:
                st.success(f"✅ {len(cands)} candidates parsed")
                cvs_text = "\n\n---\n\n".join([json.dumps(c, indent=2) for c in cands])
    else:
        cvs_text = st.text_area("Paste CVs (separate with ---):", height=180, 
            placeholder="Paste candidate resumes here, separated by ---\n\nExample:\nJohn Smith - Software Engineer\n5 years Python, AWS...\n\n---\n\nJane Doe - Backend Developer\n3 years Node.js, PostgreSQL...")
    
    st.markdown("---")
    
    st.markdown("### ⚙️ Options")
    c1, c2, c3, c4 = st.columns(4)
    with c1: bias_free = st.checkbox("🔒 Bias-Free", value=True)
    with c2: ai_detect = st.checkbox("🤖 AI Detection", value=True)
    with c3: salary_est = st.checkbox("💰 Salary Est.", value=True)
    with c4: save_hist = st.checkbox("💾 Save Results", value=True)
    
    if st.button("🔍 Screen Candidates", type="primary", use_container_width=True):
        if not jd_text: st.warning("Provide JD")
        elif not cvs_text: st.warning("Provide CVs")
        else:
            st.session_state.working_on = "Analyzing candidates..."
            with st.spinner("Screening..."):
                result, tokens = screen_candidates(jd_text, cvs_text, {'bias_free': bias_free, 'ai_detection': ai_detect, 'salary_estimate': salary_est})
                st.session_state.working_on = None
                if not result.startswith("Error"):
                    st.session_state.screening_results = result
                    st.session_state.selected_candidates = []
                    st.session_state.hm_summary = None
                    if save_hist:
                        save_screen_history(jd_id, jd_text, cvs_text, result, cvs_text.count('---')+1, tokens)
                    st.rerun()
                else:
                    st.error(result)
    
    # Results
    if st.session_state.screening_results:
        st.markdown("---")
        
        try:
            match = re.search(r'```json\s*(.*?)\s*```', st.session_state.screening_results, re.DOTALL)
            data = json.loads(match.group(1) if match else st.session_state.screening_results)
            st.session_state.parsed_data = data
            
            summary = data.get('screening_summary', {})
            candidates = data.get('candidates', [])
            insights = data.get('batch_insights', {})
            
            # Summary metrics
            st.markdown("### 📊 Results Summary")
            cols = st.columns(5)
            metrics = [
                ("Total", summary.get('total_candidates', '?'), "#fff"),
                ("Recommended", summary.get('recommended_for_interview', '?'), "#10b981"),
                ("Maybe", summary.get('maybe', '?'), "#f59e0b"),
                ("Time Saved", f"~{summary.get('time_saved_minutes', '?')}m", "#6366f1"),
                ("Bias Score", f"{summary.get('bias_mitigation_score', '?')}/100", "#8b5cf6")
            ]
            for col, (label, val, color) in zip(cols, metrics):
                with col:
                    st.markdown(f'<div class="metric-card"><p style="color:#9ca3af;margin:0;font-size:11px;">{label}</p><p style="color:{color};font-size:24px;font-weight:bold;margin:0;">{val}</p></div>', unsafe_allow_html=True)
            
            # Filters
            st.markdown("### 🔍 Filter & Sort")
            f1, f2, f3, f4 = st.columns(4)
            with f1: min_score = st.slider("Min Score", 0, 100, 0)
            with f2: salary_filter = st.selectbox("Salary", ["All", "Within Range", "Above", "Below"])
            with f3: sort_by = st.selectbox("Sort", ["Rank", "Score", "Experience"])
            with f4: compare_mode = st.checkbox("Compare Mode")
            
            # Filter
            filtered = [c for c in candidates if c.get('match_score', 0) >= min_score]
            if salary_filter != "All":
                filtered = [c for c in filtered if salary_filter.lower() in c.get('salary_alignment', '').lower()]
            
            if sort_by == "Score": filtered = sorted(filtered, key=lambda x: x.get('match_score', 0), reverse=True)
            elif sort_by == "Experience": filtered = sorted(filtered, key=lambda x: x.get('experience_years', 0), reverse=True)
            
            # Compare
            if compare_mode and len(st.session_state.selected_candidates) >= 2:
                if st.button(f"📊 Compare {len(st.session_state.selected_candidates)} Selected"):
                    st.session_state.working_on = "Comparing..."
                    selected_data = [c for c in candidates if c.get('identifier') in st.session_state.selected_candidates]
                    result, _ = compare_candidates(selected_data)
                    st.session_state.working_on = None
                    st.session_state.comparison_result = result
            
            # Candidates
            st.markdown(f"### 👥 Candidates ({len(filtered)})")
            
            for c in filtered:
                identifier = c.get('identifier', 'Unknown')
                
                if compare_mode:
                    is_sel = identifier in st.session_state.selected_candidates
                    if st.checkbox(f"Select {identifier}", value=is_sel, key=f"sel_{identifier}"):
                        if identifier not in st.session_state.selected_candidates:
                            st.session_state.selected_candidates.append(identifier)
                    else:
                        if identifier in st.session_state.selected_candidates:
                            st.session_state.selected_candidates.remove(identifier)
                
                # Rich display
                display_candidate_rich(c)
                
                # Action buttons
                score = c.get('match_score', 0)
                col_st, col_act = st.columns([1, 3])
                with col_st:
                    current_status = st.session_state.candidate_statuses.get(identifier, "Reviewing")
                    new_status = st.selectbox("Status", ["Reviewing", "Contacted", "Phone Screen", "Assessment", "Rejected", "Hired"], 
                        index=["Reviewing", "Contacted", "Phone Screen", "Assessment", "Rejected", "Hired"].index(current_status), key=f"st_{identifier}")
                    if new_status != current_status:
                        st.session_state.candidate_statuses[identifier] = new_status
                
                with col_act:
                    bc1, bc2, bc3, bc4 = st.columns(4)
                    if score >= 70:
                        with bc1:
                            if st.button("📧 Phone Screen", key=f"ps_{identifier}"):
                                st.session_state.working_on = "Drafting..."
                                email, _ = generate_email_template(c, "phone_screen", jd_text[:100])
                                st.session_state.working_on = None
                                st.session_state[f"email_{identifier}"] = email
                    if 50 <= score < 70:
                        with bc2:
                            if st.button("📝 Skills Test", key=f"sk_{identifier}"):
                                st.session_state.working_on = "Drafting..."
                                email, _ = generate_email_template(c, "skills_test", jd_text[:100])
                                st.session_state.working_on = None
                                st.session_state[f"email_{identifier}"] = email
                    if score < 50:
                        with bc3:
                            if st.button("❌ Rejection", key=f"rej_{identifier}"):
                                st.session_state.working_on = "Drafting..."
                                email, _ = generate_email_template(c, "rejection", jd_text[:100])
                                st.session_state.working_on = None
                                st.session_state[f"email_{identifier}"] = email
                    with bc4:
                        if st.button("💼 LinkedIn", key=f"li_{identifier}"):
                            st.session_state.working_on = "Drafting..."
                            msg, _ = generate_email_template(c, "linkedin", jd_text[:100])
                            st.session_state.working_on = None
                            st.session_state[f"email_{identifier}"] = msg
                
                if st.session_state.get(f"email_{identifier}"):
                    st.text_area("Generated:", st.session_state[f"email_{identifier}"], height=120, key=f"em_{identifier}")
                    if st.button("Clear", key=f"cl_{identifier}"):
                        del st.session_state[f"email_{identifier}"]
                        st.rerun()
                
                st.markdown("---")
            
            # Batch Actions
            st.markdown("### 📋 Batch Actions")
            if st.button("📊 Generate Hiring Manager Summary", use_container_width=True):
                st.session_state.working_on = "Generating summary..."
                result, _ = generate_hiring_manager_summary(candidates, bias_free)
                st.session_state.working_on = None
                st.session_state.hm_summary = result
                st.rerun()
            
            # HM Summary with rich display
            if st.session_state.get('hm_summary'):
                st.markdown("### 👔 Hiring Manager Summary")
                display_hm_summary_rich(st.session_state.hm_summary)
            
            # Export
            st.markdown("### 📥 Export Reports")
            e1, e2, e3, e4 = st.columns(4)
            
            with e1:
                pdf_bytes = generate_pdf_bytes(data, jd_text[:50] if jd_text else "Report")
                if pdf_bytes:
                    st.download_button("📄 PDF Report", pdf_bytes, "screening_report.pdf", "application/pdf", use_container_width=True)
                else:
                    st.warning("PDF requires reportlab")
            with e2:
                st.download_button("📊 CSV", generate_csv_report(data), "report.csv", "text/csv", use_container_width=True)
            with e3:
                st.download_button("🔗 JSON", json.dumps(data, indent=2), "data.json", "application/json", use_container_width=True)
            with e4:
                if st.button("🔄 New Screening", use_container_width=True):
                    st.session_state.screening_results = None
                    st.session_state.selected_candidates = []
                    st.session_state.comparison_result = None
                    st.session_state.hm_summary = None
                    st.rerun()
            
            # Insights
            if insights:
                st.markdown("### 💡 Batch Insights")
                st.info(f"**🏆 Top Candidate:** {insights.get('strongest_candidate_summary', '')}")
                st.warning(f"**📋 Common Gaps:** {', '.join(insights.get('common_gaps', []))}")
                st.success(f"**✅ Recommendation:** {insights.get('hiring_recommendation', '')}")
        
        except Exception as e:
            st.error(f"Error parsing results: {e}")
            st.markdown(f'<div class="output-box">{st.session_state.screening_results}</div>', unsafe_allow_html=True)

with tab_blind:
    st.markdown("### 👤 Blind Resume Anonymization")
    c1, c2 = st.columns([2, 1])
    with c1:
        src = st.radio("Source:", ["📝 Paste", "📄 Upload"], horizontal=True, key="anon_src")
        if src == "📄 Upload":
            f = st.file_uploader("Upload CV", type=['pdf', 'docx', 'txt'], key="anon_file")
            resume = extract_text_from_file(f) if f else ""
        else:
            resume = st.text_area("Paste resume:", height=250)
    with c2:
        time_based = st.checkbox("⏱️ Time-based dates")
        st.caption("Removes: names, emails, phones, addresses, universities, companies")
    
    if st.button("👤 Anonymize", type="primary", use_container_width=True):
        if resume:
            st.session_state.working_on = "Anonymizing..."
            result, _ = anonymize_cv(resume, {'time_based': time_based})
            st.session_state.working_on = None
            if not result.startswith("Error"):
                st.session_state.anonymized_result = result
                st.rerun()
            else:
                st.error(result)
    
    if st.session_state.get('anonymized_result'):
        try:
            match = re.search(r'```json\s*(.*?)\s*```', st.session_state.anonymized_result, re.DOTALL)
            data = json.loads(match.group(1) if match else st.session_state.anonymized_result)
            st.metric("Confidence", f"{data.get('anonymization_score', 0)}%")
            st.markdown(f'<div class="output-box">{data.get("anonymized_cv", "")}</div>', unsafe_allow_html=True)
            st.download_button("📥 Download", data.get("anonymized_cv", ""), "anonymized.txt", use_container_width=True)
        except:
            st.markdown(f'<div class="output-box">{st.session_state.anonymized_result}</div>', unsafe_allow_html=True)

with tab_salary:
    st.markdown("### 💰 Salary Estimator")
    c1, c2 = st.columns(2)
    with c1: sal_jd = st.text_area("Job Description:", height=150, key="sal_jd")
    with c2: sal_cv = st.text_area("Candidate CV:", height=150, key="sal_cv")
    location = st.selectbox("Market:", ["US - National", "US - SF Bay Area", "US - NYC", "US - Remote", "UK - London", "EU", "Canada"])
    
    if st.button("💰 Estimate Salary", type="primary", use_container_width=True):
        if sal_jd and sal_cv:
            st.session_state.working_on = "Estimating..."
            result, _ = estimate_salary(sal_cv, sal_jd, location)
            st.session_state.working_on = None
            if not result.startswith("Error"):
                st.markdown(f'<div class="output-box">{result}</div>', unsafe_allow_html=True)
            else:
                st.error(result)

# ============================================
# FLOATING FEEDBACK (Clean popover)
# ============================================
st.markdown('<div style="height:60px;"></div>', unsafe_allow_html=True)

st.markdown("""
<style>
div[data-testid="stPopover"] button {
    background: linear-gradient(135deg, #6366f1, #8b5cf6) !important;
    color: white !important;
    border: none !important;
    border-radius: 25px !important;
    padding: 10px 20px !important;
    font-weight: 600 !important;
    box-shadow: 0 4px 15px rgba(99, 102, 241, 0.3);
}
</style>
""", unsafe_allow_html=True)

_, _, _, fb_col = st.columns([4, 1, 1, 1])
with fb_col:
    with st.popover("💬 Feedback"):
        st.markdown("**Send Feedback**")
        fb_type = st.segmented_control("Type", ["🐛 Bug", "✨ Feature", "💬 General"], 
                                        default="💬 General", label_visibility="collapsed")
        fb_msg = st.text_area("Message", height=100, placeholder="What's on your mind?",
                              label_visibility="collapsed", key="fb_msg")
        if st.button("Send Feedback", type="primary", use_container_width=True, key="fb_send"):
            if fb_msg:
                t = fb_type.split()[1].lower() if fb_type else "general"
                success = submit_feedback("screen", t, fb_msg)
                if success:
                    st.success("Thanks! 🙏")
                else:
                    st.error("Failed to send")

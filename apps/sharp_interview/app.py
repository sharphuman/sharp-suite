"""Sharp Interview - AI Interview Evaluation with Multi-Candidate Comparison"""
import streamlit as st
import os
import requests
from datetime import datetime, timedelta
import secrets
import json
import re
import io
import tempfile
import zipfile
from xml.etree import ElementTree

# ============================================
# SHARED MODULE IMPORTS
# ============================================

# Try to import shared logging first
try:
    import shared_logging as log
    log.set_app_name("interview")
    HAS_LOGGING = True
except ImportError:
    HAS_LOGGING = False
    # Minimal fallback logging
    class FallbackLog:
        @staticmethod
        def debug(msg, **kwargs): print(f"[DEBUG] {msg}")
        @staticmethod
        def info(msg, **kwargs): print(f"[INFO] {msg}")
        @staticmethod
        def warn(msg, **kwargs): print(f"[WARN] {msg}")
        @staticmethod
        def error(msg, **kwargs): print(f"[ERROR] {msg}")
        @staticmethod
        def log_import_status(success, module, error_msg=None):
            if success:
                print(f"[INFO] Imported {module}")
            else:
                print(f"[WARN] Failed to import {module}: {error_msg}")
    log = FallbackLog()

# Try to import few-shot examples
try:
    from interview_examples import get_few_shot_examples
    HAS_EXAMPLES = True
    log.log_import_status(True, "interview_examples")
except ImportError:
    HAS_EXAMPLES = False
    log.log_import_status(False, "interview_examples", "File not found")
    def get_few_shot_examples():
        return ""  # Empty fallback

# Now import shared config and UI
try:
    from shared_config import (
        SUPABASE_URL, SUPABASE_ANON_KEY, SUPABASE_SERVICE_KEY,
        ANTHROPIC_API_KEY, GOD_PASSWORD, APP_URLS, CLAUDE_MODEL
    )
    from shared_ui import (
        apply_global_styles,
        render_top_banner,
        render_header,
        render_sidebar,
        render_feedback_widget,
        
        COLORS
    )
    USING_SHARED = True
    log.log_import_status(True, "shared_ui + shared_config")
except ImportError as e:
    USING_SHARED = False
    log.log_import_status(False, "shared_ui + shared_config", str(e))
    log.warn(f"Current working directory: {os.getcwd()}")
    log.warn(f"Files in cwd: {os.listdir('.')[:20]}")
    
    # Fallback values
    SUPABASE_URL = "https://qkjtprqgblnfftrotyks.supabase.co"
    SUPABASE_ANON_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InFranRwcnFnYmxuZmZ0cm90eWtzIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NjUzNTgzNDAsImV4cCI6MjA4MDkzNDM0MH0.pVzSq4M5i58zBGl7OPDhNL9qYBcg-bz8MVrBI5MQSkw"
    SUPABASE_SERVICE_KEY = os.environ.get("SUPABASE_SERVICE_KEY", SUPABASE_ANON_KEY)
    GOD_PASSWORD = os.environ.get("GOD_PASSWORD", "G0DHum@n101!!!")
    ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
    CLAUDE_MODEL = "claude-sonnet-4-20250514"
    APP_URLS = {"portal": "https://demo.sharphuman.com", "jd": "https://jd.sharphuman.com", "screen": "https://screen.sharphuman.com", "interview": "https://interview.sharphuman.com", "outreach": "https://outreach.sharphuman.com", "content": "https://content.sharphuman.com", "sales": "https://sales.sharphuman.com", "admin": "https://admin.sharphuman.com"}
    COLORS = {"primary": "#6366f1", "secondary": "#8b5cf6", "success": "#10b981", "warning": "#eab308", "error": "#ef4444"}

FOCUS_AREAS = ["Technical Skills", "Problem Solving", "Communication", "Leadership", "Cultural Fit", "Experience Depth", "Motivation", "Collaboration"]
INTERVIEW_STAGES = ["Phone Screen", "Technical Round", "Hiring Manager", "Final Round", "Culture Fit", "Panel Interview"]

# ============================================
# EVALUATOR PERSONAS (Option 1)
# ============================================
EVALUATOR_PERSONAS = {
    "balanced": {
        "name": "⚖️ Balanced Reviewer",
        "description": "Standard balanced analysis weighing pros and cons equally",
        "system_prompt": "You are a fair and balanced interview evaluator. Weigh both strengths and weaknesses equally. Provide objective assessment without bias toward optimism or skepticism."
    },
    "skeptical_recruiter": {
        "name": "🔍 Skeptical Recruiter",
        "description": "Critical, detail-oriented, flags every inconsistency",
        "system_prompt": "You are a veteran Technical Recruiter with 20 years of experience. You are cynical and hard to impress. Your job is to protect the company from bad hires. Scrutinize every claim. Flag gaps in employment, vague timelines, buzzword stuffing, and any exaggeration. If something seems too good to be true, call it out. Be direct and unforgiving in your assessment."
    },
    "hiring_manager": {
        "name": "🎯 Hiring Manager",
        "description": "Practical, urgent - can they hit the ground running?",
        "system_prompt": "You are a busy Hiring Manager who needs someone to deliver results immediately. You care about: Can they do THIS job on Day 1? Look for specific project outcomes, relevant tool/stack matches, and evidence of independent delivery. Estimate ramp-up time. You have no patience for theory without practice."
    },
    "technical_lead": {
        "name": "🛠️ Technical Lead",
        "description": "Deep technical scrutiny, BS detector activated",
        "system_prompt": "You are a Senior Technical Lead with a finely-tuned BS detector. You've interviewed hundreds of candidates who exaggerate their technical skills. Probe for DEPTH not breadth. When someone claims 'microservices architecture' - did they actually design it or just maintain it? Check for specific technical decisions, trade-offs discussed, and complexity of problems solved. Surface-level buzzword knowledge is a red flag."
    },
    "executive": {
        "name": "👔 Executive/CEO",
        "description": "Strategic fit, leadership potential, long-term value",
        "system_prompt": "You are a CEO evaluating candidates for strategic fit and long-term potential. You care less about specific syntax and more about: problem-solving approach, business acumen, leadership qualities, cultural alignment, and growth trajectory. Can this person eventually lead a team or department? Do they think like an owner?"
    }
}

# ============================================
# ANALYSIS RIGOR LEVELS (Option 2)
# ============================================
RIGOR_LEVELS = {
    "optimistic": {
        "name": "🌱 Optimistic",
        "description": "Focus on potential and strengths (good for junior roles)",
        "prompt_modifier": "Focus on the candidate's POTENTIAL and transferable skills. Give benefit of the doubt on gaps or missing experience. This is likely a junior role where growth matters more than current expertise. Highlight what they COULD become with proper mentorship."
    },
    "balanced": {
        "name": "⚖️ Balanced",
        "description": "Standard SWOT analysis",
        "prompt_modifier": "Provide a balanced assessment covering both strengths and weaknesses equally. Be fair but thorough."
    },
    "ruthless": {
        "name": "🔥 Ruthless",
        "description": "Assume CV is exaggerated, find reasons NOT to hire (executive roles)",
        "prompt_modifier": "Be EXTREMELY critical. Assume the CV contains exaggerations until proven otherwise in the transcript. Actively look for reasons NOT to hire. This is a high-stakes role where a bad hire is catastrophic. Every vague answer is a red flag. Every unverified claim should lower the score. The bar is VERY high."
    }
}

# ============================================
# QUESTION GENERATION STYLES (Tab 2 Enhancement)
# ============================================
QUESTION_STYLES = {
    "standard": {
        "name": "🎯 Standard Interviewer",
        "description": "Balanced mix of behavioral and technical questions",
        "prompt_modifier": "Generate a balanced mix of behavioral (STAR format) and technical questions. Focus on past experience and problem-solving."
    },
    "technical_deep": {
        "name": "🔬 Technical Deep-Dive",
        "description": "Heavy focus on architecture, coding, system design",
        "prompt_modifier": "Focus heavily on technical depth. Include system design questions, coding scenarios, and architecture discussions. Probe for specific technical decisions and trade-offs. Less emphasis on soft skills."
    },
    "culture_fit": {
        "name": "🤝 Culture & Soft Skills",
        "description": "Team dynamics, communication, values alignment",
        "prompt_modifier": "Focus on soft skills, team dynamics, and cultural fit. Include questions about collaboration, conflict resolution, giving/receiving feedback, and work style. Less emphasis on technical skills."
    },
    "rapid_screen": {
        "name": "🏃 Rapid Screen",
        "description": "Quick qualification questions for phone screens",
        "prompt_modifier": "Generate concise, efficient questions for a quick phone screen. Focus on must-have qualifications and deal-breakers. Questions should be answerable in 1-2 minutes each."
    },
    "stress_test": {
        "name": "🎭 Stress Test",
        "description": "Challenging scenarios to test pressure handling",
        "prompt_modifier": "Include challenging scenarios that test how candidates handle pressure, ambiguity, and difficult situations. Add curveball questions and ethical dilemmas. Test for composure and critical thinking under stress."
    }
}

DIFFICULTY_LEVELS = {
    "junior": {
        "name": "🌱 Junior-Friendly",
        "description": "Entry-level, focus on potential",
        "prompt_modifier": "Calibrate for junior/entry-level candidates. Focus on learning ability, enthusiasm, and foundational knowledge. Avoid questions requiring extensive experience. Accept theoretical knowledge over practical examples."
    },
    "mid": {
        "name": "⚖️ Mid-Level",
        "description": "2-5 years experience expected",
        "prompt_modifier": "Calibrate for mid-level candidates with 2-5 years experience. Expect concrete examples from past work. Probe for independent problem-solving and growing ownership."
    },
    "senior": {
        "name": "🔥 Senior/Lead",
        "description": "5+ years, leadership expected",
        "prompt_modifier": "Calibrate for senior candidates. Expect deep expertise, mentorship examples, and strategic thinking. Probe for technical leadership, architecture decisions, and cross-team influence."
    },
    "staff": {
        "name": "💀 Staff/Principal",
        "description": "Expert level, org-wide impact",
        "prompt_modifier": "Calibrate for staff/principal level. Expect org-wide impact, technical vision, and executive communication. Probe for building systems at scale, influencing without authority, and long-term technical strategy."
    }
}

# ============================================
# AUTH FUNCTIONS
# ============================================

def create_session(user_id, email):
    token = secrets.token_urlsafe(32)
    try:
        requests.post(f"{SUPABASE_URL}/rest/v1/sessions", 
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}", "Content-Type": "application/json", "Prefer": "return=minimal"}, 
            json={"user_id": user_id, "token": token, "ip_address": "unknown", "device_hash": "interview", "is_active": True, "expires_at": (datetime.utcnow() + timedelta(hours=24)).isoformat()}, timeout=10)
    except:
        pass
    return token

def supabase_sign_in(email, password):
    try:
        r = requests.post(f"{SUPABASE_URL}/auth/v1/token?grant_type=password", 
            headers={"apikey": SUPABASE_ANON_KEY, "Content-Type": "application/json"}, 
            json={"email": email, "password": password}, timeout=10)
        data = r.json()
        if r.status_code == 200 and data.get("access_token"):
            user = data.get("user", {})
            return {"success": True, "user": user, "session_token": create_session(user.get("id"), email)}
        return {"success": False, "message": data.get("error_description") or "Invalid"}
    except Exception as e:
        return {"success": False, "message": str(e)}

def supabase_sign_up(email, password):
    try:
        r = requests.post(f"{SUPABASE_URL}/auth/v1/signup", 
            headers={"apikey": SUPABASE_ANON_KEY, "Content-Type": "application/json"}, 
            json={"email": email, "password": password}, timeout=10)
        data = r.json()
        if r.status_code == 200 and data.get("user"):
            return {"success": True, "message": "Check email!"}
        return {"success": False, "message": data.get("error_description") or "Failed"}
    except Exception as e:
        return {"success": False, "message": str(e)}

def validate_session_token(token):
    if not token:
        return None
    try:
        r = requests.get(f"{SUPABASE_URL}/rest/v1/sessions?token=eq.{token}&is_active=eq.true", 
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}"}, timeout=10)
        if r.status_code == 200 and r.json():
            session = r.json()[0]
            expires = datetime.fromisoformat(session["expires_at"].replace("Z", "+00:00"))
            if expires.replace(tzinfo=None) > datetime.utcnow():
                ur = requests.get(f"{SUPABASE_URL}/rest/v1/user_profiles?id=eq.{session['user_id']}", 
                    headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}"}, timeout=10)
                if ur.status_code == 200 and ur.json():
                    p = ur.json()[0]
                    return {"user_id": session["user_id"], "email": p.get("email"), "plan": p.get("plan", "free"), "token": token}
    except:
        pass
    return None

def log_usage(user_id, session_id, app, action, tokens_used=0):
    try:
        requests.post(f"{SUPABASE_URL}/rest/v1/usage_logs", 
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}", "Content-Type": "application/json", "Prefer": "return=minimal"}, 
            json={"user_id": user_id, "session_id": session_id, "app": app, "action": action, "tokens_used": tokens_used}, timeout=5)
    except:
        pass

def submit_feedback(app, feedback_type, message):
    try:
        email = get_user_email()
        user_id = st.session_state.user.get("id") if st.session_state.user else None
        r = requests.post(f"{SUPABASE_URL}/rest/v1/user_feedback", 
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}", "Content-Type": "application/json", "Prefer": "return=minimal"}, 
            json={"user_id": user_id, "app": app, "feedback_type": feedback_type, "rating": 4, "message": message, "email": email}, timeout=10)
        return r.status_code in [200, 201]
    except:
        return False

def submit_analysis_feedback(result_data, feedback, reason=None, comment=None):
    """
    Submit feedback on an AI analysis result.
    Stores metadata only - no PII or transcript content.
    
    Args:
        result_data: The AI analysis result dict
        feedback: 'thumbs_up' or 'thumbs_down'
        reason: Optional reason code (e.g., 'too_generic', 'missed_red_flag')
        comment: Optional free-text comment
    """
    try:
        user_id = st.session_state.user.get("id") if st.session_state.user else None
        
        # Extract metadata from result (no PII)
        analysis_settings = result_data.get('analysis_settings', {})
        
        payload = {
            "user_id": user_id,
            "app": "interview",
            "persona": analysis_settings.get('persona', 'balanced'),
            "rigor_level": analysis_settings.get('rigor', 'balanced'),
            "overall_score": result_data.get('overall_score'),
            "recommendation": result_data.get('recommendation'),
            "feedback": feedback,
            "feedback_reason": reason,
            "feedback_comment": comment
        }
        
        r = requests.post(
            f"{SUPABASE_URL}/rest/v1/analysis_feedback", 
            headers={
                "apikey": SUPABASE_ANON_KEY, 
                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}", 
                "Content-Type": "application/json", 
                "Prefer": "return=minimal"
            }, 
            json=payload, 
            timeout=10
        )
        return r.status_code in [200, 201]
    except Exception as e:
        if HAS_LOGGING:
            log.error(f"Failed to submit analysis feedback: {e}")
        return False

# Feedback reason options
FEEDBACK_REASONS = {
    "thumbs_down": [
        ("too_generic", "Too generic / vague"),
        ("missed_red_flag", "Missed obvious red flag"),
        ("wrong_score", "Score feels wrong"),
        ("missed_strength", "Missed key strength"),
        ("hallucination", "Made up / hallucinated info"),
        ("poor_evidence", "Weak evidence cited"),
        ("other", "Other")
    ],
    "thumbs_up": [
        ("accurate", "Accurate assessment"),
        ("good_evidence", "Great evidence/quotes"),
        ("caught_red_flag", "Caught something I missed"),
        ("helpful_insights", "Helpful insights"),
        ("other", "Other")
    ]
}

def get_jd_history(limit=20):
    user_id = st.session_state.user.get("id") if st.session_state.user else None
    if not user_id or user_id == "god":
        return []
    try:
        r = requests.get(f"{SUPABASE_URL}/rest/v1/jd_history?user_id=eq.{user_id}&order=created_at.desc&limit={limit}", 
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}"}, timeout=10)
        return r.json() if r.status_code == 200 else []
    except:
        return []

def init_session():
    defaults = [
        ('authenticated', False), ('user', None), ('is_god', False), ('session_token', None),
        ('user_plan', 'free'), ('working_on', None), ('current_step', 'input'),
        ('candidates', []), ('current_candidate', None), ('comparison_result', None),
        ('questions_result', None), ('jd_text', ''), ('cv_text', ''),
        ('feedback_given', {}),  # Track which candidates have received feedback
        ('locked_settings', None),  # Lock persona/rigor/focus after first candidate
        ('coach_result', None)  # Recruiter coach result
    ]
    for k, v in defaults:
        if k not in st.session_state:
            st.session_state[k] = v

def check_url_auth():
    token = st.query_params.get("auth")
    if token and not st.session_state.authenticated:
        user_info = validate_session_token(token)
        if user_info:
            st.session_state.authenticated = True
            st.session_state.user = {"email": user_info["email"], "id": user_info["user_id"]}
            st.session_state.session_token = token
            st.session_state.user_plan = user_info.get("plan", "free")
            st.session_state.is_god = user_info.get("plan") == "god"
            return True
    return False

def get_user_email():
    if st.session_state.user:
        return st.session_state.user.get("email", "User")
    return "GOD" if st.session_state.is_god else "User"

def build_app_url(app_name):
    base = APP_URLS.get(app_name, "")
    token = st.session_state.get("session_token", "")
    return f"{base}?auth={token}" if base and token else base

# ============================================
# FILE EXTRACTION (ROBUST)
# ============================================

def is_readable_text(text):
    """Check if extracted text is actually readable (not garbled)"""
    if not text or len(text.strip()) < 10:
        return False
    # Count alphabetic characters - real text should have words
    alpha_chars = sum(1 for c in text if c.isalpha())
    alpha_ratio = alpha_chars / len(text) if len(text) > 0 else 0
    # Real text should have at least 30% letters
    return alpha_ratio > 0.3

def clean_extracted_text(text):
    """Clean up extracted text by removing problematic characters"""
    if not text:
        return ""
    # Remove null bytes and other control characters (except newlines/tabs)
    cleaned = ''.join(c if c.isprintable() or c in '\n\r\t' else ' ' for c in text)
    # Normalize whitespace
    cleaned = re.sub(r'[ \t]+', ' ', cleaned)
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
    return cleaned.strip()

def extract_text_from_file(uploaded_file):
    if uploaded_file is None:
        return ""
    
    file_type = uploaded_file.name.split('.')[-1].lower()
    content = uploaded_file.read()
    uploaded_file.seek(0)
    
    if file_type == 'txt':
        text = content.decode('utf-8', errors='ignore')
        return clean_extracted_text(text)
    
    elif file_type == 'pdf':
        extracted_text = ""
        
        # Try PyMuPDF first with multiple extraction methods
        try:
            import fitz
            pdf = fitz.open(stream=content, filetype="pdf")
            text_parts = []
            for page in pdf:
                # Try standard text extraction
                page_text = page.get_text("text")
                if not page_text or not page_text.strip():
                    # Try with different flags
                    page_text = page.get_text("text", flags=fitz.TEXT_PRESERVE_WHITESPACE)
                if not page_text or not page_text.strip():
                    # Try blocks extraction
                    blocks = page.get_text("blocks")
                    page_text = "\n".join([b[4] for b in blocks if b[6] == 0])  # type 0 = text
                if page_text and page_text.strip():
                    text_parts.append(page_text)
            pdf.close()
            extracted_text = "\n".join(text_parts)
        except Exception as e:
            pass
        
        # If we got text, clean and validate it
        if extracted_text and extracted_text.strip():
            cleaned = clean_extracted_text(extracted_text)
            if is_readable_text(cleaned):
                return cleaned
        
        # Fallback: try pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                extracted_text = "\n".join(text_parts)
                if extracted_text and extracted_text.strip():
                    cleaned = clean_extracted_text(extracted_text)
                    if is_readable_text(cleaned):
                        return cleaned
        except:
            pass
        
        # Fallback 3: Try PyPDF2 if available
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(content))
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            extracted_text = "\n".join(text_parts)
            if extracted_text and extracted_text.strip():
                cleaned = clean_extracted_text(extracted_text)
                if is_readable_text(cleaned):
                    return cleaned
        except:
            pass
        
        # Last resort: if we got ANY text from PyMuPDF, return it with a warning
        try:
            import fitz
            pdf = fitz.open(stream=content, filetype="pdf")
            text_parts = []
            for page in pdf:
                page_text = page.get_text()
                if page_text:
                    text_parts.append(page_text)
            pdf.close()
            if text_parts:
                result = clean_extracted_text("\n".join(text_parts))
                if len(result) > 50:  # If we got something substantial
                    return result
        except:
            pass
        
        return "[PDF extraction failed - the PDF may be scanned/image-based. Please paste the content directly.]"
    
    elif file_type in ['docx', 'doc']:
        # Try python-docx
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            paragraphs = []
            for p in doc.paragraphs:
                if p.text.strip():
                    paragraphs.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            if paragraphs:
                text = '\n\n'.join(paragraphs)
                cleaned = clean_extracted_text(text)
                if is_readable_text(cleaned):
                    return cleaned
        except Exception as e:
            pass
        
        # Fallback: raw XML from ZIP
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as z:
                if 'word/document.xml' in z.namelist():
                    xml_content = z.read('word/document.xml')
                    tree = ElementTree.fromstring(xml_content)
                    texts = []
                    for elem in tree.iter():
                        if elem.text and elem.text.strip():
                            texts.append(elem.text.strip())
                    if texts:
                        text = ' '.join(texts)
                        cleaned = clean_extracted_text(text)
                        if is_readable_text(cleaned):
                            return cleaned
        except:
            pass
        
        return "[DOCX extraction failed - please paste the content directly]"
    
    elif file_type in ['vtt', 'srt']:
        text = content.decode('utf-8', errors='ignore')
        text = re.sub(r'\d{2}:\d{2}:\d{2}[.,]\d{3}\s*-->\s*\d{2}:\d{2}:\d{2}[.,]\d{3}', '', text)
        text = re.sub(r'WEBVTT.*?\n', '', text)
        text = re.sub(r'<[^>]+>', '', text)
        text = re.sub(r'^\d+$', '', text, flags=re.MULTILINE)
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        return '\n'.join(lines)
    
    elif file_type in ['mp3', 'm4a', 'wav', 'mp4', 'webm', 'ogg']:
        openai_key = os.environ.get("OPENAI_API_KEY", "")
        if openai_key:
            try:
                with tempfile.NamedTemporaryFile(suffix=f".{file_type}", delete=False) as tmp:
                    tmp.write(content)
                    tmp_path = tmp.name
                with open(tmp_path, 'rb') as f:
                    r = requests.post("https://api.openai.com/v1/audio/transcriptions",
                        headers={"Authorization": f"Bearer {openai_key}"},
                        files={"file": f}, data={"model": "whisper-1"}, timeout=300)
                os.unlink(tmp_path)
                if r.status_code == 200:
                    return r.json().get("text", "")
            except:
                pass
        return "[Audio transcription requires OPENAI_API_KEY]"
    
    return content.decode('utf-8', errors='ignore')

# ============================================
# CLAUDE API
# ============================================

def call_claude(prompt, max_tokens=8000, action="interview"):
    if not ANTHROPIC_API_KEY:
        return "Error: ANTHROPIC_API_KEY not set", 0
    try:
        model = CLAUDE_MODEL if USING_SHARED else "claude-sonnet-4-20250514"
        r = requests.post("https://api.anthropic.com/v1/messages",
            headers={"x-api-key": ANTHROPIC_API_KEY, "Content-Type": "application/json", "anthropic-version": "2023-06-01"},
            json={"model": model, "max_tokens": max_tokens, "messages": [{"role": "user", "content": prompt}]}, timeout=180)
        if r.status_code == 200:
            text = r.json()["content"][0]["text"]
            if st.session_state.user:
                log_usage(st.session_state.user.get("id"), st.session_state.get("session_token"), "interview", action, (len(prompt)+len(text))//4)
            return text, (len(prompt)+len(text))//4
        elif r.status_code == 500:
            return "Error: 500 - Claude API server error. Please try again in a moment.", 0
        elif r.status_code == 529:
            return "Error: 529 - Claude API overloaded. Please wait a moment and try again.", 0
        elif r.status_code == 429:
            return "Error: 429 - Rate limited. Please wait a moment and try again.", 0
        else:
            try:
                err_detail = r.json().get("error", {}).get("message", r.text[:200])
            except:
                err_detail = r.text[:200]
            return f"Error: {r.status_code} - {err_detail}", 0
    except requests.exceptions.Timeout:
        return "Error: Request timed out. The transcript may be too long - try shortening it.", 0
    except Exception as e:
        return f"Error: {e}", 0

# ============================================
# QUESTION GENERATOR
# ============================================

def generate_questions(job_title, requirements, stage, duration, focus_areas, cv_text="", exclude_illegal=True, style="standard", difficulty="mid", company_context=""):
    """
    Generate interview questions with expected answers, CV verification questions, and JD fit assessment.
    
    Args:
        style: Key from QUESTION_STYLES
        difficulty: Key from DIFFICULTY_LEVELS
        company_context: Optional company/team context for tailored questions
    """
    
    # Get style and difficulty settings
    style_config = QUESTION_STYLES.get(style, QUESTION_STYLES["standard"])
    difficulty_config = DIFFICULTY_LEVELS.get(difficulty, DIFFICULTY_LEVELS["mid"])
    
    company_section = ""
    if company_context and company_context.strip():
        company_section = f"""
## COMPANY/TEAM CONTEXT
Tailor questions to probe for fit with this environment:
{company_context}
"""

    cv_section = ""
    cv_verification_instruction = ""
    if cv_text and cv_text.strip():
        cv_section = f"""
## CANDIDATE CV
{cv_text[:4000]}
"""
        cv_verification_instruction = """
## CV VERIFICATION QUESTIONS (CRITICAL)
Generate 2-3 specific questions to VERIFY claims made in the CV:
- If they claim "Led team of 10", ask for specific names, challenges, outcomes
- If they claim a technology, ask for specific implementation details
- If there are gaps or vague descriptions, probe for specifics
- These questions should reveal if the candidate truly did what they claim
"""

    prompt = f"""You are an expert interviewer creating a comprehensive interview guide. Generate questions WITH expected answers.

## ROLE: {job_title}

## JOB REQUIREMENTS (Use these to assess FIT):
{requirements[:4000]}

## INTERVIEW STAGE: {stage}
## DURATION: {duration} minutes
## FOCUS AREAS: {', '.join(focus_areas)}

## QUESTION STYLE: {style_config['name']}
{style_config['prompt_modifier']}

## DIFFICULTY LEVEL: {difficulty_config['name']}
{difficulty_config['prompt_modifier']}
{company_section}
{cv_section}
{cv_verification_instruction}

{"IMPORTANT: Exclude any questions that could be considered discriminatory (age, religion, family status, etc.)" if exclude_illegal else ""}

## QUESTION TYPES TO INCLUDE:

1. **JD FIT QUESTIONS**: Questions that directly assess if the candidate meets the job requirements
2. **CV VERIFICATION QUESTIONS**: Questions to verify specific claims from their CV (if provided)
3. **BEHAVIORAL QUESTIONS**: Past behavior predicts future behavior
4. **TECHNICAL/SKILL QUESTIONS**: Assess core competencies for the role

Generate a structured interview guide:

```json
{{
    "settings": {{
        "style": "{style_config['name']}",
        "difficulty": "{difficulty_config['name']}",
        "duration_minutes": {duration},
        "total_questions": <number>
    }},
    "opening": {{
        "icebreaker": "<warm-up question>",
        "icebreaker_purpose": "<why this opens the conversation well>",
        "role_intro": "<question about their understanding of the role>",
        "expected_good_answer": "<what a prepared candidate would say>"
    }},
    "jd_fit_questions": [
        {{
            "requirement_being_tested": "<specific requirement from JD>",
            "question": "<the interview question>",
            "expected_good_answer": "<detailed answer showing they meet the requirement>",
            "expected_poor_answer": "<answer that would indicate they don't meet it>",
            "follow_ups": ["<dig deeper>", "<probe for specifics>"],
            "scoring_guide": {{
                "excellent": "<what 9-10 looks like>",
                "acceptable": "<what 6-8 looks like>",
                "poor": "<what 1-5 looks like>"
            }},
            "time_estimate": "<minutes>"
        }}
    ],
    "cv_verification_questions": [
        {{
            "cv_claim_being_verified": "<specific claim from CV>",
            "question": "<question to verify this claim>",
            "expected_truthful_answer": "<what someone who really did this would say>",
            "red_flags": ["<signs they're exaggerating>", "<vague responses>"],
            "follow_up_if_suspicious": "<how to probe deeper>",
            "time_estimate": "<minutes>"
        }}
    ],
    "behavioral_questions": [
        {{
            "competency": "<what this assesses>",
            "question": "<STAR-format behavioral question>",
            "expected_good_answer": {{
                "situation": "<what context they should describe>",
                "task": "<what responsibility they should explain>",
                "action": "<specific actions they should detail>",
                "result": "<measurable outcomes they should share>"
            }},
            "red_flags": ["<warning signs>"],
            "time_estimate": "<minutes>"
        }}
    ],
    "closing": {{
        "candidate_questions": "<invite their questions>",
        "what_good_candidates_ask": ["<insightful question>", "<shows research>"],
        "next_steps": "<explain process>",
        "timeline": "<when they'll hear back>"
    }},
    "evaluation_scorecard": [
        {{
            "criterion": "<what to evaluate>",
            "weight": "<High/Medium/Low>",
            "must_have": <true/false>,
            "questions_that_assess_this": ["<question reference>"],
            "scoring_notes": "<how to score this>"
        }}
    ],
    "overall_recommendation_guide": {{
        "strong_hire_signals": ["<what indicates exceptional candidate>"],
        "hire_signals": ["<what indicates good candidate>"],
        "no_hire_signals": ["<what indicates poor fit>"],
        "red_flags_to_watch": ["<immediate disqualifiers>"]
    }}
}}
```

Generate {max(4, duration // 8)} total questions across all categories, appropriate for {duration} minutes.
Ensure every question has a detailed expected answer so interviewers know what "good" looks like."""

    return call_claude(prompt, max_tokens=6000, action="generate_questions")

def analyze_recruiter_screen(transcript, call_type="candidate_screen", candidate_role="", company_context="", key_concerns="", jd_text=""):
    """
    Analyze a recruiter's performance on a call (screening, intake, etc.)
    
    Args:
        transcript: The call transcript
        call_type: Type of call being analyzed
        candidate_role: Role being recruited for (if applicable)
        company_context: Context about the company/client
        key_concerns: Specific areas to focus on
        jd_text: Job description to analyze candidate fit against
    """
    
    # Call type specific frameworks
    RECRUITER_CALL_TYPES = {
        "candidate_screen": {
            "name": "Candidate Screening Call",
            "description": "Recruiter qualifying and selling opportunity to a candidate",
            "skills": {
                "opening": {
                    "name": "Opening & Rapport",
                    "weight": 10,
                    "criteria": ["Quick rapport building", "Set agenda/expectations", "Created comfortable environment"]
                },
                "qualification": {
                    "name": "Candidate Qualification",
                    "weight": 25,
                    "criteria": ["Salary expectations extracted", "Notice period/availability confirmed", "Motivation for move understood", "Competing offers identified", "Relocation/remote requirements clarified"]
                },
                "role_discovery": {
                    "name": "Experience & Fit Assessment",
                    "weight": 25,
                    "criteria": ["Relevant experience verified", "Skills match explored", "Cultural fit signals gathered", "Red flags identified (gaps, job hopping, evasiveness)", "Strengths/weaknesses understood"]
                },
                "role_sell": {
                    "name": "Opportunity Sell",
                    "weight": 20,
                    "criteria": ["Role benefits communicated", "Company value proposition delivered", "Candidate excitement built", "Objections to role addressed", "Differentiated from other opportunities"]
                },
                "control": {
                    "name": "Call Control & Efficiency",
                    "weight": 10,
                    "criteria": ["Kept call on track", "Managed talkative candidate", "Extracted info efficiently", "Professional pace maintained"]
                },
                "closing": {
                    "name": "Close & Next Steps",
                    "weight": 10,
                    "criteria": ["Clear next steps established", "Candidate commitment obtained", "Timeline confirmed", "Exclusivity/priority discussed"]
                }
            }
        },
        "client_intake": {
            "name": "Client Job Intake Call",
            "description": "Recruiter taking a job order from a client",
            "skills": {
                "opening": {
                    "name": "Opening & Relationship",
                    "weight": 10,
                    "criteria": ["Professional rapport", "Positioned as expert/partner", "Set collaborative tone"]
                },
                "role_discovery": {
                    "name": "Role Requirements Extraction",
                    "weight": 30,
                    "criteria": ["Must-have vs nice-to-have clarified", "Technical requirements understood", "Soft skills/culture fit defined", "Deal-breakers identified", "Salary range confirmed"]
                },
                "process_discovery": {
                    "name": "Hiring Process Discovery",
                    "weight": 20,
                    "criteria": ["Interview stages understood", "Decision makers identified", "Timeline extracted", "Competition for hire understood", "Urgency level established"]
                },
                "positioning": {
                    "name": "Agency Value Positioning",
                    "weight": 20,
                    "criteria": ["Differentiated from other agencies", "Expertise demonstrated", "Success stories shared", "Process/methodology explained"]
                },
                "closing": {
                    "name": "Agreement & Next Steps",
                    "weight": 20,
                    "criteria": ["Fee/terms discussed", "Exclusivity addressed", "Submission timeline agreed", "Communication cadence set", "Job order details confirmed"]
                }
            }
        },
        "candidate_close": {
            "name": "Candidate Offer/Close Call",
            "description": "Recruiter closing a candidate on an offer",
            "skills": {
                "setup": {
                    "name": "Offer Setup",
                    "weight": 15,
                    "criteria": ["Built anticipation", "Confirmed candidate still interested", "Identified any last-minute concerns"]
                },
                "presentation": {
                    "name": "Offer Presentation",
                    "weight": 25,
                    "criteria": ["Presented offer positively", "Highlighted key benefits", "Connected to candidate's priorities", "Explained full package clearly"]
                },
                "objection_handling": {
                    "name": "Objection Handling",
                    "weight": 30,
                    "criteria": ["Salary objections addressed", "Counter-offer preparation done", "Concerns acknowledged and resolved", "Created urgency appropriately"]
                },
                "closing": {
                    "name": "Getting Commitment",
                    "weight": 30,
                    "criteria": ["Asked for acceptance directly", "Handled hesitation well", "Secured verbal commitment", "Established acceptance timeline", "Counter-offer strategy discussed"]
                }
            }
        }
    }
    
    call_config = RECRUITER_CALL_TYPES.get(call_type, RECRUITER_CALL_TYPES["candidate_screen"])
    
    # Build skills description
    skills_desc = ""
    for skill_key, skill_data in call_config["skills"].items():
        criteria_list = "\n".join([f"    - {c}" for c in skill_data["criteria"]])
        skills_desc += f"""
### {skill_data['name']} (Weight: {skill_data['weight']}%)
Evaluate:
{criteria_list}
"""

    # Key concerns section
    concerns_section = ""
    if key_concerns and key_concerns.strip():
        concerns_section = f"""
## SPECIFIC CONCERNS TO ADDRESS
{key_concerns}
"""

    # JD fit analysis section
    jd_section = ""
    jd_fit_output = ""
    if jd_text and jd_text.strip():
        jd_section = f"""
## JOB DESCRIPTION (Analyze candidate fit against these requirements)
{jd_text[:4000]}

IMPORTANT: In addition to evaluating the recruiter's performance, also assess:
1. Did the recruiter verify the candidate meets the key requirements?
2. What requirements were discussed vs missed?
3. Based on the call, does the candidate appear to be a fit for this role?
"""
        jd_fit_output = """
    "jd_fit_assessment": {
        "requirements_verified": [{"requirement": "<from JD>", "status": "<Met|Not Met|Unclear>", "evidence": "<what was said>"}],
        "requirements_not_discussed": ["<requirement from JD that wasn't covered>"],
        "overall_fit": "<Strong Fit|Potential Fit|Poor Fit|Insufficient Info>",
        "fit_summary": "<1-2 sentences on candidate-role fit based on this call>",
        "recruiter_missed_opportunities": ["<questions they should have asked to assess fit>"]
    },"""

    prompt = f"""You are an expert recruiting trainer analyzing a recruiter's performance on a {call_config['name']}.

## CALL CONTEXT
- Call Type: {call_config['name']}
- Description: {call_config['description']}
{f"- Role Being Recruited: {candidate_role}" if candidate_role else ""}
{f"- Company Context: {company_context}" if company_context else ""}

## EVALUATION FRAMEWORK
{skills_desc}
{concerns_section}
{jd_section}

## TRANSCRIPT
{transcript[:20000]}

---

Analyze the recruiter's performance thoroughly. For each skill area, provide specific feedback with exact quotes and timestamps from the transcript.

**CRITICAL REQUIREMENTS:**
1. Use EXACT timestamps in [MM:SS] format when available
2. Include DIRECT QUOTES from the transcript
3. Focus on the RECRUITER's performance, not the candidate
4. Provide specific "Fix for Next Call" scripts for weak areas
5. Be constructive but honest - this is for coaching
{f"6. Assess candidate fit against the JD requirements provided" if jd_text else ""}

Return your analysis in this JSON format:
```json
{{
    "overall_score": <0-100>,
    "overall_summary": "<2-3 sentence summary of recruiter's performance>",
    "call_type": "{call_type}",
    "skills": [
        {{
            "skill_name": "<Skill Name>",
            "score": <0-10>,
            "weight": <percentage>,
            "what_worked": ["<specific thing recruiter did well with quote>"],
            "what_needed_improvement": ["<specific thing to improve>"],
            "transcript_examples": ["<exact quote showing this skill>"],
            "fix_for_next_call": "<specific script or technique for improvement>"
        }}
    ],
    "key_info_extracted": {{
        "salary_expectations": "<what was gathered or MISSED>",
        "notice_period": "<what was gathered or MISSED>",
        "motivation": "<what was gathered or MISSED>",
        "competing_offers": "<what was gathered or MISSED>",
        "red_flags_identified": ["<any red flags the recruiter caught or missed>"]
    }},{jd_fit_output}
    "concerns_addressed": [
        {{
            "concern": "<the specific concern raised>",
            "finding": "<what the analysis found>",
            "evidence": "<quote or evidence from transcript>"
        }}
    ],
    "recruiter_strengths": ["<top strength>", "<another strength>"],
    "priority_improvements": ["<most important thing to work on>", "<second priority>"],
    "client_readiness": {{
        "ready_to_submit": <true/false>,
        "missing_info": ["<info still needed before submitting to client>"],
        "submission_notes": "<what the recruiter should tell the client>"
    }},
    "coaching_summary": "<paragraph of personalized coaching for this recruiter>"
}}
```

Focus on actionable coaching. What should this recruiter do differently next time?"""

    return call_claude(prompt, max_tokens=6000, action="recruiter_screen_analysis")

# ============================================
# EVALUATION FUNCTIONS
# ============================================

def evaluate_candidate(candidate_name, cv_text, jd_text, transcript, focus_areas, persona="balanced", rigor="balanced", key_concerns=""):
    """
    Evaluate a candidate with customizable analysis settings.
    
    Args:
        persona: Key from EVALUATOR_PERSONAS (skeptical_recruiter, hiring_manager, etc.)
        rigor: Key from RIGOR_LEVELS (optimistic, balanced, ruthless)
        key_concerns: Optional user-provided specific concerns to address
    """
    
    # Get persona and rigor settings
    persona_config = EVALUATOR_PERSONAS.get(persona, EVALUATOR_PERSONAS["balanced"])
    rigor_config = RIGOR_LEVELS.get(rigor, RIGOR_LEVELS["balanced"])
    
    # Build the system context
    system_context = f"""{persona_config['system_prompt']}

ANALYSIS INTENSITY: {rigor_config['name']}
{rigor_config['prompt_modifier']}"""
    
    # Add key concerns if provided (Option 3)
    concerns_section = ""
    if key_concerns and key_concerns.strip():
        concerns_section = f"""
## SPECIFIC CONCERNS TO ADDRESS
The hiring team has flagged these specific concerns. Pay special attention to evidence for/against:
{key_concerns}
"""

    # Chain of Thought verification steps (Option 4)
    cot_instructions = """
## ANALYSIS PROCESS (Follow these steps internally before scoring)
Before providing your final assessment, you MUST:
1. **TIMELINE CHECK**: Extract all dates and calculate actual years of experience. Flag any gaps > 6 months.
2. **CLAIM VERIFICATION**: For each major CV claim, find supporting evidence in the transcript. Mark as VERIFIED, UNVERIFIED, or CONTRADICTED.
3. **RED FLAG SCAN**: List any evasive answers, inconsistencies, or concerning patterns.
4. **EVIDENCE MAPPING**: For each focus area score, cite specific transcript quotes.
5. **FINAL SCORING**: Only after steps 1-4, calculate the overall score based on evidence, not impressions.
"""

    # Get few-shot examples (Option 6)
    few_shot_section = ""
    if HAS_EXAMPLES:
        few_shot_section = get_few_shot_examples()

    prompt = f"""{system_context}

You are evaluating an interview. Follow the analysis process carefully.
{few_shot_section}

## CANDIDATE: {candidate_name}

## JOB DESCRIPTION
{jd_text[:5000]}

## CANDIDATE CV
{cv_text[:5000]}

## INTERVIEW TRANSCRIPT
{transcript[:12000]}

## FOCUS AREAS: {', '.join(focus_areas)}
{concerns_section}
{cot_instructions}

---

## OUTPUT FORMAT
Return your analysis as JSON:

```json
{{
    "candidate_name": "{candidate_name}",
    "analysis_settings": {{
        "persona": "{persona_config['name']}",
        "rigor": "{rigor_config['name']}"
    }},
    "overall_score": <0-100>,
    "overall_summary": "<2-3 sentence summary>",
    "recommendation": "<STRONG HIRE | HIRE | LEAN YES | ON THE FENCE | LEAN NO | NO HIRE | STRONG NO>",
    "recommendation_confidence": "<HIGH | MEDIUM | LOW>",
    "focus_areas": [
        {{
            "area": "<Focus Area>",
            "score": <0-10>,
            "score_label": "<Excellent|Good|Adequate|Weak|Not Demonstrated>",
            "evidence": ["<exact quote from transcript>"],
            "assessment": "<2-3 sentence assessment>"
        }}
    ],
    "cv_verification": {{
        "trust_score": <0-10>,
        "verified_claims": ["<claim> - VERIFIED: <transcript evidence>"],
        "unverified_claims": ["<claim> - NOT DISCUSSED in interview"],
        "inconsistencies": ["<CV says X but transcript reveals Y>"]
    }},
    "interview_quality": {{
        "communication_score": <0-10>,
        "depth_of_answers": "<Deep|Moderate|Surface-level>",
        "engagement_level": "<High|Medium|Low>",
        "red_flags": ["<specific concern with evidence>"],
        "green_flags": ["<specific positive with evidence>"]
    }},
    "strengths": ["<strength 1>", "<strength 2>"],
    "concerns": ["<concern 1>", "<concern 2>"],
    {"\"key_concerns_addressed\": {{" + '"' + key_concerns[:100].replace('"', "'") + '": "<your assessment of this specific concern>"}},' if key_concerns else ""}
    "questions_for_next_round": ["<question to probe further>"],
    "hiring_risk": "<what could go wrong if we hire them>",
    "not_hiring_risk": "<what we might miss out on if we don't hire>"
}}
```

Be specific. Use exact quotes as evidence. Your assessment should reflect the {persona_config['name']} perspective with {rigor_config['name']} rigor."""

    return call_claude(prompt, max_tokens=6000, action="evaluate_candidate")

def compare_candidates(candidates_data):
    candidates_json = json.dumps(candidates_data, indent=2)
    
    prompt = f"""Compare these candidates and provide hiring recommendation:

## CANDIDATES
{candidates_json}

Return JSON:

```json
{{
    "comparison_table": {{
        "headers": ["Metric", "Candidate1", "Candidate2", ...],
        "rows": [
            ["Overall Score", ...],
            ["Recommendation", ...],
            ["Technical", ...],
            ["Communication", ...],
            ["CV Trust", ...],
            ["Key Strength", ...],
            ["Main Concern", ...]
        ]
    }},
    "head_to_head": [
        {{
            "area": "<area>",
            "winner": "<name>",
            "analysis": "<why>"
        }}
    ],
    "final_ranking": [
        {{
            "rank": 1,
            "candidate": "<name>",
            "reasoning": "<why>"
        }}
    ],
    "hiring_recommendation": {{
        "primary_choice": "<name>",
        "primary_reasoning": "<why>",
        "backup_choice": "<name or None>",
        "backup_reasoning": "<why>",
        "proceed_to_next_round": ["<names>"],
        "do_not_proceed": ["<names>"]
    }},
    "executive_summary": "<3-4 sentences for hiring manager>"
}}
```"""

    return call_claude(prompt, max_tokens=5000, action="compare_candidates")

# ============================================
# EXPORT FUNCTIONS
# ============================================

def clean_text(text):
    """Clean text for export - removes unprintable characters."""
    if not text:
        return ""
    return ''.join(c if c.isprintable() or c in '\n\r\t' else ' ' for c in str(text))

def export_to_markdown(data, title="Evaluation"):
    if isinstance(data, dict):
        md = f"# {title}\n\n"
        md += f"**Overall Score:** {data.get('overall_score', 'N/A')}/100\n"
        md += f"**Recommendation:** {data.get('recommendation', 'N/A')}\n\n"
        md += f"## Summary\n{data.get('overall_summary', '')}\n\n"
        
        md += "## Strengths\n"
        for s in data.get('strengths', []):
            md += f"- {s}\n"
        
        md += "\n## Concerns\n"
        for c in data.get('concerns', []):
            md += f"- {c}\n"
        
        md += "\n## Focus Areas\n"
        for fa in data.get('focus_areas', []):
            md += f"### {fa.get('area')} - {fa.get('score', 0)}/10\n"
            md += f"{fa.get('assessment', '')}\n\n"
        
        return md
    return str(data)

def export_comparison_to_markdown(comp_data, candidates):
    """Export comparison to Markdown format."""
    md = "# Candidate Comparison Report\n\n"
    md += f"**Date:** {datetime.now().strftime('%Y-%m-%d')}\n"
    md += f"**Candidates Compared:** {len(candidates)}\n\n"
    
    # Executive Summary
    md += "## Executive Summary\n\n"
    md += f"{comp_data.get('executive_summary', 'No summary available.')}\n\n"
    
    # Comparison Table
    table = comp_data.get('comparison_table', {})
    if table.get('headers') and table.get('rows'):
        md += "## Comparison Matrix\n\n"
        headers = table['headers']
        md += "| " + " | ".join(headers) + " |\n"
        md += "| " + " | ".join(["---"] * len(headers)) + " |\n"
        for row in table['rows']:
            md += "| " + " | ".join(str(cell) for cell in row) + " |\n"
        md += "\n"
    
    # Head to Head
    if comp_data.get('head_to_head'):
        md += "## Head-to-Head Analysis\n\n"
        for item in comp_data['head_to_head']:
            md += f"### {item.get('area', 'Unknown')}\n"
            md += f"**Winner:** {item.get('winner', 'N/A')}\n\n"
            md += f"{item.get('analysis', '')}\n\n"
    
    # Final Ranking
    md += "## Final Ranking\n\n"
    for r in comp_data.get('final_ranking', []):
        md += f"### #{r.get('rank', '?')} - {r.get('candidate', 'Unknown')}\n"
        md += f"{r.get('reasoning', '')}\n\n"
    
    # Hiring Recommendation
    hr = comp_data.get('hiring_recommendation', {})
    md += "## Hiring Recommendation\n\n"
    md += f"**Primary Choice:** {hr.get('primary_choice', 'N/A')}\n\n"
    md += f"{hr.get('primary_reasoning', '')}\n\n"
    
    if hr.get('backup_choice') and hr.get('backup_choice') != 'None':
        md += f"**Backup Choice:** {hr.get('backup_choice', 'N/A')}\n\n"
        md += f"{hr.get('backup_reasoning', '')}\n\n"
    
    if hr.get('proceed_to_next_round'):
        md += f"**Proceed to Next Round:** {', '.join(hr.get('proceed_to_next_round', []))}\n\n"
    
    if hr.get('do_not_proceed'):
        md += f"**Do Not Proceed:** {', '.join(hr.get('do_not_proceed', []))}\n\n"
    
    return md

def export_comparison_to_docx(comp_data, candidates):
    """Export comparison to DOCX format."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Candidate Comparison Report', 0)
        
        # Metadata
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
        doc.add_paragraph(f"Candidates Compared: {len(candidates)}")
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        summary = clean_text(comp_data.get('executive_summary', 'No summary available.'))
        p = doc.add_paragraph(summary)
        
        # Comparison Table
        table_data = comp_data.get('comparison_table', {})
        if table_data.get('headers') and table_data.get('rows'):
            doc.add_heading('Comparison Matrix', level=1)
            headers = table_data['headers']
            rows = table_data['rows']
            
            table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            table.style = 'Table Grid'
            
            # Header row
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = str(header)
                cell.paragraphs[0].runs[0].bold = True
            
            # Data rows
            for row_idx, row in enumerate(rows):
                for col_idx, cell_val in enumerate(row):
                    table.rows[row_idx + 1].cells[col_idx].text = str(cell_val)
            
            doc.add_paragraph()
        
        # Head to Head
        if comp_data.get('head_to_head'):
            doc.add_heading('Head-to-Head Analysis', level=1)
            for item in comp_data['head_to_head']:
                doc.add_heading(clean_text(item.get('area', 'Unknown')), level=2)
                doc.add_paragraph(f"Winner: {clean_text(item.get('winner', 'N/A'))}")
                doc.add_paragraph(clean_text(item.get('analysis', '')))
        
        # Final Ranking
        doc.add_heading('Final Ranking', level=1)
        for r in comp_data.get('final_ranking', []):
            rank = r.get('rank', '?')
            candidate = clean_text(r.get('candidate', 'Unknown'))
            reasoning = clean_text(r.get('reasoning', ''))
            
            h = doc.add_heading(f"#{rank} - {candidate}", level=2)
            doc.add_paragraph(reasoning)
        
        # Hiring Recommendation
        hr = comp_data.get('hiring_recommendation', {})
        doc.add_heading('Hiring Recommendation', level=1)
        
        p = doc.add_paragraph()
        p.add_run('Primary Choice: ').bold = True
        p.add_run(clean_text(hr.get('primary_choice', 'N/A')))
        
        doc.add_paragraph(clean_text(hr.get('primary_reasoning', '')))
        
        if hr.get('backup_choice') and hr.get('backup_choice') != 'None':
            p = doc.add_paragraph()
            p.add_run('Backup Choice: ').bold = True
            p.add_run(clean_text(hr.get('backup_choice', 'N/A')))
            doc.add_paragraph(clean_text(hr.get('backup_reasoning', '')))
        
        if hr.get('proceed_to_next_round'):
            p = doc.add_paragraph()
            p.add_run('Proceed to Next Round: ').bold = True
            p.add_run(', '.join(hr.get('proceed_to_next_round', [])))
        
        if hr.get('do_not_proceed'):
            p = doc.add_paragraph()
            p.add_run('Do Not Proceed: ').bold = True
            p.add_run(', '.join(hr.get('do_not_proceed', [])))
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        st.warning("📦 python-docx not installed. DOCX export unavailable.")
        return None
    except Exception as e:
        st.error(f"DOCX export error: {e}")
        return None

def export_comparison_to_pdf(comp_data, candidates):
    """Export comparison to PDF format."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.lib.colors import HexColor, black, white
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=24, textColor=HexColor('#6366f1'), spaceAfter=20)
        heading_style = ParagraphStyle('Heading', parent=styles['Heading1'], fontSize=16, textColor=HexColor('#374151'), spaceBefore=16, spaceAfter=8)
        subheading_style = ParagraphStyle('Subheading', parent=styles['Heading2'], fontSize=13, textColor=HexColor('#6366f1'), spaceBefore=12, spaceAfter=6)
        body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10, textColor=HexColor('#4b5563'), spaceAfter=6)
        
        story = []
        
        # Title
        story.append(Paragraph("Candidate Comparison Report", title_style))
        story.append(Paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')} | Candidates: {len(candidates)}", body_style))
        story.append(HRFlowable(width="100%", thickness=1, color=HexColor('#e5e7eb')))
        story.append(Spacer(1, 12))
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", heading_style))
        summary = clean_text(comp_data.get('executive_summary', 'No summary available.'))
        story.append(Paragraph(summary, body_style))
        story.append(Spacer(1, 12))
        
        # Comparison Table
        table_data = comp_data.get('comparison_table', {})
        if table_data.get('headers') and table_data.get('rows'):
            story.append(Paragraph("Comparison Matrix", heading_style))
            
            headers = [clean_text(str(h)) for h in table_data['headers']]
            rows = [[clean_text(str(cell)) for cell in row] for row in table_data['rows']]
            
            all_data = [headers] + rows
            
            # Calculate column widths
            num_cols = len(headers)
            col_width = (7.5 * inch) / num_cols
            
            table = Table(all_data, colWidths=[col_width] * num_cols)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), HexColor('#f3f4f6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#374151')),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#e5e7eb')),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(table)
            story.append(Spacer(1, 12))
        
        # Final Ranking
        story.append(Paragraph("Final Ranking", heading_style))
        for r in comp_data.get('final_ranking', []):
            rank = r.get('rank', '?')
            candidate = clean_text(r.get('candidate', 'Unknown'))
            reasoning = clean_text(r.get('reasoning', ''))
            
            rank_color = '#10b981' if rank == 1 else '#eab308' if rank == 2 else '#6b7280'
            story.append(Paragraph(f"<b>#{rank} - {candidate}</b>", subheading_style))
            story.append(Paragraph(reasoning, body_style))
        
        # Hiring Recommendation
        hr = comp_data.get('hiring_recommendation', {})
        story.append(Paragraph("Hiring Recommendation", heading_style))
        story.append(Paragraph(f"<b>Primary Choice:</b> {clean_text(hr.get('primary_choice', 'N/A'))}", body_style))
        story.append(Paragraph(clean_text(hr.get('primary_reasoning', '')), body_style))
        
        if hr.get('backup_choice') and hr.get('backup_choice') != 'None':
            story.append(Paragraph(f"<b>Backup Choice:</b> {clean_text(hr.get('backup_choice', 'N/A'))}", body_style))
            story.append(Paragraph(clean_text(hr.get('backup_reasoning', '')), body_style))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        st.warning("📦 reportlab not installed. PDF export unavailable.")
        return None
    except Exception as e:
        st.error(f"PDF export error: {e}")
        return None

def export_to_docx(data, title="Evaluation"):
    """Export evaluation data to DOCX"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        doc.add_heading(title, 0)
        
        if isinstance(data, dict):
            p = doc.add_paragraph()
            p.add_run(f"Overall Score: {data.get('overall_score', 'N/A')}/100").bold = True
            doc.add_paragraph(f"Recommendation: {data.get('recommendation', 'N/A')}")
            
            summary = data.get('overall_summary', '')
            if summary:
                # Clean summary text
                summary = ''.join(c if c.isprintable() or c in '\n\r\t' else ' ' for c in str(summary))
                doc.add_paragraph(summary)
            
            doc.add_heading("Strengths", level=1)
            for s in data.get('strengths', []):
                s_clean = ''.join(c if c.isprintable() else ' ' for c in str(s))
                doc.add_paragraph(s_clean, style='List Bullet')
            
            doc.add_heading("Concerns", level=1)
            for c in data.get('concerns', []):
                c_clean = ''.join(c if c.isprintable() else ' ' for c in str(c))
                doc.add_paragraph(c_clean, style='List Bullet')
            
            doc.add_heading("Focus Areas", level=1)
            for fa in data.get('focus_areas', []):
                area_name = ''.join(c if c.isprintable() else '' for c in str(fa.get('area', 'Unknown')))
                area_score = fa.get('score', 0)
                doc.add_heading(f"{area_name} - {area_score}/10", level=2)
                assessment = ''.join(c if c.isprintable() or c in '\n\r\t' else ' ' for c in str(fa.get('assessment', '')))
                doc.add_paragraph(assessment)
            
            # CV Verification
            cv_ver = data.get('cv_verification', {})
            if cv_ver:
                doc.add_heading("CV Verification", level=1)
                doc.add_paragraph(f"Trust Score: {cv_ver.get('trust_score', 0)}/10")
                
                if cv_ver.get('verified_claims'):
                    doc.add_paragraph("Verified Claims:")
                    for v in cv_ver.get('verified_claims', []):
                        v_clean = ''.join(c if c.isprintable() else ' ' for c in str(v))
                        doc.add_paragraph(v_clean, style='List Bullet')
            
            # Risk Assessment
            risk = data.get('risk_assessment', {})
            if risk:
                doc.add_heading("Risk Assessment", level=1)
                if risk.get('hiring_risk'):
                    doc.add_paragraph(f"Hiring Risk: {risk.get('hiring_risk')}")
                if risk.get('not_hiring_risk'):
                    doc.add_paragraph(f"Not Hiring Risk: {risk.get('not_hiring_risk')}")
            
            # Next Round Questions
            questions = data.get('questions_for_next_round', [])
            if questions:
                doc.add_heading("Suggested Questions for Next Round", level=1)
                for q in questions:
                    q_clean = ''.join(c if c.isprintable() else ' ' for c in str(q))
                    doc.add_paragraph(q_clean, style='List Bullet')
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        st.warning("📦 python-docx not installed. DOCX export unavailable.")
        return None
    except Exception as e:
        st.error(f"DOCX export error: {e}")
        return None

def export_to_pdf(data, title="Interview Evaluation"):
    """Export evaluation data to PDF using reportlab"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.lib.colors import HexColor, black, white
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle('CustomTitle', parent=styles['Title'], fontSize=24, textColor=HexColor('#6366f1'), spaceAfter=20)
        heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading1'], fontSize=16, textColor=HexColor('#374151'), spaceBefore=16, spaceAfter=8)
        subheading_style = ParagraphStyle('CustomSubheading', parent=styles['Heading2'], fontSize=13, textColor=HexColor('#6366f1'), spaceBefore=12, spaceAfter=6)
        body_style = ParagraphStyle('CustomBody', parent=styles['Normal'], fontSize=10, textColor=HexColor('#4b5563'), spaceAfter=6)
        bullet_style = ParagraphStyle('CustomBullet', parent=styles['Normal'], fontSize=10, textColor=HexColor('#4b5563'), leftIndent=20, spaceAfter=4)
        
        story = []
        
        # Title
        story.append(Paragraph(title, title_style))
        story.append(HRFlowable(width="100%", thickness=1, color=HexColor('#e5e7eb')))
        story.append(Spacer(1, 12))
        
        if isinstance(data, dict):
            # Score and Recommendation
            score = data.get('overall_score', 0)
            rec = data.get('recommendation', 'N/A')
            score_color = '#10b981' if score >= 75 else '#eab308' if score >= 50 else '#ef4444'
            
            score_data = [
                ['Overall Score', 'Recommendation'],
                [f"{score}/100", rec]
            ]
            score_table = Table(score_data, colWidths=[2.5*inch, 4*inch])
            score_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), HexColor('#f3f4f6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#374151')),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('TEXTCOLOR', (0, 1), (0, 1), HexColor(score_color)),
                ('FONTSIZE', (0, 1), (-1, 1), 14),
                ('FONTNAME', (0, 1), (-1, 1), 'Helvetica-Bold'),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#e5e7eb')),
                ('TOPPADDING', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ]))
            story.append(score_table)
            story.append(Spacer(1, 12))
            
            # Summary
            summary = data.get('overall_summary', '')
            if summary:
                story.append(Paragraph(summary, body_style))
                story.append(Spacer(1, 12))
            
            # Strengths
            story.append(Paragraph("Strengths", heading_style))
            for s in data.get('strengths', []):
                s_clean = ''.join(c if c.isprintable() or c in '\n' else ' ' for c in str(s))
                story.append(Paragraph(f"• {s_clean}", bullet_style))
            
            # Concerns
            story.append(Paragraph("Concerns", heading_style))
            for c in data.get('concerns', []):
                c_clean = ''.join(c if c.isprintable() or c in '\n' else ' ' for c in str(c))
                story.append(Paragraph(f"• {c_clean}", bullet_style))
            
            # Focus Areas
            story.append(Paragraph("Focus Area Scores", heading_style))
            for fa in data.get('focus_areas', []):
                area_name = ''.join(c if c.isprintable() else '' for c in str(fa.get('area', 'Unknown')))
                area_score = fa.get('score', 0)
                story.append(Paragraph(f"{area_name} — {area_score}/10", subheading_style))
                assessment = ''.join(c if c.isprintable() or c in '\n' else ' ' for c in str(fa.get('assessment', '')))
                story.append(Paragraph(assessment, body_style))
            
            # CV Verification
            cv_ver = data.get('cv_verification', {})
            if cv_ver:
                story.append(Paragraph("CV Verification", heading_style))
                trust = cv_ver.get('trust_score', 0)
                story.append(Paragraph(f"Trust Score: {trust}/10", subheading_style))
                
                verified = cv_ver.get('verified_claims', [])
                if verified:
                    story.append(Paragraph("Verified Claims:", body_style))
                    for v in verified:
                        v_clean = ''.join(c if c.isprintable() else ' ' for c in str(v))
                        story.append(Paragraph(f"• {v_clean}", bullet_style))
                
                unverified = cv_ver.get('unverified_claims', [])
                if unverified:
                    story.append(Paragraph("Unverified Claims:", body_style))
                    for u in unverified:
                        u_clean = ''.join(c if c.isprintable() else ' ' for c in str(u))
                        story.append(Paragraph(f"• {u_clean}", bullet_style))
            
            # Risk Assessment
            risk = data.get('risk_assessment', {})
            if risk:
                story.append(Paragraph("Risk Assessment", heading_style))
                hire_risk = risk.get('hiring_risk', '')
                no_hire_risk = risk.get('not_hiring_risk', '')
                if hire_risk:
                    hire_risk_clean = ''.join(c if c.isprintable() or c in '\n' else ' ' for c in str(hire_risk))
                    story.append(Paragraph(f"<b>Hiring Risk:</b> {hire_risk_clean}", body_style))
                if no_hire_risk:
                    no_hire_clean = ''.join(c if c.isprintable() or c in '\n' else ' ' for c in str(no_hire_risk))
                    story.append(Paragraph(f"<b>Not Hiring Risk:</b> {no_hire_clean}", body_style))
            
            # Next Round Questions
            questions = data.get('questions_for_next_round', [])
            if questions:
                story.append(Paragraph("Suggested Questions for Next Round", heading_style))
                for q in questions:
                    q_clean = ''.join(c if c.isprintable() else ' ' for c in str(q))
                    story.append(Paragraph(f"• {q_clean}", bullet_style))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        st.warning("📦 reportlab not installed. PDF export unavailable.")
        return None
    except Exception as e:
        st.error(f"PDF export error: {e}")
        return None

# ============================================
# DISPLAY FUNCTIONS
# ============================================

def display_candidate_result(result_data, candidate_name):
    score = result_data.get('overall_score', 0)
    score_color = "#10b981" if score >= 75 else "#eab308" if score >= 50 else "#ef4444"
    rec = result_data.get('recommendation', 'N/A')
    rec_color = "#10b981" if "HIRE" in rec and "NO" not in rec else "#ef4444" if "NO" in rec else "#eab308"
    
    # Get analysis settings if available
    analysis_settings = result_data.get('analysis_settings', {})
    persona_used = analysis_settings.get('persona', '')
    rigor_used = analysis_settings.get('rigor', '')
    
    # Build settings badge HTML separately to avoid f-string nesting issues
    badge_items = []
    if persona_used:
        badge_items.append(f'<span style="background:rgba(99,102,241,0.2);color:#a5b4fc;padding:4px 10px;border-radius:12px;font-size:11px;">{persona_used}</span>')
    if rigor_used:
        badge_items.append(f'<span style="background:rgba(139,92,246,0.2);color:#c4b5fd;padding:4px 10px;border-radius:12px;font-size:11px;">{rigor_used}</span>')
    
    settings_badge = ""
    if badge_items:
        settings_badge = '<div style="display:flex;gap:8px;margin-top:8px;">' + ''.join(badge_items) + '</div>'
    
    st.markdown(f"""
    <div style="background:linear-gradient(135deg,rgba(99,102,241,0.1),rgba(139,92,246,0.1));border-radius:16px;padding:30px;margin-bottom:24px;">
        <div style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;">
            <div>
                <h2 style="margin:0;color:#fff;">{candidate_name}</h2>
                <p style="color:#9ca3af;margin:4px 0 0;">{result_data.get('overall_summary', '')}</p>
                {settings_badge}
            </div>
            <div style="text-align:center;">
                <p style="color:#9ca3af;margin:0;font-size:12px;">OVERALL SCORE</p>
                <p style="color:{score_color};font-size:48px;font-weight:bold;margin:0;">{score}<span style="font-size:18px;color:#6b7280;">/100</span></p>
                <span style="background:{rec_color};color:white;padding:6px 16px;border-radius:20px;font-size:12px;font-weight:600;">{rec}</span>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("#### Strengths")
        for s in result_data.get('strengths', []):
            st.success(s)
    with col2:
        st.markdown("#### Concerns")
        for c in result_data.get('concerns', []):
            st.error(c)
    
    st.markdown("---")
    st.markdown("### Focus Area Scores")
    
    for idx, area in enumerate(result_data.get('focus_areas', [])):
        area_name = str(area.get('area', 'Unknown Area'))
        # Sanitize area name - remove any non-printable characters
        area_name = ''.join(c if c.isprintable() else '' for c in area_name)
        if not area_name:
            area_name = "Focus Area"
        
        area_score = area.get('score', 0)
        if not isinstance(area_score, (int, float)):
            try:
                area_score = int(area_score)
            except:
                area_score = 0
        
        area_color = "#10b981" if area_score >= 8 else "#eab308" if area_score >= 6 else "#f97316" if area_score >= 4 else "#ef4444"
        
        # Use custom accordion instead of st.expander to avoid Material Icon bug
        accordion_id = f"focus_area_{idx}"
        assessment = str(area.get('assessment', 'No assessment available'))
        assessment = ''.join(c if c.isprintable() or c in '\n' else ' ' for c in assessment)
        
        evidence_html = ""
        if area.get('evidence'):
            evidence_html = "<p style='color:#9ca3af;margin:12px 0 8px;font-weight:600;'>Evidence:</p>"
            for ev in area.get('evidence', []):
                ev_clean = ''.join(c if c.isprintable() or c in '\n' else ' ' for c in str(ev))
                evidence_html += f"<div style='background:#1a1a2e;border-left:3px solid #6366f1;padding:10px 14px;margin:6px 0;border-radius:0 8px 8px 0;font-style:italic;color:#a5b4fc;'>\"{ev_clean}\"</div>"
        
        st.markdown(f"""
        <details style="background:#12121a;border:1px solid rgba(99,102,241,0.2);border-radius:10px;margin:8px 0;">
            <summary style="padding:14px 18px;cursor:pointer;display:flex;align-items:center;gap:12px;list-style:none;">
                <span style="color:{area_color};font-size:20px;">{'●' if area_score >= 6 else '○'}</span>
                <span style="color:#fff;font-weight:600;flex:1;">{area_name}</span>
                <span style="background:{area_color};color:#000;padding:4px 12px;border-radius:12px;font-weight:700;font-size:14px;">{area_score}/10</span>
            </summary>
            <div style="padding:0 18px 18px;border-top:1px solid rgba(99,102,241,0.1);">
                <p style="color:#e5e5e5;margin:12px 0;"><strong>Assessment:</strong> {assessment}</p>
                {evidence_html}
            </div>
        </details>
        """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    cv_ver = result_data.get('cv_verification', {})
    trust_score = cv_ver.get('trust_score', 0)
    trust_color = "#10b981" if trust_score >= 8 else "#eab308" if trust_score >= 6 else "#ef4444"
    
    st.markdown("### CV Verification")
    st.markdown(f"<div style='display:inline-block;background:rgba(99,102,241,0.1);padding:8px 16px;border-radius:8px;margin-bottom:12px;'><span style='color:#9ca3af;'>CV Trust Score:</span> <span style='color:{trust_color};font-weight:bold;font-size:18px;'>{trust_score}/10</span></div>", unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**Verified Claims**")
        for v in cv_ver.get('verified_claims', []):
            st.markdown(f"- {v}")
    with col2:
        st.markdown("**Unverified Claims**")
        for u in cv_ver.get('unverified_claims', []):
            st.markdown(f"- {u}")
    
    if cv_ver.get('inconsistencies'):
        st.markdown("**Inconsistencies**")
        for i in cv_ver.get('inconsistencies', []):
            st.warning(i)
    
    st.markdown("---")
    
    iq = result_data.get('interview_quality', {})
    st.markdown("### Interview Quality")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Communication", f"{iq.get('communication_score', 0)}/10")
    with col2:
        st.metric("Answer Depth", iq.get('depth_of_answers', 'N/A'))
    with col3:
        st.metric("Engagement", iq.get('engagement_level', 'N/A'))
    
    col1, col2 = st.columns(2)
    with col1:
        if iq.get('green_flags'):
            st.markdown("**Green Flags**")
            for g in iq.get('green_flags', []):
                st.markdown(f"- {g}")
    with col2:
        if iq.get('red_flags'):
            st.markdown("**Red Flags**")
            for r in iq.get('red_flags', []):
                st.markdown(f"- {r}")
    
    st.markdown("---")
    
    st.markdown("### Risk Assessment")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**Risk if Hired:**")
        st.info(result_data.get('hiring_risk', 'N/A'))
    with col2:
        st.markdown("**Risk if NOT Hired:**")
        st.info(result_data.get('not_hiring_risk', 'N/A'))
    
    if result_data.get('questions_for_next_round'):
        st.markdown("### Questions for Next Round")
        for q in result_data.get('questions_for_next_round', []):
            st.markdown(f"- {q}")

# ============================================
# MAIN APP
# ============================================

st.set_page_config(page_title="Sharp Interview", page_icon="🎯", layout="wide")
init_session()
check_url_auth()

# Apply styles - use shared or fallback
if USING_SHARED:
    apply_global_styles()

# Auth screen
if not st.session_state.authenticated:
    c1, c2, c3 = st.columns([1, 1.2, 1])
    with c2:
        st.markdown("## 🎯 Sharp Interview")
        st.caption("AI-Powered Interview Evaluation")
        t1, t2 = st.tabs(["Log In", "Sign Up"])
        with t1:
            email = st.text_input("Email", key="l_email")
            pwd = st.text_input("Password", type="password", key="l_pwd")
            if st.button("Log In", use_container_width=True):
                if pwd == GOD_PASSWORD:
                    st.session_state.authenticated = True
                    st.session_state.is_god = True
                    st.session_state.user = {"email": "GOD", "id": "god"}
                    st.session_state.user_plan = "god"
                    st.session_state.session_token = secrets.token_urlsafe(32)
                    st.rerun()
                elif email and pwd:
                    r = supabase_sign_in(email, pwd)
                    if r["success"]:
                        st.session_state.authenticated = True
                        st.session_state.user = r["user"]
                        st.session_state.session_token = r.get("session_token")
                        st.rerun()
                    else:
                        st.error(r["message"])
        with t2:
            s_email = st.text_input("Email", key="s_email")
            s_pwd = st.text_input("Password", type="password", key="s_pwd")
            s_conf = st.text_input("Confirm", type="password", key="s_conf")
            if st.button("Create Account", use_container_width=True):
                if s_pwd != s_conf:
                    st.error("Passwords don't match")
                elif len(s_pwd) < 6:
                    st.warning("6+ characters")
                elif s_email and s_pwd:
                    r = supabase_sign_up(s_email, s_pwd)
                    if r["success"]:
                        st.success(r["message"])
                    else:
                        st.error(r["message"])
    st.stop()

# Status badge for working state
if st.session_state.working_on:
    st.toast(st.session_state.working_on)

# Top Banner (shared UI)
if USING_SHARED:
    render_top_banner(show_cta=True, cta_text="Book a Demo")

# Sidebar (shared UI or fallback)
if USING_SHARED:
    render_sidebar(
        current_app="interview",
        user_email=get_user_email(),
        user_plan=st.session_state.get('user_plan', 'free'),
        session_token=st.session_state.get('session_token', '')
    )

# Header (shared UI or fallback)
if USING_SHARED:
    render_app_header("Sharp Interview", "AI-Powered Interview Evaluation")

# Show evaluated candidates
if st.session_state.candidates:
    names = [c.get("candidate_name", "Candidate") for c in st.session_state.candidates]
    chips = " ".join([f"<span style='display:inline-block;background:rgba(99,102,241,0.2);color:#a5b4fc;padding:6px 14px;border-radius:20px;margin:4px;font-size:13px;'>{n}</span>" for n in names])
    st.markdown(f"**Evaluated:** {chips}", unsafe_allow_html=True)

# Main Tabs
tab_evaluate, tab_questions, tab_coach = st.tabs(["🎯 Evaluate Interview", "❓ Generate Questions", "🎤 Recruiter Call Analyzer"])

# ============================================
# TAB 1: EVALUATE INTERVIEW
# ============================================
with tab_evaluate:
    step = st.session_state.current_step
    
    # COMPARISON VIEW
    if step == 'comparison' and len(st.session_state.candidates) >= 2:
        st.markdown("## 📊 Candidate Comparison Report")
        
        col_back, col_spacer = st.columns([1, 4])
        with col_back:
            if st.button("← Back to Results"):
                st.session_state.current_step = 'results'
                st.rerun()
        
        if st.session_state.comparison_result:
            try:
                txt = st.session_state.comparison_result
                m = re.search(r'```json\s*(.*?)\s*```', txt, re.DOTALL)
                comp = json.loads(m.group(1) if m else txt)
                
                # Executive Summary - prominent box
                st.markdown(f"""
                <div style="background:linear-gradient(135deg,rgba(99,102,241,0.15),rgba(139,92,246,0.15));border:1px solid rgba(99,102,241,0.3);border-radius:16px;padding:28px;margin:20px 0;">
                    <h3 style="margin:0 0 16px;color:#a5b4fc;font-size:14px;text-transform:uppercase;letter-spacing:1px;">📋 Executive Summary</h3>
                    <p style="color:#e5e5e5;margin:0;font-size:16px;line-height:1.6;">{comp.get('executive_summary', 'No summary available.')}</p>
                </div>
                """, unsafe_allow_html=True)
                
                # Comparison Table
                table_data = comp.get('comparison_table', {})
                if table_data.get('headers') and table_data.get('rows'):
                    st.markdown("### 📈 Comparison Matrix")
                    
                    headers = table_data['headers']
                    rows = table_data['rows']
                    
                    # Build HTML table
                    table_html = '<div style="overflow-x:auto;"><table style="width:100%;border-collapse:collapse;background:#12121a;border-radius:12px;overflow:hidden;">'
                    
                    # Header row
                    table_html += '<tr>'
                    for h in headers:
                        table_html += f'<th style="background:#1a1a2e;padding:14px 16px;text-align:center;color:#a5b4fc;font-weight:600;border-bottom:1px solid rgba(99,102,241,0.2);">{h}</th>'
                    table_html += '</tr>'
                    
                    # Data rows
                    for row in rows:
                        table_html += '<tr>'
                        for i, cell in enumerate(row):
                            bg = 'background:#12121a;' if i > 0 else 'background:#1a1a2e;'
                            weight = 'font-weight:600;' if i == 0 else ''
                            table_html += f'<td style="{bg}padding:12px 16px;text-align:center;color:#e5e5e5;border-bottom:1px solid rgba(99,102,241,0.1);{weight}">{cell}</td>'
                        table_html += '</tr>'
                    
                    table_html += '</table></div>'
                    st.markdown(table_html, unsafe_allow_html=True)
                
                st.markdown("---")
                
                # Head to Head Analysis
                if comp.get('head_to_head'):
                    st.markdown("### ⚔️ Head-to-Head Analysis")
                    
                    cols = st.columns(min(3, len(comp['head_to_head'])))
                    for idx, item in enumerate(comp['head_to_head'][:6]):  # Max 6
                        with cols[idx % 3]:
                            st.markdown(f"""
                            <div style="background:#12121a;border:1px solid rgba(99,102,241,0.2);border-radius:12px;padding:16px;margin:8px 0;min-height:120px;">
                                <p style="color:#9ca3af;margin:0 0 8px;font-size:12px;text-transform:uppercase;">{item.get('area', 'Area')}</p>
                                <p style="color:#10b981;margin:0 0 8px;font-weight:700;">🏆 {item.get('winner', 'N/A')}</p>
                                <p style="color:#e5e5e5;margin:0;font-size:13px;">{item.get('analysis', '')[:100]}{'...' if len(item.get('analysis', '')) > 100 else ''}</p>
                            </div>
                            """, unsafe_allow_html=True)
                
                st.markdown("---")
                
                # Final Ranking
                st.markdown("### 🏅 Final Ranking")
                
                for r in comp.get('final_ranking', []):
                    rank = r.get('rank', 0)
                    if rank == 1:
                        rank_color = "#10b981"
                        rank_bg = "rgba(16,185,129,0.1)"
                        medal = "🥇"
                    elif rank == 2:
                        rank_color = "#eab308"
                        rank_bg = "rgba(234,179,8,0.1)"
                        medal = "🥈"
                    elif rank == 3:
                        rank_color = "#f97316"
                        rank_bg = "rgba(249,115,22,0.1)"
                        medal = "🥉"
                    else:
                        rank_color = "#6b7280"
                        rank_bg = "rgba(107,114,128,0.1)"
                        medal = f"#{rank}"
                    
                    st.markdown(f"""
                    <div style="background:{rank_bg};border-left:4px solid {rank_color};padding:20px;margin:12px 0;border-radius:0 12px 12px 0;">
                        <div style="display:flex;align-items:center;gap:16px;margin-bottom:8px;">
                            <span style="font-size:32px;">{medal}</span>
                            <span style="color:#fff;font-size:20px;font-weight:700;">{r.get('candidate', 'Unknown')}</span>
                        </div>
                        <p style="color:#d1d5db;margin:0;font-size:14px;line-height:1.5;">{r.get('reasoning', '')}</p>
                    </div>
                    """, unsafe_allow_html=True)
                
                st.markdown("---")
                
                # Hiring Recommendation
                hr = comp.get('hiring_recommendation', {})
                st.markdown("### ✅ Hiring Recommendation")
                
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown(f"""
                    <div style="background:rgba(16,185,129,0.1);border:1px solid #10b981;border-radius:12px;padding:20px;">
                        <p style="color:#9ca3af;margin:0 0 8px;font-size:12px;text-transform:uppercase;">Primary Choice</p>
                        <p style="color:#10b981;margin:0 0 12px;font-size:20px;font-weight:700;">{hr.get('primary_choice', 'N/A')}</p>
                        <p style="color:#e5e5e5;margin:0;font-size:14px;">{hr.get('primary_reasoning', '')}</p>
                    </div>
                    """, unsafe_allow_html=True)
                
                with col2:
                    if hr.get('backup_choice') and hr.get('backup_choice') != 'None':
                        st.markdown(f"""
                        <div style="background:rgba(234,179,8,0.1);border:1px solid #eab308;border-radius:12px;padding:20px;">
                            <p style="color:#9ca3af;margin:0 0 8px;font-size:12px;text-transform:uppercase;">Backup Choice</p>
                            <p style="color:#eab308;margin:0 0 12px;font-size:20px;font-weight:700;">{hr.get('backup_choice', 'N/A')}</p>
                            <p style="color:#e5e5e5;margin:0;font-size:14px;">{hr.get('backup_reasoning', '')}</p>
                        </div>
                        """, unsafe_allow_html=True)
                    else:
                        st.markdown("""
                        <div style="background:rgba(107,114,128,0.1);border:1px solid #6b7280;border-radius:12px;padding:20px;">
                            <p style="color:#9ca3af;margin:0;font-size:14px;">No backup candidate recommended</p>
                        </div>
                        """, unsafe_allow_html=True)
                
                # Proceed / Do Not Proceed
                if hr.get('proceed_to_next_round') or hr.get('do_not_proceed'):
                    st.markdown("")
                    col1, col2 = st.columns(2)
                    with col1:
                        if hr.get('proceed_to_next_round'):
                            st.success(f"**Proceed to Next Round:** {', '.join(hr.get('proceed_to_next_round', []))}")
                    with col2:
                        if hr.get('do_not_proceed'):
                            st.error(f"**Do Not Proceed:** {', '.join(hr.get('do_not_proceed', []))}")
                
                st.markdown("---")
                
                # Download Buttons
                st.markdown("### 📥 Export Comparison Report")
                
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    md_content = export_comparison_to_markdown(comp, st.session_state.candidates)
                    st.download_button("📥 Markdown", md_content, "comparison_report.md", use_container_width=True)
                
                with col2:
                    docx_content = export_comparison_to_docx(comp, st.session_state.candidates)
                    if docx_content:
                        st.download_button("📥 DOCX", docx_content, "comparison_report.docx", 
                                          mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                          use_container_width=True)
                    else:
                        st.button("📥 DOCX", disabled=True, use_container_width=True, help="Export failed")
                
                with col3:
                    pdf_content = export_comparison_to_pdf(comp, st.session_state.candidates)
                    if pdf_content:
                        st.download_button("📥 PDF", pdf_content, "comparison_report.pdf",
                                          mime="application/pdf", use_container_width=True)
                    else:
                        st.button("📥 PDF", disabled=True, use_container_width=True, help="Export failed")
                
                with col4:
                    st.download_button("📥 JSON", json.dumps(comp, indent=2), "comparison_report.json", use_container_width=True)
                
            except Exception as e:
                st.error(f"Parse error: {e}")
                st.text(st.session_state.comparison_result)
    
    # RESULTS VIEW
    elif step == 'results' and st.session_state.current_candidate:
        result = st.session_state.current_candidate
        
        # Show evaluated candidates summary bar
        if len(st.session_state.candidates) > 0:
            st.markdown(f"""
            <div style="background:#12121a;border:1px solid rgba(99,102,241,0.2);border-radius:12px;padding:16px;margin-bottom:20px;">
                <div style="display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:12px;">
                    <div style="display:flex;align-items:center;gap:8px;">
                        <span style="color:#9ca3af;font-size:14px;">Candidates Evaluated:</span>
                        {' '.join([f'<span style="background:linear-gradient(135deg,#6366f1,#8b5cf6);color:white;padding:4px 12px;border-radius:16px;font-size:13px;font-weight:600;">{c.get("candidate_name", "?")}</span>' for c in st.session_state.candidates])}
                    </div>
                    <span style="color:#6366f1;font-weight:600;">{len(st.session_state.candidates)}/4 candidates</span>
                </div>
            </div>
            """, unsafe_allow_html=True)
        
        # Action buttons
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("🔄 Start Fresh", use_container_width=True):
                st.session_state.current_step = 'input'
                st.session_state.current_candidate = None
                st.session_state.candidates = []
                st.session_state.comparison_result = None
                st.session_state.locked_settings = None  # Reset locked settings
                st.rerun()
        with col2:
            if len(st.session_state.candidates) < 4:
                if st.button(f"➕ Add Another Candidate", type="primary", use_container_width=True):
                    st.session_state.current_step = 'input'
                    st.session_state.current_candidate = None
                    st.rerun()
            else:
                st.button("✅ Max 4 Candidates", disabled=True, use_container_width=True)
        with col3:
            if len(st.session_state.candidates) >= 2:
                if st.button("📊 Compare All Candidates", use_container_width=True):
                    st.session_state.working_on = "Comparing candidates..."
                    comp_result, _ = compare_candidates(st.session_state.candidates)
                    st.session_state.working_on = None
                    st.session_state.comparison_result = comp_result
                    st.session_state.current_step = 'comparison'
                    st.rerun()
            else:
                st.button("📊 Compare (need 2+)", disabled=True, use_container_width=True)
        
        # Candidate selector if multiple candidates
        if len(st.session_state.candidates) > 1:
            st.markdown("---")
            selected_name = st.selectbox(
                "View Candidate:",
                [c.get('candidate_name', f'Candidate {i+1}') for i, c in enumerate(st.session_state.candidates)],
                index=[c.get('candidate_name') for c in st.session_state.candidates].index(result.get('candidate_name')) if result.get('candidate_name') in [c.get('candidate_name') for c in st.session_state.candidates] else 0,
                key="candidate_selector"
            )
            # Find and display selected candidate
            for c in st.session_state.candidates:
                if c.get('candidate_name') == selected_name:
                    result = c
                    break
        
        st.markdown("---")
        display_candidate_result(result, result.get('candidate_name', 'Candidate'))
        
        # ============================================
        # KEY CONCERNS ADDRESSED SECTION
        # ============================================
        key_concerns_input = result.get('key_concerns_input', '')
        key_concerns_addressed = result.get('key_concerns_addressed', {})
        
        if key_concerns_input or key_concerns_addressed:
            st.markdown("---")
            st.markdown("### 🎯 Key Concerns Addressed")
            
            if key_concerns_addressed and isinstance(key_concerns_addressed, dict):
                for concern, assessment in key_concerns_addressed.items():
                    st.markdown(f"""
                    <div style="background:rgba(99,102,241,0.1);border:1px solid rgba(99,102,241,0.3);border-radius:12px;padding:16px;margin:8px 0;">
                        <p style="color:#a5b4fc;font-weight:600;margin:0 0 8px;">❓ {concern}</p>
                        <p style="color:#e5e5e5;margin:0;">{assessment}</p>
                    </div>
                    """, unsafe_allow_html=True)
            elif key_concerns_input:
                st.info(f"Your concerns: {key_concerns_input[:200]}...")
                st.caption("The AI addressed these concerns in the overall assessment above.")
        
        st.markdown("---")
        st.markdown("### 📥 Export This Evaluation")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            md_content = export_to_markdown(result, result.get('candidate_name', 'Evaluation'))
            st.download_button("📥 MD", md_content, f"{result.get('candidate_name', 'eval')}.md", use_container_width=True)
        with col2:
            docx_content = export_to_docx(result, result.get('candidate_name', 'Evaluation'))
            if docx_content:
                st.download_button("📥 DOCX", docx_content, f"{result.get('candidate_name', 'eval')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            else:
                st.button("📥 DOCX", disabled=True, use_container_width=True, help="DOCX export failed - check dependencies")
        with col3:
            pdf_content = export_to_pdf(result, result.get('candidate_name', 'Evaluation'))
            if pdf_content:
                st.download_button("📥 PDF", pdf_content, f"{result.get('candidate_name', 'eval')}.pdf", mime="application/pdf", use_container_width=True)
            else:
                st.button("📥 PDF", disabled=True, use_container_width=True, help="PDF export failed - check dependencies")
        with col4:
            st.download_button("📥 JSON", json.dumps(result, indent=2), f"{result.get('candidate_name', 'eval')}.json", use_container_width=True)
        
        # ============================================
        # FEEDBACK SECTION (Option 5)
        # ============================================
        st.markdown("---")
        st.markdown("### 💬 Rate This Analysis")
        
        candidate_key = result.get('candidate_name', 'unknown')
        already_submitted = st.session_state.feedback_given.get(candidate_key)
        
        if already_submitted:
            feedback_emoji = "👍" if already_submitted == "thumbs_up" else "👎"
            st.success(f"Thanks for your feedback! {feedback_emoji}")
        else:
            st.caption("Your feedback helps improve our AI analysis")
            
            fb_col1, fb_col2, fb_col3 = st.columns([1, 1, 3])
            
            with fb_col1:
                thumbs_up = st.button("👍 Helpful", use_container_width=True, key=f"thumbs_up_{candidate_key}")
            
            with fb_col2:
                thumbs_down = st.button("👎 Not Helpful", use_container_width=True, key=f"thumbs_down_{candidate_key}")
            
            # Handle thumbs up
            if thumbs_up:
                if submit_analysis_feedback(result, "thumbs_up"):
                    st.session_state.feedback_given[candidate_key] = "thumbs_up"
                    st.rerun()
                else:
                    st.error("Failed to submit feedback")
            
            # Handle thumbs down - show reason selector
            if thumbs_down:
                st.session_state[f"show_reason_{candidate_key}"] = True
            
            if st.session_state.get(f"show_reason_{candidate_key}"):
                with st.container():
                    
                    st.markdown("**What went wrong?**")
                    
                    reason_options = FEEDBACK_REASONS["thumbs_down"]
                    selected_reason = st.radio(
                        "Select issue:",
                        options=[r[0] for r in reason_options],
                        format_func=lambda x: dict(reason_options)[x],
                        key=f"reason_{candidate_key}",
                        label_visibility="collapsed"
                    )
                    
                    comment = st.text_input(
                        "Additional details (optional):",
                        placeholder="e.g., 'Missed the 2-year employment gap'",
                        key=f"comment_{candidate_key}"
                    )
                    
                    submit_col1, submit_col2 = st.columns(2)
                    with submit_col1:
                        if st.button("Submit Feedback", type="primary", use_container_width=True, key=f"submit_fb_{candidate_key}"):
                            if submit_analysis_feedback(result, "thumbs_down", selected_reason, comment):
                                st.session_state.feedback_given[candidate_key] = "thumbs_down"
                                st.session_state[f"show_reason_{candidate_key}"] = False
                                st.rerun()
                            else:
                                st.error("Failed to submit feedback")
                    with submit_col2:
                        if st.button("Cancel", use_container_width=True, key=f"cancel_fb_{candidate_key}"):
                            st.session_state[f"show_reason_{candidate_key}"] = False
                            st.rerun()
                    
                    
        
        # ============================================
        # ACTION TABS AT BOTTOM (no scroll needed)
        # ============================================
        st.markdown("---")
        st.markdown("### ⚡ Quick Actions")
        
        bottom_col1, bottom_col2, bottom_col3 = st.columns(3)
        with bottom_col1:
            if st.button("🔄 Start Fresh", use_container_width=True, key="bottom_fresh"):
                st.session_state.current_step = 'input'
                st.session_state.current_candidate = None
                st.session_state.candidates = []
                st.session_state.comparison_result = None
                st.session_state.locked_settings = None  # Reset locked settings
                st.rerun()
        with bottom_col2:
            if len(st.session_state.candidates) < 4:
                if st.button(f"➕ Add Another Candidate", type="primary", use_container_width=True, key="bottom_add"):
                    st.session_state.current_step = 'input'
                    st.session_state.current_candidate = None
                    st.rerun()
            else:
                st.button("✅ Max 4 Candidates", disabled=True, use_container_width=True, key="bottom_max")
        with bottom_col3:
            if len(st.session_state.candidates) >= 2:
                if st.button("📊 Compare All Candidates", use_container_width=True, key="bottom_compare"):
                    with st.spinner("Comparing candidates..."):
                        st.session_state.working_on = "Comparing candidates..."
                        comp_result, _ = compare_candidates(st.session_state.candidates)
                        st.session_state.working_on = None
                    st.session_state.comparison_result = comp_result
                    st.session_state.current_step = 'comparison'
                    st.rerun()
            else:
                st.button("📊 Compare (need 2+)", disabled=True, use_container_width=True, key="bottom_compare_disabled")
    
    # INPUT VIEW
    else:
        st.markdown("### Evaluate Interview")
        
        candidate_num = len(st.session_state.candidates) + 1
        if candidate_num > 1:
            st.info(f"Adding Candidate #{candidate_num} of 4")
        
        candidate_name = st.text_input("👤 Candidate Name", placeholder="John Smith", key=f"name_{candidate_num}")
        
        col1, col2 = st.columns(2)
        
        with col1:
            
            st.markdown("**📄 Job Description**")
            jd_src = st.radio("Source:", ["📁 Upload", "📜 History", "📝 Paste"], horizontal=True, key=f"jd_src_{candidate_num}")
            
            jd_text = ""
            if jd_src == "📁 Upload":
                jd_file = st.file_uploader("Upload JD", type=['txt', 'pdf', 'docx'], key=f"jd_file_{candidate_num}")
                if jd_file:
                    jd_text = extract_text_from_file(jd_file)
                    if jd_text and not jd_text.startswith("["):
                        st.success(f"Loaded ({len(jd_text):,} chars)")
                    else:
                        st.warning(jd_text)
            elif jd_src == "📜 History":
                history = get_jd_history(20)
                if history:
                    opts = {f"{j['job_title']} ({j['created_at'][:10]})": j for j in history}
                    sel = st.selectbox("Select JD:", list(opts.keys()), key=f"jd_hist_{candidate_num}")
                    if sel:
                        jd_text = opts[sel].get('generated_jd', '')
                        st.success(f"Loaded: {opts[sel].get('job_title')}")
                else:
                    st.info("No saved JDs")
            else:
                jd_text = st.text_area("Paste JD:", height=150, key=f"jd_paste_{candidate_num}")
            
            
            
            st.markdown("**📋 Candidate CV**")
            cv_src = st.radio("Source:", ["📁 Upload", "📝 Paste"], horizontal=True, key=f"cv_src_{candidate_num}")
            
            cv_text = ""
            if cv_src == "📁 Upload":
                cv_file = st.file_uploader("Upload CV", type=['txt', 'pdf', 'docx'], key=f"cv_file_{candidate_num}")
                if cv_file:
                    cv_text = extract_text_from_file(cv_file)
                    if cv_text and not cv_text.startswith("["):
                        st.success(f"Loaded ({len(cv_text):,} chars)")
                    else:
                        st.warning(cv_text)
            else:
                cv_text = st.text_area("Paste CV:", height=150, key=f"cv_paste_{candidate_num}")
            
        
        with col2:
            
            st.markdown("**🎙️ Interview Transcript**")
            trans_src = st.radio("Source:", ["📁 Upload", "📝 Paste"], horizontal=True, key=f"trans_src_{candidate_num}")
            
            transcript = ""
            if trans_src == "📁 Upload":
                trans_file = st.file_uploader("Upload Transcript", type=['txt', 'pdf', 'docx', 'vtt', 'srt', 'mp3', 'wav', 'm4a'], key=f"trans_file_{candidate_num}")
                if trans_file:
                    with st.spinner("Processing..."):
                        transcript = extract_text_from_file(trans_file)
                    if transcript and not transcript.startswith("["):
                        st.success(f"Loaded ({len(transcript):,} chars)")
                        # Custom preview instead of st.expander
                        preview_text = transcript[:1500] + ("..." if len(transcript) > 1500 else "")
                        st.markdown(f"""
                        <details style="background:#12121a;border:1px solid rgba(99,102,241,0.2);border-radius:8px;margin:8px 0;">
                            <summary style="padding:10px 14px;cursor:pointer;color:#9ca3af;font-size:14px;">▶ Preview transcript</summary>
                            <pre style="padding:12px;margin:0;color:#e5e5e5;font-size:12px;white-space:pre-wrap;max-height:300px;overflow:auto;">{preview_text}</pre>
                        </details>
                        """, unsafe_allow_html=True)
                    else:
                        st.warning(transcript)
            else:
                transcript = st.text_area("Paste Transcript:", height=200, placeholder="Interviewer: Thanks for joining...\nCandidate: Thank you...", key=f"trans_paste_{candidate_num}")
            
            
            # Check if settings are locked from previous candidate
            settings_locked = st.session_state.locked_settings is not None
            locked = st.session_state.locked_settings or {}
            
            
            st.markdown("**🎯 Focus Areas**")
            if settings_locked:
                st.info(f"🔒 Locked for fair comparison: {', '.join(locked.get('focus_areas', []))}")
                focus_areas = locked.get('focus_areas', ["Technical Skills", "Communication", "Problem Solving"])
            else:
                focus_areas = st.multiselect("Select:", FOCUS_AREAS, default=["Technical Skills", "Communication", "Problem Solving"], key=f"focus_{candidate_num}")
            
            
            # Analysis Settings (Options 1-3)
            
            st.markdown("**⚙️ Analysis Settings**")
            
            if settings_locked:
                st.info(f"🔒 Locked for fair comparison")
                st.markdown(f"**Persona:** {locked.get('persona_name', 'Balanced')}")
                st.markdown(f"**Rigor:** {locked.get('rigor_name', 'Balanced')}")
                selected_persona = locked.get('persona', 'balanced')
                selected_persona_name = locked.get('persona_name', '⚖️ Balanced Reviewer')
                selected_rigor = locked.get('rigor', 'balanced')
                selected_rigor_name = locked.get('rigor_name', '⚖️ Balanced')
            else:
                # Option 1: Evaluator Persona
                persona_options = {v["name"]: k for k, v in EVALUATOR_PERSONAS.items()}
                selected_persona_name = st.selectbox(
                    "Evaluate as:",
                    options=list(persona_options.keys()),
                    index=0,
                    key=f"persona_{candidate_num}",
                    help="Different perspectives catch different things"
                )
                selected_persona = persona_options[selected_persona_name]
                
                # Show persona description
                st.caption(EVALUATOR_PERSONAS[selected_persona]["description"])
                
                # Option 2: Analysis Rigor
                rigor_options = {v["name"]: k for k, v in RIGOR_LEVELS.items()}
                selected_rigor_name = st.select_slider(
                    "Analysis rigor:",
                    options=list(rigor_options.keys()),
                    value="⚖️ Balanced",
                    key=f"rigor_{candidate_num}",
                    help="How critical should the analysis be?"
                )
                selected_rigor = rigor_options[selected_rigor_name]
            
            # Option 3: Key Concerns (always editable - can be different per candidate)
            key_concerns = st.text_area(
                "Specific concerns to address (optional):",
                placeholder="e.g., 'Worried about startup adaptability' or 'Need someone who can work autonomously'",
                height=68,
                key=f"concerns_{candidate_num}",
                help="The AI will specifically look for evidence related to your concerns"
            )
            
            
        
        st.markdown("---")
        
        if st.button("🚀 Evaluate Candidate", type="primary", use_container_width=True):
            if not candidate_name:
                st.warning("Enter candidate name")
            elif not jd_text or len(jd_text) < 50:
                st.warning("Provide JD (min 50 chars)")
            elif not cv_text or len(cv_text) < 50:
                st.warning("Provide CV (min 50 chars)")
            elif not transcript or len(transcript) < 100:
                st.warning("Provide transcript (min 100 chars)")
            elif not focus_areas:
                st.warning("Select at least one focus area")
            else:
                # Show immediate spinner feedback
                with st.spinner(f"🔄 Analyzing {candidate_name}... This takes 15-30 seconds."):
                    st.session_state.working_on = f"Evaluating {candidate_name} as {selected_persona_name}..."
                    result, _ = evaluate_candidate(
                        candidate_name, 
                        cv_text, 
                        jd_text, 
                        transcript, 
                        focus_areas,
                        persona=selected_persona,
                        rigor=selected_rigor,
                        key_concerns=key_concerns
                    )
                    st.session_state.working_on = None
                
                if not str(result).startswith("Error"):
                    try:
                        m = re.search(r'```json\s*(.*?)\s*```', result, re.DOTALL)
                        parsed = json.loads(m.group(1) if m else result)
                        parsed['candidate_name'] = candidate_name
                        # Store key concerns in the result for display
                        parsed['key_concerns_input'] = key_concerns
                        st.session_state.candidates.append(parsed)
                        st.session_state.current_candidate = parsed
                        st.session_state.current_step = 'results'
                        
                        # Lock settings after first candidate for fair comparison
                        if st.session_state.locked_settings is None:
                            st.session_state.locked_settings = {
                                'persona': selected_persona,
                                'persona_name': selected_persona_name,
                                'rigor': selected_rigor,
                                'rigor_name': selected_rigor_name,
                                'focus_areas': focus_areas
                            }
                        
                        st.rerun()
                    except Exception as e:
                        st.error(f"Parse error: {e}")
                        st.text(result)
                else:
                    st.error(result)

# ============================================
# TAB 2: GENERATE QUESTIONS
# ============================================
with tab_questions:
    st.markdown("### Generate Interview Questions")
    
    col1, col2 = st.columns(2)
    
    with col1:
        
        st.markdown("**Job Details**")
        job_title = st.text_input("Job Title", placeholder="Senior Software Engineer", key="q_title")
        
        q_jd_src = st.radio("JD Source:", ["📁 Upload", "📜 History", "📝 Paste"], horizontal=True, key="q_jd_src")
        
        requirements = ""
        if q_jd_src == "📁 Upload":
            q_jd_file = st.file_uploader("Upload JD", type=['txt', 'pdf', 'docx'], key="q_jd_file")
            if q_jd_file:
                requirements = extract_text_from_file(q_jd_file)
                if requirements and not requirements.startswith("["):
                    st.success("Loaded")
                else:
                    st.warning(requirements)
        elif q_jd_src == "📜 History":
            history = get_jd_history(20)
            if history:
                opts = {f"{j['job_title']} ({j['created_at'][:10]})": j for j in history}
                sel = st.selectbox("Select:", list(opts.keys()), key="q_jd_hist")
                if sel:
                    requirements = opts[sel].get('generated_jd', '')
                    st.success(f"Loaded: {opts[sel].get('job_title')}")
            else:
                st.info("No saved JDs")
        else:
            requirements = st.text_area("Requirements:", height=150, key="q_requirements")
        
    
    with col2:
        
        st.markdown("**Interview Settings**")
        stage = st.selectbox("Stage", INTERVIEW_STAGES, key="q_stage")
        duration = st.slider("Duration (min)", 15, 90, 45, 5, key="q_duration")
        q_focus = st.multiselect("Focus Areas", FOCUS_AREAS, default=["Technical Skills", "Communication"], key="q_focus")
        exclude_illegal = st.checkbox("Exclude illegal questions", value=True, key="q_exclude")
        
        
        
        st.markdown("**Candidate CV (optional)**")
        q_cv_src = st.radio("Source:", ["None", "📁 Upload", "📝 Paste"], horizontal=True, key="q_cv_src")
        
        q_cv_text = ""
        if q_cv_src == "📁 Upload":
            q_cv_file = st.file_uploader("Upload CV", type=['txt', 'pdf', 'docx'], key="q_cv_file")
            if q_cv_file:
                q_cv_text = extract_text_from_file(q_cv_file)
                if q_cv_text and not q_cv_text.startswith("["):
                    st.success("Loaded")
        elif q_cv_src == "📝 Paste":
            q_cv_text = st.text_area("Paste CV:", height=100, key="q_cv_paste")
        
        
        # New: Question Style and Difficulty Settings
        
        st.markdown("**⚙️ Question Settings**")
        
        # Question Style
        style_options = {v["name"]: k for k, v in QUESTION_STYLES.items()}
        selected_style_name = st.selectbox(
            "Question Style:",
            options=list(style_options.keys()),
            index=0,
            key="q_style",
            help="Adjust the focus and tone of questions"
        )
        selected_style = style_options[selected_style_name]
        st.caption(QUESTION_STYLES[selected_style]["description"])
        
        # Difficulty Level
        difficulty_options = {v["name"]: k for k, v in DIFFICULTY_LEVELS.items()}
        selected_difficulty_name = st.select_slider(
            "Difficulty Level:",
            options=list(difficulty_options.keys()),
            value="⚖️ Mid-Level",
            key="q_difficulty"
        )
        selected_difficulty = difficulty_options[selected_difficulty_name]
        
        # Company Context
        company_context = st.text_area(
            "Company/Team Context (optional):",
            placeholder="e.g., 'Fast-paced startup, remote-first, need someone autonomous'",
            height=68,
            key="q_company_context",
            help="Questions will probe for fit with your specific environment"
        )
        
        
    
    st.markdown("---")
    
    if st.button("❓ Generate Questions", type="primary", use_container_width=True):
        if not job_title:
            st.warning("Enter job title")
        elif not requirements or len(requirements) < 50:
            st.warning("Provide requirements (min 50 chars)")
        elif not q_focus:
            st.warning("Select focus areas")
        else:
            st.session_state.working_on = f"Generating {selected_style_name} questions..."
            result, _ = generate_questions(
                job_title, requirements, stage, duration, q_focus, q_cv_text, exclude_illegal,
                style=selected_style,
                difficulty=selected_difficulty,
                company_context=company_context
            )
            st.session_state.working_on = None
            st.session_state.questions_result = result
    
    if st.session_state.get('questions_result'):
        result = st.session_state.questions_result
        
        if not str(result).startswith("Error"):
            try:
                m = re.search(r'```json\s*(.*?)\s*```', result, re.DOTALL)
                data = json.loads(m.group(1) if m else result)
                
                st.markdown("---")
                st.markdown("### 📋 Interview Guide")
                
                # Show settings used
                settings = data.get('settings', {})
                if settings:
                    st.markdown(f"""
                    <div style="display:flex;gap:8px;margin-bottom:16px;flex-wrap:wrap;">
                        <span style="background:rgba(99,102,241,0.2);color:#a5b4fc;padding:4px 10px;border-radius:12px;font-size:11px;">{settings.get('style', '')}</span>
                        <span style="background:rgba(139,92,246,0.2);color:#c4b5fd;padding:4px 10px;border-radius:12px;font-size:11px;">{settings.get('difficulty', '')}</span>
                        <span style="background:rgba(16,185,129,0.2);color:#6ee7b7;padding:4px 10px;border-radius:12px;font-size:11px;">⏱ {settings.get('duration_minutes', 45)} min</span>
                        <span style="background:rgba(234,179,8,0.2);color:#fcd34d;padding:4px 10px;border-radius:12px;font-size:11px;">📝 {settings.get('total_questions', 'N/A')} questions</span>
                    </div>
                    """, unsafe_allow_html=True)
                
                # Opening
                opening = data.get('opening', {})
                with st.expander("🎬 **Opening**", expanded=True):
                    st.markdown(f"**Icebreaker:** {opening.get('icebreaker', 'N/A')}")
                    if opening.get('icebreaker_purpose'):
                        st.caption(f"💡 Purpose: {opening.get('icebreaker_purpose')}")
                    st.markdown("---")
                    st.markdown(f"**Role Understanding:** {opening.get('role_intro', 'N/A')}")
                    if opening.get('expected_good_answer'):
                        st.success(f"✅ **Good Answer:** {opening.get('expected_good_answer')}")
                
                # JD Fit Questions
                jd_questions = data.get('jd_fit_questions', [])
                if jd_questions:
                    st.markdown("#### 🎯 JD Fit Questions")
                    st.caption("These assess if the candidate meets the job requirements")
                    for i, q in enumerate(jd_questions, 1):
                        with st.expander(f"**Q{i}: {q.get('requirement_being_tested', 'Requirement')[:50]}...**"):
                            st.markdown(f"**❓ Question:** {q.get('question', '')}")
                            st.markdown(f"**⏱ Time:** {q.get('time_estimate', '5')} min")
                            
                            if q.get('follow_ups'):
                                st.markdown("**🔄 Follow-ups:**")
                                for f in q.get('follow_ups', []):
                                    st.markdown(f"  → {f}")
                            
                            st.markdown("---")
                            st.success(f"✅ **Good Answer:** {q.get('expected_good_answer', 'N/A')}")
                            st.error(f"❌ **Poor Answer:** {q.get('expected_poor_answer', 'N/A')}")
                            
                            scoring = q.get('scoring_guide', {})
                            if scoring:
                                st.markdown("**📊 Scoring Guide:**")
                                cols = st.columns(3)
                                with cols[0]:
                                    st.markdown(f"<div style='background:rgba(16,185,129,0.1);padding:10px;border-radius:8px;'><strong style='color:#10b981;'>9-10</strong><br>{scoring.get('excellent', '')}</div>", unsafe_allow_html=True)
                                with cols[1]:
                                    st.markdown(f"<div style='background:rgba(234,179,8,0.1);padding:10px;border-radius:8px;'><strong style='color:#eab308;'>6-8</strong><br>{scoring.get('acceptable', '')}</div>", unsafe_allow_html=True)
                                with cols[2]:
                                    st.markdown(f"<div style='background:rgba(239,68,68,0.1);padding:10px;border-radius:8px;'><strong style='color:#ef4444;'>1-5</strong><br>{scoring.get('poor', '')}</div>", unsafe_allow_html=True)
                
                # CV Verification Questions
                cv_questions = data.get('cv_verification_questions', [])
                if cv_questions:
                    st.markdown("#### 🔍 CV Verification Questions")
                    st.caption("These verify specific claims from the candidate's CV")
                    for i, q in enumerate(cv_questions, 1):
                        with st.expander(f"**Verify: {q.get('cv_claim_being_verified', 'CV Claim')[:50]}...**"):
                            st.markdown(f"**❓ Question:** {q.get('question', '')}")
                            st.success(f"✅ **Truthful Answer:** {q.get('expected_truthful_answer', 'N/A')}")
                            
                            if q.get('red_flags'):
                                st.markdown("**🚩 Red Flags:**")
                                for rf in q.get('red_flags', []):
                                    st.markdown(f"  ⚠️ {rf}")
                            
                            if q.get('follow_up_if_suspicious'):
                                st.warning(f"**🔎 If suspicious, ask:** {q.get('follow_up_if_suspicious')}")
                
                # Behavioral Questions
                behavioral = data.get('behavioral_questions', [])
                if behavioral:
                    st.markdown("#### 💭 Behavioral Questions")
                    st.caption("Past behavior predicts future performance (STAR format)")
                    for i, q in enumerate(behavioral, 1):
                        with st.expander(f"**{q.get('competency', 'Competency')}**"):
                            st.markdown(f"**❓ Question:** {q.get('question', '')}")
                            
                            expected = q.get('expected_good_answer', {})
                            if isinstance(expected, dict):
                                st.markdown("**✅ Expected STAR Answer:**")
                                st.markdown(f"  • **S**ituation: {expected.get('situation', 'N/A')}")
                                st.markdown(f"  • **T**ask: {expected.get('task', 'N/A')}")
                                st.markdown(f"  • **A**ction: {expected.get('action', 'N/A')}")
                                st.markdown(f"  • **R**esult: {expected.get('result', 'N/A')}")
                            
                            if q.get('red_flags'):
                                st.markdown("**🚩 Red Flags:**")
                                for rf in q.get('red_flags', []):
                                    st.markdown(f"  ⚠️ {rf}")
                
                # Fallback for old format questions
                old_questions = data.get('questions', [])
                if old_questions and not jd_questions and not cv_questions:
                    st.markdown("#### Questions")
                    for i, q in enumerate(old_questions, 1):
                        with st.expander(f"**Q{i}: {q.get('category', '')}** ({q.get('time_estimate', '')} min)"):
                            st.markdown(f"**Question:** {q.get('question', '')}")
                            if q.get('follow_ups'):
                                st.markdown("**Follow-ups:**")
                                for f in q.get('follow_ups', []):
                                    st.markdown(f"  → {f}")
                            st.success(f"**Look for:** {q.get('what_to_look_for', '')}")
                            st.error(f"**Red flags:** {q.get('red_flags', '')}")
                
                # Closing
                closing = data.get('closing', {})
                with st.expander("🎬 **Closing**"):
                    st.markdown(f"**Invite Questions:** {closing.get('candidate_questions', 'N/A')}")
                    if closing.get('what_good_candidates_ask'):
                        st.markdown("**Good candidates typically ask:**")
                        for q in closing.get('what_good_candidates_ask', []):
                            st.markdown(f"  💡 {q}")
                    st.markdown(f"**Next Steps:** {closing.get('next_steps', 'N/A')}")
                    st.markdown(f"**Timeline:** {closing.get('timeline', 'N/A')}")
                
                # Evaluation Scorecard
                scorecard = data.get('evaluation_scorecard', [])
                if scorecard:
                    st.markdown("#### 📊 Evaluation Scorecard")
                    for ec in scorecard:
                        must_have = "🔴 MUST HAVE" if ec.get('must_have') else "🟡 Nice to have"
                        st.markdown(f"**{ec.get('criterion')}** ({ec.get('weight')}) - {must_have}")
                        if ec.get('scoring_notes'):
                            st.caption(ec.get('scoring_notes'))
                
                # Recommendation Guide
                rec_guide = data.get('overall_recommendation_guide', {})
                if rec_guide:
                    st.markdown("#### 🎯 Recommendation Guide")
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown("**🟢 STRONG HIRE signals:**")
                        for s in rec_guide.get('strong_hire_signals', []):
                            st.markdown(f"  ✅ {s}")
                        st.markdown("**🟡 HIRE signals:**")
                        for s in rec_guide.get('hire_signals', []):
                            st.markdown(f"  ✓ {s}")
                    with col2:
                        st.markdown("**🔴 NO HIRE signals:**")
                        for s in rec_guide.get('no_hire_signals', []):
                            st.markdown(f"  ❌ {s}")
                        st.markdown("**⚠️ Red Flags:**")
                        for s in rec_guide.get('red_flags_to_watch', []):
                            st.markdown(f"  🚩 {s}")
                
                st.markdown("---")
                st.download_button("📥 Download Interview Guide", json.dumps(data, indent=2), "interview_guide.json", use_container_width=True)
                
            except Exception as e:
                st.error(f"Parse error: {e}")
                st.text(result)
        else:
            st.error(result)

# ============================================
# TAB 3: RECRUITER CALL ANALYZER
# ============================================
with tab_coach:
    st.markdown("### 🎤 Recruiter Call Analyzer")
    st.markdown("Analyze your screening calls to improve candidate qualification and conversion.")
    
    # Initialize session state for this tab
    if 'recruiter_analysis_result' not in st.session_state:
        st.session_state.recruiter_analysis_result = None
    
    # Show results if available
    if st.session_state.recruiter_analysis_result:
        result = st.session_state.recruiter_analysis_result
        
        if not str(result).startswith("Error"):
            try:
                m = re.search(r'```json\s*(.*?)\s*```', result, re.DOTALL)
                a = json.loads(m.group(1) if m else result)
                
                # Back button
                if st.button("← Analyze Another Call", key="back_recruiter"):
                    st.session_state.recruiter_analysis_result = None
                    st.rerun()
                
                st.markdown("---")
                
                # Overall Score
                score = a.get('overall_score', 0)
                color = "#10b981" if score >= 70 else "#eab308" if score >= 50 else "#ef4444"
                call_type_name = {
                    "candidate_screen": "Candidate Screen",
                    "client_intake": "Client Intake",
                    "candidate_close": "Offer/Close Call"
                }.get(a.get('call_type', ''), 'Recruiter Call')
                
                st.markdown(f"""
                <div style="background:linear-gradient(135deg,rgba(99,102,241,0.1),rgba(139,92,246,0.1));border-radius:16px;padding:30px;text-align:center;margin-bottom:24px;">
                    <span style="background:rgba(99,102,241,0.2);color:#a5b4fc;padding:4px 12px;border-radius:12px;font-size:12px;">{call_type_name}</span>
                    <p style="color:#9ca3af;margin:16px 0 0;">YOUR PERFORMANCE</p>
                    <p style="color:{color};font-size:64px;font-weight:bold;margin:10px 0;">{score}<span style="font-size:24px;color:#6b7280;">/100</span></p>
                    <p style="color:#e5e5e5;">{a.get('overall_summary', '')}</p>
                </div>
                """, unsafe_allow_html=True)
                
                # Strengths & Improvements
                c1, c2 = st.columns(2)
                with c1:
                    st.markdown("#### 💪 Your Strengths")
                    for s in a.get('recruiter_strengths', []):
                        st.success(s)
                with c2:
                    st.markdown("#### 🎯 Priority Improvements")
                    for i in a.get('priority_improvements', []):
                        st.warning(i)
                
                # Skills Breakdown
                st.markdown("---")
                st.markdown("## 📊 Skills Breakdown")
                
                for skill in a.get('skills', []):
                    skill_score = skill.get('score', 0)
                    skill_color = "#10b981" if skill_score >= 8 else "#eab308" if skill_score >= 6 else "#ef4444"
                    
                    with st.expander(f"**{skill.get('skill_name')}** — {skill_score}/10 ({skill.get('weight', 0)}% weight)"):
                        if skill.get('what_worked'):
                            st.markdown("**✅ What Worked:**")
                            for w in skill.get('what_worked', []):
                                st.markdown(f"→ {w}")
                        
                        if skill.get('what_needed_improvement'):
                            st.markdown("**❌ What Needed Improvement:**")
                            for w in skill.get('what_needed_improvement', []):
                                st.markdown(f"→ {w}")
                        
                        if skill.get('transcript_examples'):
                            st.markdown("**📍 From the Call:**")
                            for ex in skill.get('transcript_examples', []):
                                st.markdown(f'<div style="background:#1a1a2e;border-left:3px solid #6366f1;padding:12px;margin:8px 0;border-radius:0 8px 8px 0;font-style:italic;color:#a5b4fc;">"{ex}"</div>', unsafe_allow_html=True)
                        
                        if skill.get('fix_for_next_call'):
                            st.markdown(f"""
                            <div style="background:rgba(16,185,129,0.1);border:1px solid rgba(16,185,129,0.3);border-radius:8px;padding:12px;margin-top:12px;">
                                <p style="color:#10b981;font-weight:600;margin:0 0 8px;">⚡ Fix for Next Call:</p>
                                <p style="color:#e5e5e5;margin:0;font-style:italic;">"{skill.get('fix_for_next_call')}"</p>
                            </div>
                            """, unsafe_allow_html=True)
                
                # Key Info Extracted
                key_info = a.get('key_info_extracted', {})
                if key_info:
                    st.markdown("---")
                    st.markdown("## 📋 Key Info Extracted")
                    st.caption("Did you gather what you need to submit this candidate?")
                    
                    info_items = [
                        ("💰 Salary Expectations", key_info.get('salary_expectations', 'Not discussed')),
                        ("📅 Notice Period", key_info.get('notice_period', 'Not discussed')),
                        ("🎯 Motivation", key_info.get('motivation', 'Not discussed')),
                        ("⚔️ Competing Offers", key_info.get('competing_offers', 'Not discussed'))
                    ]
                    
                    cols = st.columns(2)
                    for idx, (label, value) in enumerate(info_items):
                        with cols[idx % 2]:
                            is_missed = 'MISSED' in str(value).upper() or 'NOT DISCUSSED' in str(value).upper()
                            bg_color = "rgba(239,68,68,0.1)" if is_missed else "rgba(16,185,129,0.1)"
                            border_color = "#ef4444" if is_missed else "#10b981"
                            st.markdown(f"""
                            <div style="background:{bg_color};border-left:3px solid {border_color};padding:12px;margin:8px 0;border-radius:0 8px 8px 0;">
                                <p style="color:#9ca3af;font-size:12px;margin:0;">{label}</p>
                                <p style="color:#e5e5e5;margin:4px 0 0;font-weight:500;">{value}</p>
                            </div>
                            """, unsafe_allow_html=True)
                    
                    # Red flags identified
                    red_flags = key_info.get('red_flags_identified', [])
                    if red_flags:
                        st.markdown("**🚩 Red Flags Identified:**")
                        for flag in red_flags:
                            st.markdown(f"- {flag}")
                
                # JD Fit Assessment (NEW)
                jd_fit = a.get('jd_fit_assessment', {})
                if jd_fit:
                    st.markdown("---")
                    st.markdown("## 🎯 Candidate-Role Fit Assessment")
                    
                    overall_fit = jd_fit.get('overall_fit', 'Unknown')
                    fit_color = "#10b981" if "Strong" in overall_fit else "#eab308" if "Potential" in overall_fit else "#ef4444"
                    
                    st.markdown(f"""
                    <div style="background:linear-gradient(135deg,rgba(99,102,241,0.1),rgba(139,92,246,0.1));border-radius:12px;padding:20px;margin-bottom:16px;">
                        <p style="color:#9ca3af;font-size:12px;margin:0;">OVERALL FIT</p>
                        <p style="color:{fit_color};font-size:24px;font-weight:bold;margin:8px 0;">{overall_fit}</p>
                        <p style="color:#e5e5e5;margin:0;">{jd_fit.get('fit_summary', '')}</p>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Requirements verified
                    reqs_verified = jd_fit.get('requirements_verified', [])
                    if reqs_verified:
                        st.markdown("**📋 Requirements Discussed:**")
                        for req in reqs_verified:
                            status = req.get('status', 'Unknown')
                            status_icon = "✅" if status == "Met" else "❌" if status == "Not Met" else "❓"
                            st.markdown(f"{status_icon} **{req.get('requirement', '')}**: {status}")
                            if req.get('evidence'):
                                st.caption(f"   Evidence: \"{req.get('evidence')}\"")
                    
                    # Requirements not discussed
                    not_discussed = jd_fit.get('requirements_not_discussed', [])
                    if not_discussed:
                        st.warning("**⚠️ Requirements NOT Discussed:**")
                        for req in not_discussed:
                            st.markdown(f"- {req}")
                    
                    # Missed opportunities
                    missed = jd_fit.get('recruiter_missed_opportunities', [])
                    if missed:
                        st.markdown("**💡 Questions You Should Have Asked:**")
                        for m in missed:
                            st.markdown(f"- {m}")
                
                # Concerns Addressed (NEW)
                concerns = a.get('concerns_addressed', [])
                if concerns:
                    st.markdown("---")
                    st.markdown("## 🔍 Your Concerns Addressed")
                    for concern in concerns:
                        st.markdown(f"""
                        <div style="background:#1a1a2e;border:1px solid rgba(99,102,241,0.3);border-radius:12px;padding:16px;margin:12px 0;">
                            <p style="color:#a5b4fc;font-weight:600;margin:0 0 8px;">❓ {concern.get('concern', '')}</p>
                            <p style="color:#e5e5e5;margin:0 0 8px;">→ {concern.get('finding', '')}</p>
                            {f"<p style='color:#9ca3af;font-style:italic;margin:0;font-size:13px;'>Evidence: \"{concern.get('evidence', '')}\"</p>" if concern.get('evidence') else ""}
                        </div>
                        """, unsafe_allow_html=True)
                
                # Client Readiness
                client_ready = a.get('client_readiness', {})
                if client_ready:
                    st.markdown("---")
                    st.markdown("## 📤 Client Submission Readiness")
                    
                    ready = client_ready.get('ready_to_submit', False)
                    if ready:
                        st.success("✅ **Ready to Submit** - You have the key info needed")
                    else:
                        st.warning("⚠️ **Not Ready** - Missing critical information")
                        missing = client_ready.get('missing_info', [])
                        if missing:
                            st.markdown("**Missing Info:**")
                            for m in missing:
                                st.markdown(f"- ❌ {m}")
                    
                    notes = client_ready.get('submission_notes', '')
                    if notes:
                        st.markdown(f"""
                        <div style="background:rgba(99,102,241,0.1);border:1px solid rgba(99,102,241,0.3);border-radius:8px;padding:16px;margin-top:12px;">
                            <p style="color:#a5b4fc;font-weight:600;margin:0 0 8px;">📝 Submission Notes:</p>
                            <p style="color:#e5e5e5;margin:0;">{notes}</p>
                        </div>
                        """, unsafe_allow_html=True)
                
                # Coaching Summary
                st.markdown("---")
                st.markdown("## 🎓 Coaching Summary")
                st.markdown(f"""
                <div style="background:linear-gradient(135deg,rgba(99,102,241,0.1),rgba(139,92,246,0.1));border-radius:12px;padding:24px;border-left:4px solid #6366f1;">
                    {a.get('coaching_summary', '')}
                </div>
                """, unsafe_allow_html=True)
                
                # Export
                st.markdown("---")
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button("📥 Download Analysis", result, "recruiter_analysis.json", use_container_width=True)
                with col2:
                    if st.button("🔄 Analyze Another Call", use_container_width=True, key="another_recruiter"):
                        st.session_state.recruiter_analysis_result = None
                        st.rerun()
            
            except Exception as e:
                st.error(f"Parse error: {e}")
                st.text(result)
        else:
            st.error(result)
    
    else:
        # Input Form
        col1, col2 = st.columns(2)
        
        with col1:
            
            st.markdown("**Call Details**")
            
            call_type = st.selectbox(
                "Call Type:",
                options=["candidate_screen", "client_intake", "candidate_close"],
                format_func=lambda x: {
                    "candidate_screen": "📞 Candidate Screening Call",
                    "client_intake": "📋 Client Job Intake Call",
                    "candidate_close": "🎯 Candidate Offer/Close Call"
                }.get(x, x),
                key="recruiter_call_type"
            )
            
            # Show description based on call type
            call_descriptions = {
                "candidate_screen": "Analyze how well you qualified the candidate and sold the opportunity",
                "client_intake": "Analyze how well you took the job order and positioned your services",
                "candidate_close": "Analyze how well you presented the offer and handled objections"
            }
            st.caption(call_descriptions.get(call_type, ""))
            
            candidate_role = st.text_input(
                "Role Being Recruited:",
                placeholder="e.g., Senior Software Engineer, Sales Manager",
                key="recruiter_role"
            )
            
            company_context = st.text_input(
                "Company/Client (optional):",
                placeholder="e.g., Tech startup, Fortune 500",
                key="recruiter_company"
            )
            
            key_concerns = st.text_area(
                "Specific areas to focus on (optional):",
                placeholder="e.g., 'Did I handle the salary discussion well?' or 'Watch for red flags I missed'",
                height=68,
                key="recruiter_concerns"
            )
            
            
            # JD Upload for fit analysis
            
            st.markdown("**📋 Job Description (optional)**")
            st.caption("Upload JD to analyze if candidate fits the role requirements")
            
            rc_jd_src = st.radio("JD Source:", ["None", "📁 Upload", "📝 Paste"], horizontal=True, key="rc_jd_src")
            
            rc_jd_text = ""
            if rc_jd_src == "📁 Upload":
                rc_jd_file = st.file_uploader("Upload JD", type=['txt', 'pdf', 'docx'], key="rc_jd_file")
                if rc_jd_file:
                    rc_jd_text = extract_text_from_file(rc_jd_file)
                    if rc_jd_text and not rc_jd_text.startswith("["):
                        st.success("✅ JD Loaded - Will analyze candidate fit")
            elif rc_jd_src == "📝 Paste":
                rc_jd_text = st.text_area("Paste JD:", height=100, key="rc_jd_paste", placeholder="Paste job requirements here...")
                if rc_jd_text:
                    st.success("✅ JD provided - Will analyze candidate fit")
            
            
        
        with col2:
            
            st.markdown("**Call Transcript**")
            
            input_method = st.radio(
                "Input Method:",
                ["📝 Paste Transcript", "📁 Upload File"],
                horizontal=True,
                key="recruiter_input_method"
            )
            
            recruiter_transcript = ""
            
            if input_method == "📝 Paste Transcript":
                recruiter_transcript = st.text_area(
                    "Paste your transcript:",
                    height=300,
                    placeholder="Recruiter: Hi Sarah, thanks for taking the time to chat today...\nCandidate: Of course, I'm excited to learn more...",
                    key="recruiter_transcript_paste"
                )
            else:
                uploaded = st.file_uploader(
                    "Upload transcript",
                    type=['txt', 'pdf', 'docx', 'vtt', 'srt'],
                    key="recruiter_transcript_upload"
                )
                if uploaded:
                    with st.spinner("Processing..."):
                        recruiter_transcript = extract_text_from_file(uploaded)
                    if recruiter_transcript and not recruiter_transcript.startswith("["):
                        st.success(f"✅ Loaded: {len(recruiter_transcript):,} characters")
            
            
        
        st.markdown("---")
        
        if st.button("🎤 Analyze My Performance", type="primary", use_container_width=True):
            if not recruiter_transcript or len(recruiter_transcript.strip()) < 100:
                st.warning("Please provide a transcript (minimum 100 characters)")
            else:
                with st.spinner("🔄 Analyzing your recruiter performance... 20-40 seconds"):
                    result, _ = analyze_recruiter_screen(
                        transcript=recruiter_transcript,
                        call_type=call_type,
                        candidate_role=candidate_role,
                        company_context=company_context,
                        key_concerns=key_concerns,
                        jd_text=rc_jd_text
                    )
                
                if not str(result).startswith("Error"):
                    st.session_state.recruiter_analysis_result = result
                    st.rerun()
                else:
                    st.error(result)

# Feedback Widget
if USING_SHARED:
    render_feedback_widget("interview", submit_feedback)
else:
    st.markdown('<div style="height:60px;"></div>', unsafe_allow_html=True)
    _, _, _, fb = st.columns([4, 1, 1, 1])
    with fb:
        with st.popover("💬 Feedback"):
            st.markdown("**Send Feedback**")
            ft = st.segmented_control("Type", ["🐛 Bug", "✨ Feature", "💬 General"], default="💬 General", label_visibility="collapsed")
            fm = st.text_area("Message", height=100, placeholder="...", label_visibility="collapsed", key="fb_msg")
            if st.button("Send", type="primary", use_container_width=True, key="fb_send"):
                if fm:
                    fb_type = ft.split()[1].lower() if ft else "general"
                    if submit_feedback("interview", fb_type, fm):
                        st.success("Thanks!")
                    else:
                        st.error("Failed")

"""
Sharp [AppName] - [Description]
================================
TEMPLATE: Copy this when creating new apps or updating existing ones.

STANDARDS ENFORCED:
1. Uses shared_config.py for all configuration
2. Uses shared_ui.py for consistent styling
3. Custom HTML <details> instead of st.expander (fixes Material Icons bug)
4. Proper export functions with error handling
5. StatusIndicator for showing AI progress
6. Top banner + sidebar from shared_ui
"""
import streamlit as st
import os
import requests
from datetime import datetime, timedelta
import secrets
import json
import re
import io
import sys

# ============================================
# SHARED MODULE IMPORTS (REQUIRED)
# ============================================
# Add parent directory for shared modules
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..'))

try:
    from shared_config import (
        SUPABASE_URL, SUPABASE_ANON_KEY, SUPABASE_SERVICE_KEY,
        ANTHROPIC_API_KEY, GOD_PASSWORD, APP_URLS, CLAUDE_MODEL
    )
    from shared_ui import (
        apply_global_styles,
        render_top_banner,
        render_promo_banner,
        render_header,
        render_sidebar,
        render_feedback_widget,
        StatusIndicator,
        COLORS, TYPOGRAPHY, EXTERNAL_LINKS
    )
    USING_SHARED = True
except ImportError:
    # Fallback if shared modules not available (for local testing)
    USING_SHARED = False
    SUPABASE_URL = "https://qkjtprqgblnfftrotyks.supabase.co"
    SUPABASE_ANON_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InFranRwcnFnYmxuZmZ0cm90eWtzIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NjUzNTgzNDAsImV4cCI6MjA4MDkzNDM0MH0.pVzSq4M5i58zBGl7OPDhNL9qYBcg-bz8MVrBI5MQSkw"
    SUPABASE_SERVICE_KEY = os.environ.get("SUPABASE_SERVICE_KEY", SUPABASE_ANON_KEY)
    ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
    GOD_PASSWORD = os.environ.get("GOD_PASSWORD", "G0DHum@n101!!!")
    CLAUDE_MODEL = "claude-sonnet-4-20250514"
    APP_URLS = {
        "portal": "https://demo.sharphuman.com", "jd": "https://jd.sharphuman.com",
        "screen": "https://screen.sharphuman.com", "interview": "https://hire.sharphuman.com",
        "source": "https://outreach.sharphuman.com", "content": "https://content.sharphuman.com",
        "sales": "https://sales.sharphuman.com", "reach": "https://reach.sharphuman.com",
        "assistant": "https://assistant.sharphuman.com", "admin": "https://admin.sharphuman.com",
    }
    # Fallback colors for when shared_ui not available
    COLORS = {
        "primary": "#6366f1", "secondary": "#8b5cf6", "bg_dark": "#0a0a0f",
        "bg_card": "#12121a", "text_primary": "#ffffff", "text_secondary": "#e5e5e5",
        "text_muted": "#9ca3af", "border": "rgba(99, 102, 241, 0.2)",
        "success": "#10b981", "warning": "#eab308", "error": "#ef4444"
    }

# ============================================
# APP CONFIGURATION (CHANGE THESE)
# ============================================
APP_KEY = "template"  # Change: portal, jd, screen, interview, source, content, sales, reach, assistant
APP_NAME = "Sharp Template"
APP_SUBTITLE = "Template App Description"
APP_ICON = "🚀"

# ============================================
# STANDARD AUTH FUNCTIONS (DO NOT MODIFY)
# ============================================

def create_session(user_id, email):
    """Create a session in Supabase and return the token."""
    token = secrets.token_urlsafe(32)
    try:
        requests.post(f"{SUPABASE_URL}/rest/v1/sessions",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                     "Content-Type": "application/json", "Prefer": "return=minimal"},
            json={"user_id": user_id, "token": token, "ip_address": "unknown", "device_hash": APP_KEY,
                  "is_active": True, "expires_at": (datetime.utcnow() + timedelta(hours=24)).isoformat()}, timeout=10)
    except:
        pass
    return token

def supabase_sign_in(email, password):
    try:
        r = requests.post(f"{SUPABASE_URL}/auth/v1/token?grant_type=password",
                          headers={"apikey": SUPABASE_ANON_KEY, "Content-Type": "application/json"},
                          json={"email": email, "password": password}, timeout=10)
        data = r.json()
        if r.status_code == 200 and data.get("access_token"):
            user = data.get("user", {})
            session_token = create_session(user.get("id"), email)
            return {"success": True, "user": user, "session_token": session_token}
        return {"success": False, "message": data.get("error_description") or "Invalid credentials"}
    except Exception as e:
        return {"success": False, "message": str(e)}

def supabase_sign_up(email, password):
    try:
        r = requests.post(f"{SUPABASE_URL}/auth/v1/signup",
                          headers={"apikey": SUPABASE_ANON_KEY, "Content-Type": "application/json"},
                          json={"email": email, "password": password}, timeout=10)
        data = r.json()
        if r.status_code == 200 and data.get("user"):
            return {"success": True, "message": "Check your email to confirm!"}
        return {"success": False, "message": data.get("error_description") or "Sign up failed"}
    except Exception as e:
        return {"success": False, "message": str(e)}

def supabase_magic_link(email):
    try:
        r = requests.post(f"{SUPABASE_URL}/auth/v1/magiclink",
                          headers={"apikey": SUPABASE_ANON_KEY, "Content-Type": "application/json"},
                          json={"email": email}, timeout=10)
        return {"success": r.status_code == 200, "message": "Magic link sent!" if r.status_code == 200 else "Failed"}
    except Exception as e:
        return {"success": False, "message": str(e)}

def validate_session_token(token):
    """Validate a session token from URL and return user info."""
    if not token:
        return None
    try:
        r = requests.get(f"{SUPABASE_URL}/rest/v1/sessions?token=eq.{token}&is_active=eq.true",
                         headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}"}, timeout=10)
        if r.status_code == 200 and r.json():
            session = r.json()[0]
            expires = datetime.fromisoformat(session["expires_at"].replace("Z", "+00:00"))
            if expires.replace(tzinfo=None) > datetime.utcnow():
                user_r = requests.get(f"{SUPABASE_URL}/rest/v1/user_profiles?id=eq.{session['user_id']}",
                                      headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}"}, timeout=10)
                if user_r.status_code == 200 and user_r.json():
                    profile = user_r.json()[0]
                    return {"user_id": session["user_id"], "email": profile.get("email"),
                            "plan": profile.get("plan", "free"), "token": token}
    except:
        pass
    return None

def log_usage(user_id, session_token, app, action, tokens_used=0):
    """Log usage to Supabase."""
    try:
        requests.post(f"{SUPABASE_URL}/rest/v1/usage_logs",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                     "Content-Type": "application/json", "Prefer": "return=minimal"},
            json={"user_id": user_id, "session_id": session_token, "app": app, "action": action, "tokens_used": tokens_used}, timeout=5)
    except:
        pass

def submit_feedback(app, feedback_type, message):
    """Submit feedback to Supabase."""
    try:
        requests.post(f"{SUPABASE_URL}/rest/v1/user_feedback",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                     "Content-Type": "application/json", "Prefer": "return=minimal"},
            json={"user_id": st.session_state.user.get("id") if st.session_state.user else None,
                  "app": app, "feedback_type": feedback_type, "rating": 4, "message": message,
                  "email": get_user_email()}, timeout=5)
        return True
    except:
        return False

def init_session():
    """Initialize session state with defaults."""
    defaults = {
        'authenticated': False, 'user': None, 'is_god': False, 
        'session_token': None, 'user_plan': 'free',
        'promo_dismissed': False,
        'working_on': None  # For status indicator
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

def check_url_auth():
    """Check for auth token in URL params."""
    token = st.query_params.get("auth")
    if token and not st.session_state.authenticated:
        user_info = validate_session_token(token)
        if user_info:
            st.session_state.authenticated = True
            st.session_state.user = {"email": user_info["email"], "id": user_info["user_id"]}
            st.session_state.session_token = token
            st.session_state.user_plan = user_info.get("plan", "free")
            st.session_state.is_god = user_info.get("plan") == "god"
            return True
    return False

def get_user_email():
    """Get current user's email."""
    if st.session_state.user:
        return st.session_state.user.get("email", "User")
    return "GOD" if st.session_state.is_god else "User"

def build_app_url(app_name):
    """Build URL with auth token for cross-app navigation."""
    base_url = APP_URLS.get(app_name, "")
    token = st.session_state.get("session_token", "")
    return f"{base_url}?auth={token}" if base_url and token else base_url

# ============================================
# STANDARD CLAUDE API
# ============================================

def call_claude(prompt, max_tokens=4000, action="generate"):
    """Call Claude API with usage logging."""
    if not ANTHROPIC_API_KEY:
        return "Error: ANTHROPIC_API_KEY not set", 0
    try:
        r = requests.post("https://api.anthropic.com/v1/messages",
            headers={"x-api-key": ANTHROPIC_API_KEY, "Content-Type": "application/json", "anthropic-version": "2023-06-01"},
            json={"model": CLAUDE_MODEL, "max_tokens": max_tokens, "messages": [{"role": "user", "content": prompt}]}, timeout=180)
        if r.status_code == 200:
            text = r.json()["content"][0]["text"]
            tokens = (len(prompt) + len(text)) // 4
            if st.session_state.user:
                log_usage(st.session_state.user.get("id"), st.session_state.get("session_token"), APP_KEY, action, tokens)
            return text, tokens
        return f"Error: {r.status_code}", 0
    except Exception as e:
        return f"Error: {e}", 0

# ============================================
# STANDARD TEXT HELPERS
# ============================================

def clean_text(text):
    """Clean text for display/export - removes unprintable characters."""
    if not text:
        return ""
    return ''.join(c if c.isprintable() or c in '\n\r\t' else ' ' for c in str(text))

def is_readable_text(text):
    """Check if extracted text is actually readable (not garbled)."""
    if not text or len(text.strip()) < 10:
        return False
    alpha_chars = sum(1 for c in text if c.isalpha())
    alpha_ratio = alpha_chars / len(text) if len(text) > 0 else 0
    return alpha_ratio > 0.3

# ============================================
# STANDARD FILE EXTRACTION (PDF/DOCX/TXT)
# ============================================

def extract_text_from_file(uploaded_file):
    """Extract text from uploaded file with multi-layer fallbacks."""
    if uploaded_file is None:
        return ""
    
    file_type = uploaded_file.name.split('.')[-1].lower()
    content = uploaded_file.read()
    uploaded_file.seek(0)
    
    if file_type == 'txt':
        return clean_text(content.decode('utf-8', errors='ignore'))
    
    elif file_type == 'pdf':
        extracted = ""
        
        # Layer 1: PyMuPDF
        try:
            import fitz
            pdf = fitz.open(stream=content, filetype="pdf")
            parts = []
            for page in pdf:
                text = page.get_text("text")
                if not text.strip():
                    text = page.get_text("text", flags=fitz.TEXT_PRESERVE_WHITESPACE)
                if not text.strip():
                    blocks = page.get_text("blocks")
                    text = "\n".join([b[4] for b in blocks if b[6] == 0])
                if text.strip():
                    parts.append(text)
            pdf.close()
            extracted = "\n".join(parts)
            if extracted.strip() and is_readable_text(extracted):
                return clean_text(extracted)
        except:
            pass
        
        # Layer 2: pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                parts = [page.extract_text() or "" for page in pdf.pages]
                extracted = "\n".join(parts)
                if extracted.strip() and is_readable_text(extracted):
                    return clean_text(extracted)
        except:
            pass
        
        # Layer 3: PyPDF2
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(content))
            parts = [page.extract_text() or "" for page in reader.pages]
            extracted = "\n".join(parts)
            if extracted.strip() and is_readable_text(extracted):
                return clean_text(extracted)
        except:
            pass
        
        return "[PDF extraction failed - the PDF may be scanned/image-based. Please paste the content directly.]"
    
    elif file_type in ['docx', 'doc']:
        # Try python-docx
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text.strip())
            if parts:
                return clean_text('\n\n'.join(parts))
        except:
            pass
        
        # Fallback: raw XML
        try:
            import zipfile
            from xml.etree import ElementTree
            with zipfile.ZipFile(io.BytesIO(content)) as z:
                if 'word/document.xml' in z.namelist():
                    xml_content = z.read('word/document.xml')
                    tree = ElementTree.fromstring(xml_content)
                    texts = [elem.text.strip() for elem in tree.iter() if elem.text and elem.text.strip()]
                    if texts:
                        return clean_text(' '.join(texts))
        except:
            pass
        
        return "[DOCX extraction failed - please paste the content directly]"
    
    return clean_text(content.decode('utf-8', errors='ignore'))

# ============================================
# STANDARD EXPORT FUNCTIONS
# ============================================

def export_to_markdown(data, title="Export"):
    """Export data to Markdown format."""
    if isinstance(data, dict):
        lines = [f"# {title}\n"]
        for key, value in data.items():
            if isinstance(value, list):
                lines.append(f"\n## {key.replace('_', ' ').title()}\n")
                for item in value:
                    lines.append(f"- {clean_text(str(item))}")
            elif isinstance(value, dict):
                lines.append(f"\n## {key.replace('_', ' ').title()}\n")
                for k, v in value.items():
                    lines.append(f"**{k}:** {clean_text(str(v))}")
            else:
                lines.append(f"**{key.replace('_', ' ').title()}:** {clean_text(str(value))}")
        return "\n".join(lines)
    return str(data)

def export_to_docx(data, title="Export"):
    """Export data to DOCX format with error handling."""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(title, 0)
        
        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, list):
                    doc.add_heading(key.replace('_', ' ').title(), level=1)
                    for item in value:
                        doc.add_paragraph(clean_text(str(item)), style='List Bullet')
                elif isinstance(value, dict):
                    doc.add_heading(key.replace('_', ' ').title(), level=1)
                    for k, v in value.items():
                        doc.add_paragraph(f"{k}: {clean_text(str(v))}")
                else:
                    doc.add_paragraph(f"{key.replace('_', ' ').title()}: {clean_text(str(value))}")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        st.warning("📦 python-docx not installed. DOCX export unavailable.")
        return None
    except Exception as e:
        st.error(f"DOCX export error: {e}")
        return None

def export_to_pdf(data, title="Export"):
    """Export data to PDF format with error handling."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.colors import HexColor
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=20, textColor=HexColor('#6366f1'))
        heading_style = ParagraphStyle('Heading', parent=styles['Heading1'], fontSize=14, textColor=HexColor('#374151'))
        body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10, textColor=HexColor('#4b5563'))
        
        story = [Paragraph(title, title_style), Spacer(1, 12)]
        
        if isinstance(data, dict):
            for key, value in data.items():
                story.append(Paragraph(key.replace('_', ' ').title(), heading_style))
                if isinstance(value, list):
                    for item in value:
                        story.append(Paragraph(f"• {clean_text(str(item))}", body_style))
                elif isinstance(value, dict):
                    for k, v in value.items():
                        story.append(Paragraph(f"<b>{k}:</b> {clean_text(str(v))}", body_style))
                else:
                    story.append(Paragraph(clean_text(str(value)), body_style))
                story.append(Spacer(1, 8))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        st.warning("📦 reportlab not installed. PDF export unavailable.")
        return None
    except Exception as e:
        st.error(f"PDF export error: {e}")
        return None

# ============================================
# STANDARD UI HELPERS
# ============================================

def render_accordion(title, content, score=None, score_color=None, expanded=False):
    """
    Render a custom accordion (replaces st.expander to avoid Material Icons bug).
    
    Args:
        title: Accordion header text
        content: HTML content inside
        score: Optional score to display (e.g., "8/10")
        score_color: Color for score badge
        expanded: Whether to start expanded
    """
    open_attr = "open" if expanded else ""
    score_html = ""
    if score:
        color = score_color or COLORS.get("primary", "#6366f1")
        score_html = f'<span style="background:{color};color:white;padding:4px 12px;border-radius:12px;font-weight:600;font-size:13px;">{score}</span>'
    
    st.markdown(f"""
    <details {open_attr} style="background:{COLORS.get('bg_card', '#12121a')};border:1px solid {COLORS.get('border', 'rgba(99,102,241,0.2)')};border-radius:10px;margin:8px 0;">
        <summary style="padding:14px 18px;cursor:pointer;display:flex;align-items:center;gap:12px;list-style:none;">
            <span style="color:{COLORS.get('text_primary', '#fff')};font-weight:600;flex:1;">{title}</span>
            {score_html}
        </summary>
        <div style="padding:0 18px 18px;border-top:1px solid {COLORS.get('border', 'rgba(99,102,241,0.1)')};">
            {content}
        </div>
    </details>
    """, unsafe_allow_html=True)

def render_download_buttons(data, filename_base, formats=["md", "docx", "pdf", "json"]):
    """
    Render a row of download buttons with proper error handling.
    
    Args:
        data: Dictionary of data to export
        filename_base: Base filename without extension
        formats: List of formats to offer
    """
    cols = st.columns(len(formats))
    
    for i, fmt in enumerate(formats):
        with cols[i]:
            if fmt == "md":
                content = export_to_markdown(data, filename_base)
                st.download_button(f"📥 MD", content, f"{filename_base}.md", use_container_width=True)
            elif fmt == "docx":
                content = export_to_docx(data, filename_base)
                if content:
                    st.download_button(f"📥 DOCX", content, f"{filename_base}.docx", 
                                      mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                      use_container_width=True)
                else:
                    st.button("📥 DOCX", disabled=True, use_container_width=True, help="Export failed")
            elif fmt == "pdf":
                content = export_to_pdf(data, filename_base)
                if content:
                    st.download_button(f"📥 PDF", content, f"{filename_base}.pdf",
                                      mime="application/pdf", use_container_width=True)
                else:
                    st.button("📥 PDF", disabled=True, use_container_width=True, help="Export failed")
            elif fmt == "json":
                st.download_button(f"📥 JSON", json.dumps(data, indent=2), f"{filename_base}.json",
                                  use_container_width=True)

# ============================================
# STANDARD AUTH UI
# ============================================

def render_auth():
    """Render login/signup page with consistent styling."""
    if USING_SHARED:
        apply_global_styles()
    else:
        st.markdown("""<style>
        @import url('https://fonts.googleapis.com/css2?family=Nunito:wght@400;600;700&display=swap');
        .stApp { background: linear-gradient(135deg, #0a0a0f, #0f0f1a); }
        * { font-family: 'Nunito', sans-serif !important; }
        h1,h2,h3 { color: white !important; }
        p, label { color: #e5e5e5 !important; }
        .stTextInput>div>div>input { background: #12121a !important; border: 1px solid rgba(99,102,241,0.3) !important; color: white !important; }
        .stButton>button { background: linear-gradient(135deg, #6366f1, #8b5cf6) !important; color: white !important; border: none !important; }
        .stTabs [data-baseweb="tab-list"] { background: #12121a; }
        .stTabs [aria-selected="true"] { background: rgba(99,102,241,0.3) !important; color: white !important; }
        </style>""", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 1.3, 1])
    with col2:
        st.markdown(f"""
        <div style="text-align:center;padding:50px 0 40px;">
            <img src="https://sharphuman.com/logo1-3.png" style="width:80px;margin-bottom:20px;">
            <h1 style="color:white;font-size:32px;margin:0 0 8px;">{APP_NAME}</h1>
            <p style="color:#9ca3af;font-size:15px;">{APP_SUBTITLE}</p>
        </div>
        """, unsafe_allow_html=True)
        
        tab1, tab2 = st.tabs(["🔐 Log In", "✨ Sign Up"])
        
        with tab1:
            email = st.text_input("Email", key="login_email")
            password = st.text_input("Password", type="password", key="login_password")
            
            if st.button("Log In", use_container_width=True):
                if password == GOD_PASSWORD:
                    token = secrets.token_urlsafe(32)
                    st.session_state.authenticated = True
                    st.session_state.is_god = True
                    st.session_state.user = {"email": "GOD", "id": "god"}
                    st.session_state.user_plan = "god"
                    st.session_state.session_token = token
                    log_usage("god", token, APP_KEY, "login_god")
                    st.rerun()
                elif email and password:
                    result = supabase_sign_in(email, password)
                    if result["success"]:
                        st.session_state.authenticated = True
                        st.session_state.user = result["user"]
                        st.session_state.session_token = result.get("session_token")
                        log_usage(result["user"].get("id"), result.get("session_token"), APP_KEY, "login")
                        st.rerun()
                    else:
                        st.error(result["message"])
            
            st.markdown("<p style='text-align:center;color:#6b7280;margin:16px 0;'>— or —</p>", unsafe_allow_html=True)
            
            magic_email = st.text_input("Email for magic link", key="magic_email", label_visibility="collapsed", placeholder="Email for magic link")
            if st.button("✨ Send Magic Link", use_container_width=True, key="magic_btn"):
                if magic_email:
                    result = supabase_magic_link(magic_email)
                    st.success(result["message"]) if result["success"] else st.error(result["message"])
        
        with tab2:
            signup_email = st.text_input("Email", key="signup_email")
            signup_password = st.text_input("Password", type="password", key="signup_password")
            signup_confirm = st.text_input("Confirm Password", type="password", key="signup_confirm")
            
            if st.button("Create Account", use_container_width=True):
                if signup_password != signup_confirm:
                    st.error("Passwords don't match")
                elif len(signup_password) < 6:
                    st.warning("Password must be at least 6 characters")
                elif signup_email and signup_password:
                    result = supabase_sign_up(signup_email, signup_password)
                    st.success(result["message"]) if result["success"] else st.error(result["message"])

# ============================================
# MAIN APP LAYOUT
# ============================================

def render_main_app():
    """Render the main app content after authentication."""
    
    # 1. Apply global styles (REQUIRED)
    if USING_SHARED:
        apply_global_styles()
    else:
        # Fallback styles when shared_ui not available
        st.markdown("""<style>
        @import url('https://fonts.googleapis.com/css2?family=Nunito:wght@400;600;700&display=swap');
        * { font-family: 'Nunito', sans-serif !important; }
        .stApp { background: #0a0a0f !important; }
        h1,h2,h3 { color: #fff !important; }
        p,span,label { color: #e5e5e5 !important; }
        .stButton>button { background: linear-gradient(135deg, #6366f1, #8b5cf6) !important; color: white !important; border: none !important; }
        </style>""", unsafe_allow_html=True)
    
    # 2. Top banner (links to website, blog, calendly)
    if USING_SHARED:
        render_top_banner(show_cta=True, cta_text="Book a Demo")
    
    # 3. Sidebar (navigation, user info, external links)
    if USING_SHARED:
        render_sidebar(
            current_app=APP_KEY,
            user_email=get_user_email(),
            user_plan=st.session_state.get('user_plan', 'free'),
            session_token=st.session_state.get('session_token', '')
        )
    else:
        # Fallback sidebar
        with st.sidebar:
            st.markdown(f"**{get_user_email()}**")
            st.markdown(f"*{st.session_state.get('user_plan', 'free')} plan*")
            st.markdown("---")
            for key, label in [("portal", "🏠 Portal"), ("jd", "📝 JD"), ("screen", "🔍 Screen"), ("interview", "🎯 Interview")]:
                if key == APP_KEY:
                    st.button(f"{label} ◀", disabled=True, use_container_width=True)
                else:
                    st.link_button(label, build_app_url(key), use_container_width=True)
            st.markdown("---")
            if st.button("🚪 Logout", use_container_width=True):
                for k in list(st.session_state.keys()):
                    del st.session_state[k]
                st.rerun()
    
    # 4. Optional promo banner (can be dismissed)
    if USING_SHARED and not st.session_state.get('promo_dismissed'):
        render_promo_banner(
            message="Need custom features or bespoke implementation?",
            cta_text="Let's Chat"
        )
    
    # 5. Header
    if USING_SHARED:
        render_header(APP_NAME, APP_SUBTITLE, APP_ICON)
    else:
        st.markdown(f"""
        <div style="display:flex;align-items:center;gap:16px;padding:20px 0;border-bottom:1px solid rgba(99,102,241,0.2);margin-bottom:30px;">
            <img src="https://sharphuman.com/logo1-3.png" style="width:50px;">
            <div><h1 style="margin:0;font-size:28px;">{APP_NAME}</h1><p style="color:#9ca3af;margin:0;">{APP_SUBTITLE}</p></div>
        </div>
        """, unsafe_allow_html=True)
    
    # ==========================================
    # YOUR APP CONTENT GOES HERE
    # ==========================================
    
    st.markdown("### Your App Content")
    st.info("Replace this section with your app's functionality")
    
    # Example: Using StatusIndicator for long-running AI tasks
    if st.button("🚀 Run Example Task", type="primary"):
        if USING_SHARED:
            status = StatusIndicator(preset="general")  # Options: screening, interview, jd, sales, content, general
            status.start()
            
            import time
            time.sleep(1)
            status.next()  # Move to next step
            time.sleep(1)
            status.next()
            time.sleep(1)
            
            status.finish("✅ Task completed!")
        else:
            with st.spinner("Working..."):
                import time
                time.sleep(3)
            st.success("Done!")
    
    # Example: Using custom accordion (replaces st.expander)
    st.markdown("### Example Accordion")
    render_accordion(
        title="Click to expand",
        content="<p style='color:#e5e5e5;margin:12px 0;'>This is content inside the accordion.</p>",
        score="8/10",
        score_color="#10b981"
    )
    
    # Example: Download buttons
    st.markdown("### Example Downloads")
    example_data = {"title": "Example", "items": ["Item 1", "Item 2"], "score": 85}
    render_download_buttons(example_data, "example_export", formats=["md", "docx", "pdf", "json"])
    
    # ==========================================
    # END APP CONTENT
    # ==========================================
    
    # 6. Feedback widget (bottom right)
    if USING_SHARED:
        render_feedback_widget(APP_KEY, submit_feedback)

# ============================================
# MAIN ENTRY POINT
# ============================================

st.set_page_config(page_title=APP_NAME, page_icon=APP_ICON, layout="wide")
init_session()
check_url_auth()

if not st.session_state.authenticated:
    render_auth()
else:
    render_main_app()
